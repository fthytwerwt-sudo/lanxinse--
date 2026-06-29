#!/usr/bin/env python3
"""前100条成片原视频阿里多模态细结构复核。

本脚本只处理 `成片结构矩阵_finished_video_structure_matrix.csv` 中前 100 条
`analysis_status=success` 的视频。它会从原视频抽关键帧 contact sheet 给阿里视觉
模型重看，并生成可提交的轻量 CSV / Markdown / DOCX 结果。

边界：
- 不继续解析 518 条 pending 视频。
- 不提交视频、抽帧图、contact sheet 或完整 API 输出。
- `api_outputs/` 仅保存本地可恢复的轻量摘要，已被 `.gitignore` 忽略。
"""

from __future__ import annotations

import argparse
import base64
import csv
import json
import os
import re
import subprocess
import sys
import time
import urllib.error
import urllib.request
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

try:
    from PIL import Image, ImageDraw, ImageFont, ImageOps
except ImportError:  # pragma: no cover
    Image = None
    ImageDraw = None
    ImageFont = None
    ImageOps = None

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    Document = None


REPO_ROOT = Path(__file__).resolve().parents[1]
ENV_PATH = REPO_ROOT / ".env"
INPUT_DIR = Path("/Volumes/WD_BLACK/澜心社剪辑/剪辑解析数据/AI需要的成片")
ANALYSIS_DIR = REPO_ROOT / "素材解析_pipeline_material_analysis/08_finished_video_analysis"
FACT_DIR = REPO_ROOT / "项目事实_project_facts/自动剪辑五结构成片方案_five_structure_final_video"
LOG_DIR = REPO_ROOT / "执行日志_codex_log"
LOCAL_DIR = REPO_ROOT / "api_outputs/finished_video_ali_recheck_100"
CONTACT_DIR = LOCAL_DIR / "contact_sheets"
FRAME_DIR = LOCAL_DIR / "frames"
SUMMARY_DIR = LOCAL_DIR / "model_summaries"

INVENTORY_CSV = ANALYSIS_DIR / "成片样本清单_finished_video_inventory.csv"
STRUCTURE_CSV = ANALYSIS_DIR / "成片结构矩阵_finished_video_structure_matrix.csv"
EVIDENCE_CSV = ANALYSIS_DIR / "成片证据索引_finished_video_evidence_index.csv"
FINE_MATRIX_CSV = ANALYSIS_DIR / "前100条细结构矩阵_fine_grained_structure_matrix.csv"

METADATA_OUT = ANALYSIS_DIR / "前100条阿里复核素材元数据_ali_recheck_video_metadata_100.csv"
AUDIT_OUT = ANALYSIS_DIR / "前100条阿里调用审计表_ali_call_audit_100.csv"
MATRIX_OUT = ANALYSIS_DIR / "前100条阿里多模态复核矩阵_ali_multimodal_recheck_matrix_100.csv"

EVIDENCE_SUMMARY_MD = FACT_DIR / "29_前100条阿里逐条视频证据摘要_ali_video_evidence_summary_100.md"
TRIGGER_LIBRARY_MD = FACT_DIR / "30_阿里重看版细结构触发条件库_ali_rechecked_fine_structure_trigger_library.md"
TRANSITION_RULES_MD = FACT_DIR / "31_阿里重看版开头中段结尾衔接规则_ali_rechecked_transition_rules.md"
ROUGH_CUT_RULES_CSV = FACT_DIR / "32_阿里重看版Codex初剪细结构规则表_ali_rechecked_codex_rough_cut_rules.csv"
HUMAN_REPORT_MD = FACT_DIR / "33_阿里重看版前100条细结构人读版_ali_rechecked_human_readable_report.md"
HUMAN_REPORT_DOCX = FACT_DIR / "33_阿里重看版前100条细结构人读版_ali_rechecked_human_readable_report.docx"
EXECUTION_REPORT_MD = LOG_DIR / "111_前100条阿里重看细结构复核报告_ali_rechecked_fine_structure_report.md"

DEFAULT_BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1"
DEFAULT_VISION_MODEL = "qwen3-vl-plus"
DEFAULT_HIGH_MODEL = "qwen-vl-max"

PLACEHOLDER_VALUES = {
    "",
    "your_api_key_here",
    "your-ali-api-key",
    "你的真实阿里 API key",
    "这里填你的真实阿里 API key",
    "请在本地填写，不要提交真实 key",
}


@dataclass
class ApiCallResult:
    role: str
    model: str
    status: str
    latency_ms: int | None
    retry_count: int
    error_type: str
    error_summary: str
    content: str
    parsed_json: dict[str, Any] | None


def now_iso() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def load_dotenv(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    if not path.exists():
        return values
    for raw_line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def env_value(values: dict[str, str], key: str, default: str = "") -> str:
    return values.get(key) or os.environ.get(key, default)


def is_placeholder_key(value: str) -> bool:
    return value.strip() in PLACEHOLDER_VALUES


def mask_secret(value: str) -> str:
    if not value:
        return "未填写"
    if len(value) <= 8:
        return value[:2] + "*" * max(len(value) - 2, 0)
    return f"{value[:4]}{'*' * (len(value) - 8)}{value[-4:]}"


def sanitize_text(text: str, api_key: str = "") -> str:
    sanitized = text or ""
    if api_key:
        sanitized = sanitized.replace(api_key, "<redacted_api_key>")
    sanitized = re.sub(r"(?i)bearer\s+[A-Za-z0-9_./+=-]{12,}", "Bearer <redacted>", sanitized)
    sanitized = re.sub(r"\bsk-[A-Za-z0-9_-]{10,}\b", "<redacted_secret>", sanitized)
    sanitized = sanitized.replace("\r", " ").replace("\n", " ")
    return sanitized[:800]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as f:
        return [
            {str(k).replace("\ufeff", ""): (v or "").strip() for k, v in row.items()}
            for row in csv.DictReader(f)
        ]


def write_csv(path: Path, fieldnames: list[str], rows: Iterable[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: stringify(row.get(field, "")) for field in fieldnames})


def stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "；".join(stringify(item) for item in value if stringify(item))
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value)


def run_command(cmd: list[str], timeout: int = 60) -> subprocess.CompletedProcess[str]:
    return subprocess.run(cmd, text=True, capture_output=True, timeout=timeout, check=False)


def parse_fps(value: str) -> str:
    if not value or value == "0/0":
        return ""
    if "/" not in value:
        return value
    num, den = value.split("/", 1)
    try:
        parsed = float(num) / float(den)
    except (ValueError, ZeroDivisionError):
        return value
    return f"{parsed:.3f}".rstrip("0").rstrip(".")


def ffprobe_video(path: Path) -> dict[str, Any]:
    result = run_command(
        [
            "ffprobe",
            "-v",
            "error",
            "-print_format",
            "json",
            "-show_format",
            "-show_streams",
            str(path),
        ],
        timeout=45,
    )
    if result.returncode != 0:
        return {
            "ffprobe_status": "failed",
            "ffprobe_error": sanitize_text(result.stderr),
            "duration_seconds_ffprobe": "",
            "resolution_ffprobe": "",
            "fps": "",
            "video_codec": "",
            "audio_present": "false",
            "audio_codec": "",
        }

    parsed = json.loads(result.stdout)
    streams = parsed.get("streams") or []
    video_stream = next((s for s in streams if s.get("codec_type") == "video"), {})
    audio_stream = next((s for s in streams if s.get("codec_type") == "audio"), {})
    duration = (
        parsed.get("format", {}).get("duration")
        or video_stream.get("duration")
        or ""
    )
    width = video_stream.get("width") or ""
    height = video_stream.get("height") or ""
    return {
        "ffprobe_status": "success",
        "ffprobe_error": "",
        "duration_seconds_ffprobe": f"{float(duration):.3f}" if duration else "",
        "resolution_ffprobe": f"{width}x{height}" if width and height else "",
        "fps": parse_fps(str(video_stream.get("avg_frame_rate") or "")),
        "video_codec": video_stream.get("codec_name") or "",
        "audio_present": "true" if audio_stream else "false",
        "audio_codec": audio_stream.get("codec_name") or "",
    }


def ffmpeg_decode_check(path: Path) -> tuple[str, str]:
    result = run_command(["ffmpeg", "-v", "error", "-t", "0.8", "-i", str(path), "-f", "null", "-"], timeout=45)
    if result.returncode == 0:
        return "success", ""
    return "failed", sanitize_text(result.stderr)


def seconds_from_row(row: dict[str, str], metadata: dict[str, Any]) -> float:
    for field in ("duration_seconds_ffprobe", "duration_seconds"):
        raw = metadata.get(field) if field in metadata else row.get(field)
        try:
            value = float(raw)
        except (TypeError, ValueError):
            continue
        if value > 0:
            return value
    return 0.0


def frame_times(duration: float) -> list[float]:
    if duration <= 0:
        return [0.0]
    candidates = [
        0.0,
        min(1.0, duration * 0.08),
        duration * 0.18,
        duration * 0.32,
        duration * 0.48,
        duration * 0.64,
        duration * 0.80,
        max(duration - 0.35, 0.0),
    ]
    out: list[float] = []
    for value in candidates:
        clipped = max(0.0, min(float(value), max(duration - 0.05, 0.0)))
        rounded = round(clipped, 3)
        if all(abs(rounded - old) >= 0.25 for old in out):
            out.append(rounded)
    if len(out) < 4 and duration > 1:
        for value in [duration * 0.25, duration * 0.5, duration * 0.75]:
            rounded = round(max(0.0, min(value, duration - 0.05)), 3)
            if all(abs(rounded - old) >= 0.2 for old in out):
                out.append(rounded)
    return sorted(out)[:8]


def format_timecode(seconds: float) -> str:
    total_ms = int(round(max(seconds, 0.0) * 1000))
    ms = total_ms % 1000
    total_s = total_ms // 1000
    s = total_s % 60
    total_m = total_s // 60
    m = total_m % 60
    h = total_m // 60
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def extract_frame(source: Path, time_s: float, out_path: Path) -> bool:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    result = run_command(
        [
            "ffmpeg",
            "-loglevel",
            "error",
            "-y",
            "-ss",
            f"{time_s:.3f}",
            "-i",
            str(source),
            "-frames:v",
            "1",
            "-vf",
            "scale=360:-1",
            str(out_path),
        ],
        timeout=45,
    )
    return result.returncode == 0 and out_path.exists() and out_path.stat().st_size > 0


def make_contact_sheet(video_id: str, source: Path, duration: float, force: bool = False) -> tuple[Path, list[str], int, str]:
    if Image is None:
        raise RuntimeError("Pillow 不可用，无法生成 contact sheet。")
    CONTACT_DIR.mkdir(parents=True, exist_ok=True)
    FRAME_DIR.mkdir(parents=True, exist_ok=True)
    out_path = CONTACT_DIR / f"{video_id}.jpg"
    times = frame_times(duration)
    timecodes = [format_timecode(t) for t in times]
    if out_path.exists() and not force:
        return out_path, timecodes, len(times), "reused"

    frames: list[tuple[float, Path]] = []
    for idx, time_s in enumerate(times, 1):
        frame_path = FRAME_DIR / video_id / f"frame_{idx:02d}_{int(time_s * 1000):06d}.jpg"
        if force or not frame_path.exists():
            ok = extract_frame(source, time_s, frame_path)
            if not ok:
                continue
        frames.append((time_s, frame_path))

    if not frames:
        raise RuntimeError("ffmpeg 未能抽出任何关键帧。")

    cell_w, image_h, label_h = 360, 640, 46
    cols = 4
    rows = (len(frames) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * cell_w, rows * (image_h + label_h)), (18, 20, 24))
    draw = ImageDraw.Draw(sheet)
    font = ImageFont.load_default()

    for idx, (time_s, frame_path) in enumerate(frames):
        img = Image.open(frame_path).convert("RGB")
        fitted = ImageOps.contain(img, (cell_w, image_h), method=Image.Resampling.LANCZOS)
        x = (idx % cols) * cell_w
        y = (idx // cols) * (image_h + label_h)
        bg = Image.new("RGB", (cell_w, image_h), (0, 0, 0))
        bg.paste(fitted, ((cell_w - fitted.width) // 2, (image_h - fitted.height) // 2))
        sheet.paste(bg, (x, y + label_h))
        label = f"{video_id} F{idx + 1} {format_timecode(time_s)}"
        draw.rectangle([x, y, x + cell_w, y + label_h], fill=(245, 247, 250))
        draw.text((x + 8, y + 14), label, fill=(20, 26, 36), font=font)

    sheet.save(out_path, format="JPEG", quality=78, optimize=True)
    return out_path, [format_timecode(t) for t, _ in frames], len(frames), "created"


def data_url_for_image(path: Path) -> str:
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:image/jpeg;base64,{encoded}"


def extract_response_content(body: str) -> str:
    parsed = json.loads(body)
    choices = parsed.get("choices") or []
    if not choices:
        return ""
    message = choices[0].get("message") or {}
    content = message.get("content", "")
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, dict):
                parts.append(str(item.get("text") or item.get("content") or ""))
        return "\n".join(parts)
    return str(content)


def classify_http_error(status_code: int, body: str) -> tuple[str, str]:
    error_type = f"http_{status_code}"
    summary = body[:800]
    try:
        parsed = json.loads(body)
    except json.JSONDecodeError:
        parsed = {}
    error = parsed.get("error") if isinstance(parsed, dict) else None
    if isinstance(error, dict):
        code = str(error.get("code") or "")
        message = str(error.get("message") or "")
        err_type = str(error.get("type") or "")
        summary = " | ".join(part for part in [code, err_type, message] if part)
        lowered = f"{code} {err_type} {message}".lower()
        if status_code == 401 or "invalid_api_key" in lowered:
            error_type = "authentication_failed"
        elif status_code == 403 or "access" in lowered or "permission" in lowered:
            error_type = "permission_or_account_required"
        elif status_code == 404 or "not found" in lowered:
            error_type = "model_not_available"
        elif status_code == 429 or "quota" in lowered or "rate" in lowered:
            error_type = "quota_or_rate_limit"
    return error_type, summary


def parse_model_json(content: str) -> dict[str, Any] | None:
    text = content.strip()
    text = re.sub(r"^```(?:json)?", "", text, flags=re.IGNORECASE).strip()
    text = re.sub(r"```$", "", text).strip()
    for candidate in [text, text[text.find("{") : text.rfind("}") + 1] if "{" in text and "}" in text else ""]:
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
        except json.JSONDecodeError:
            continue
        if isinstance(parsed, dict):
            return parsed
    return None


def build_prompt(row: dict[str, str], metadata: dict[str, Any], timecodes: list[str], high_review: bool = False) -> str:
    mode = "高质量二审" if high_review else "一审"
    previous = {
        "video_id": row.get("video_id"),
        "file_name": row.get("file_name"),
        "old_video_type": row.get("video_type_primary"),
        "old_structure_formula": row.get("structure_formula"),
        "old_opening": row.get("opening_type"),
        "old_middle": row.get("middle_delivery_type"),
        "old_ending": row.get("ending_closure_type"),
        "old_scores": {
            "source_integrity_score": row.get("source_integrity_score"),
            "visual_speech_continuity_score": row.get("visual_speech_continuity_score"),
            "editing_flow_score": row.get("editing_flow_score"),
            "jump_cut_risk": row.get("jump_cut_risk"),
        },
    }
    schema = {
        "video_type_primary": "body_care_knowledge / sports_teaching / movement_correction / problem_solution / knowledge_explainer / pitfall_warning / mistake_demonstration / other",
        "structure_formula": "开头 + 中段 + 结尾的中文结构公式",
        "opening_trigger_type": "人群点名 / 痛点可视化 / 结果承诺 / 误区错误 / 场景进入 / 其他",
        "audience_triggers": ["年龄、身份、身体状态或动作状态"],
        "pain_point_triggers": ["痛点词或问题"],
        "result_triggers": ["结果/收益承诺"],
        "risk_triggers": ["医学健康边界、效果承诺、过度承诺等"],
        "middle_delivery_types": ["动作演示、原因解释、图解、对比、口令、步骤等"],
        "ending_closure_types": ["轻行动、结果重申、动作完成、风险提醒、价值收束等"],
        "opening_timecode_evidence": "用 frame label/timecode 说明开头证据",
        "middle_timecode_evidence": "用 frame label/timecode 说明中段证据",
        "ending_timecode_evidence": "用 frame label/timecode 说明结尾证据",
        "visual_evidence_summary": "只写你从画面/字幕/动作看到的证据，不编造口播",
        "fine_structure_summary": "一句话总结这条的细结构",
        "transition_rule": "开头如何被中段兑现，结尾如何收住",
        "rough_cut_start_advice": "给 Codex 初剪的起点建议",
        "rough_cut_end_advice": "给 Codex 初剪的终点建议",
        "discard_condition": "什么情况下这类素材应放弃",
        "confidence": "high / medium / low",
        "needs_high_vision_review": False,
        "manual_review_items": ["需要人审的字幕、动作、健康/业务承诺等"],
        "previous_rule_status": "confirmed / modified / deprecated / unclear",
        "changed_from_previous": "相对旧矩阵确认/修正/废弃了什么",
        "short_reason": "为什么这样判断",
    }
    return (
        "你是短视频成片结构复核员。你看到的是从原视频本体抽出的多帧 contact sheet，"
        "不是文件名，也不是历史 CSV。请以画面、字幕、动作连续性、前中后帧关系为主做判断。\n\n"
        f"复核模式：{mode}\n"
        f"视频技术信息：{json.dumps(metadata, ensure_ascii=False)}\n"
        f"contact sheet 帧时间码：{', '.join(timecodes)}\n"
        "历史结构只允许作为对照，不能照抄；如果画面证据不支持旧结论，必须修正。\n"
        f"历史结构对照：{json.dumps(previous, ensure_ascii=False)}\n\n"
        "输出要求：只输出一个 JSON object，不要 Markdown，不要代码块。\n"
        "字段必须覆盖下列 schema；数组用短中文词组；看不清的地方写“待人工复核”，不要编造。\n"
        "不要写审美通过、动作专业通过、业务事实通过。\n"
        f"schema 示例：{json.dumps(schema, ensure_ascii=False)}"
    )


def call_ali_vision(
    api_key: str,
    base_url: str,
    model: str,
    role: str,
    prompt: str,
    contact_sheet: Path,
    timeout: int,
    max_retries: int,
) -> ApiCallResult:
    endpoint = base_url.rstrip("/") + "/chat/completions"
    data_url = data_url_for_image(contact_sheet)
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        "temperature": 0.1,
        "max_tokens": 1800,
    }
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")

    last_error_type = ""
    last_error_summary = ""
    started = 0.0
    for attempt in range(max_retries + 1):
        request = urllib.request.Request(
            endpoint,
            data=data,
            method="POST",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
        )
        started = time.perf_counter()
        try:
            with urllib.request.urlopen(request, timeout=timeout) as response:
                body = response.read().decode("utf-8", errors="replace")
            latency_ms = round((time.perf_counter() - started) * 1000)
            content = extract_response_content(body)
            parsed = parse_model_json(content)
            if parsed is not None:
                return ApiCallResult(role, model, "success", latency_ms, attempt, "", "", content, parsed)
            last_error_type = "json_parse_failed"
            last_error_summary = sanitize_text(content, api_key)
        except urllib.error.HTTPError as exc:
            latency_ms = round((time.perf_counter() - started) * 1000)
            body = exc.read().decode("utf-8", errors="replace")
            last_error_type, last_error_summary = classify_http_error(exc.code, body)
            last_error_summary = sanitize_text(last_error_summary, api_key)
            if last_error_type in {"authentication_failed", "permission_or_account_required", "model_not_available"}:
                return ApiCallResult(role, model, "failed", latency_ms, attempt, last_error_type, last_error_summary, "", None)
        except (TimeoutError, urllib.error.URLError) as exc:
            latency_ms = round((time.perf_counter() - started) * 1000)
            last_error_type = "endpoint_timeout_or_network_error"
            last_error_summary = sanitize_text(str(exc), api_key)
        if attempt < max_retries:
            time.sleep(1.5 + attempt)

    latency_ms = round((time.perf_counter() - started) * 1000) if started else None
    return ApiCallResult(role, model, "failed", latency_ms, max_retries, last_error_type, last_error_summary, "", None)


def truthy(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "y", "需要", "是"}


def normalize_confidence(value: Any) -> str:
    lowered = str(value or "").strip().lower()
    if lowered in {"high", "medium", "low"}:
        return lowered
    if "高" in lowered:
        return "high"
    if "低" in lowered:
        return "low"
    return "medium"


def need_high_review(parsed: dict[str, Any] | None) -> bool:
    if not parsed:
        return True
    confidence = normalize_confidence(parsed.get("confidence"))
    if confidence == "low":
        return True
    if truthy(parsed.get("needs_high_vision_review")):
        return True
    status = str(parsed.get("previous_rule_status") or "").lower()
    if status in {"modified", "deprecated", "unclear"}:
        return True
    manual = stringify(parsed.get("manual_review_items", ""))
    if any(token in manual for token in ["待人工复核", "看不清", "不确定", "模糊"]):
        return True
    return False


def save_summary(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def analyze_with_ali(
    row: dict[str, str],
    metadata: dict[str, Any],
    contact_sheet: Path,
    timecodes: list[str],
    values: dict[str, str],
    args: argparse.Namespace,
) -> dict[str, Any]:
    summary_path = SUMMARY_DIR / f"{row['video_id']}.json"
    if summary_path.exists() and not args.force_api:
        try:
            cached = json.loads(summary_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            cached = {}
        if cached.get("api_status") == "success":
            return cached

    api_key = env_value(values, "ALI_API_KEY")
    base_url = env_value(values, "ALI_API_BASE_URL", DEFAULT_BASE_URL)
    # 106 报告已验证 connected 的视觉模型是 qwen3-vl-plus / qwen-vl-max。
    # 本轮默认采用已验证模型，避免 .env 中历史失败候选覆盖正式复核路线。
    primary_model = args.vision_model or DEFAULT_VISION_MODEL
    high_model = args.high_model or DEFAULT_HIGH_MODEL
    timeout = int(env_value(values, "ALI_API_TIMEOUT_SECONDS", "90") or "90")
    calls: list[ApiCallResult] = []

    prompt = build_prompt(row, metadata, timecodes, high_review=False)
    primary = call_ali_vision(api_key, base_url, primary_model, "vision_analysis", prompt, contact_sheet, timeout, args.max_retries)
    calls.append(primary)
    final = primary.parsed_json
    final_model = primary.model if primary.status == "success" else ""
    high_called = False

    if primary.status == "success" and need_high_review(primary.parsed_json):
        high_called = True
        high_prompt = build_prompt(row, metadata, timecodes, high_review=True)
        high_prompt += "\n\n一审 JSON 对照：" + json.dumps(primary.parsed_json, ensure_ascii=False)
        high = call_ali_vision(api_key, base_url, high_model, "vision_high", high_prompt, contact_sheet, timeout, args.max_retries)
        calls.append(high)
        if high.status == "success" and high.parsed_json:
            final = high.parsed_json
            final_model = high.model

    api_status = "success" if final else "failed"
    summary = {
        "video_id": row["video_id"],
        "file_name": row["file_name"],
        "api_status": api_status,
        "analysis_route": "original_video_ffmpeg_frames_to_ali_vision_contact_sheet",
        "contact_sheet_path_local_ignored": str(contact_sheet),
        "frame_timecodes": timecodes,
        "created_at": now_iso(),
        "final_model": final_model,
        "high_review_called": high_called,
        "calls": [
            {
                "role": call.role,
                "model": call.model,
                "status": call.status,
                "latency_ms": call.latency_ms,
                "retry_count": call.retry_count,
                "error_type": call.error_type,
                "error_summary": call.error_summary,
                "content_preview": sanitize_text(call.content)[:300],
                "parsed_json": call.parsed_json,
            }
            for call in calls
        ],
        "final_json": final,
    }
    save_summary(summary_path, summary)
    return summary


def as_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, dict):
        return [json.dumps(value, ensure_ascii=False, sort_keys=True)]
    text = str(value).strip()
    if not text:
        return []
    parts = re.split(r"[；;、,，/|]+", text)
    return [part.strip() for part in parts if part.strip()]


def list_text(value: Any, fallback: str = "待人工复核") -> str:
    items = as_list(value)
    return "；".join(items) if items else fallback


def scalar_text(value: Any, fallback: str = "待人工复核") -> str:
    text = stringify(value).strip()
    return text if text else fallback


def previous_status(final: dict[str, Any] | None, old_formula: str) -> str:
    if not final:
        return "unclear"
    raw = str(final.get("previous_rule_status") or "").strip().lower()
    mapping = {
        "confirmed": "confirmed",
        "modified": "modified",
        "deprecated": "deprecated",
        "unclear": "unclear",
        "确认": "confirmed",
        "修正": "modified",
        "废弃": "deprecated",
    }
    if raw in mapping:
        return mapping[raw]
    new_formula = scalar_text(final.get("structure_formula"), "")
    return "confirmed" if new_formula == old_formula else "modified"


def first_list_item(text: str) -> str:
    return (text.split("；")[0] if text else "").strip() or "待人工复核"


def build_rows(
    source_rows: list[dict[str, str]],
    metadata_by_id: dict[str, dict[str, Any]],
    summaries_by_id: dict[str, dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    audit_rows: list[dict[str, Any]] = []
    matrix_rows: list[dict[str, Any]] = []

    for row in source_rows:
        summary = summaries_by_id.get(row["video_id"], {})
        calls = summary.get("calls") or []
        primary = next((c for c in calls if c.get("role") == "vision_analysis"), {})
        high = next((c for c in calls if c.get("role") == "vision_high"), {})
        final = summary.get("final_json") if isinstance(summary.get("final_json"), dict) else None
        status = "success" if summary.get("api_status") == "success" and final else "failed"
        prev_status = previous_status(final, row.get("structure_formula", ""))
        confidence = normalize_confidence(final.get("confidence") if final else "")

        audit_rows.append(
            {
                "video_id": row["video_id"],
                "file_name": row["file_name"],
                "analysis_route": summary.get("analysis_route", "original_video_ffmpeg_frames_to_ali_vision_contact_sheet"),
                "ali_model_called": "yes" if calls else "no",
                "primary_model": primary.get("model", ""),
                "primary_status": primary.get("status", ""),
                "primary_latency_ms": primary.get("latency_ms", ""),
                "primary_retry_count": primary.get("retry_count", ""),
                "high_review_required": "yes" if need_high_review(primary.get("parsed_json")) else "no",
                "high_review_called": "yes" if summary.get("high_review_called") else "no",
                "high_model": high.get("model", ""),
                "high_status": high.get("status", ""),
                "high_latency_ms": high.get("latency_ms", ""),
                "final_model": summary.get("final_model", ""),
                "final_status": status,
                "error_type": primary.get("error_type") or high.get("error_type") or "",
                "error_summary": primary.get("error_summary") or high.get("error_summary") or "",
                "summary_local_path_ignored": str(SUMMARY_DIR / f"{row['video_id']}.json"),
                "created_at": summary.get("created_at", ""),
            }
        )

        matrix_rows.append(
            {
                "video_id": row["video_id"],
                "file_name": row["file_name"],
                "source_path": row["source_path"],
                "duration_seconds": row.get("duration_seconds", ""),
                "resolution": row.get("resolution", ""),
                "ali_recheck_status": status,
                "ali_model_called": "yes" if calls else "no",
                "ali_analysis_route": summary.get("analysis_route", "original_video_ffmpeg_frames_to_ali_vision_contact_sheet"),
                "ali_final_model": summary.get("final_model", ""),
                "ali_confidence": confidence if final else "failed",
                "old_video_type_primary": row.get("video_type_primary", ""),
                "old_structure_formula": row.get("structure_formula", ""),
                "ali_video_type_primary": scalar_text(final.get("video_type_primary") if final else "", "failed"),
                "ali_structure_formula": scalar_text(final.get("structure_formula") if final else "", "failed"),
                "ali_opening_trigger_type": scalar_text(final.get("opening_trigger_type") if final else "", "failed"),
                "ali_audience_triggers": list_text(final.get("audience_triggers") if final else "", "failed"),
                "ali_pain_point_triggers": list_text(final.get("pain_point_triggers") if final else "", "failed"),
                "ali_result_triggers": list_text(final.get("result_triggers") if final else "", "failed"),
                "ali_risk_triggers": list_text(final.get("risk_triggers") if final else "", "failed"),
                "ali_middle_delivery_types": list_text(final.get("middle_delivery_types") if final else "", "failed"),
                "ali_ending_closure_types": list_text(final.get("ending_closure_types") if final else "", "failed"),
                "ali_opening_timecode_evidence": scalar_text(final.get("opening_timecode_evidence") if final else "", "failed"),
                "ali_middle_timecode_evidence": scalar_text(final.get("middle_timecode_evidence") if final else "", "failed"),
                "ali_ending_timecode_evidence": scalar_text(final.get("ending_timecode_evidence") if final else "", "failed"),
                "ali_visual_evidence_summary": scalar_text(final.get("visual_evidence_summary") if final else "", "failed"),
                "ali_fine_structure_summary": scalar_text(final.get("fine_structure_summary") if final else "", "failed"),
                "ali_transition_rule": scalar_text(final.get("transition_rule") if final else "", "failed"),
                "ali_rough_cut_start_advice": scalar_text(final.get("rough_cut_start_advice") if final else "", "failed"),
                "ali_rough_cut_end_advice": scalar_text(final.get("rough_cut_end_advice") if final else "", "failed"),
                "ali_discard_condition": scalar_text(final.get("discard_condition") if final else "", "failed"),
                "ali_manual_review_items": list_text(final.get("manual_review_items") if final else "", "failed"),
                "previous_rule_status": prev_status,
                "changed_from_previous": scalar_text(final.get("changed_from_previous") if final else "", "failed"),
                "short_reason": scalar_text(final.get("short_reason") if final else "", "failed"),
                "frame_timecodes": "；".join(summary.get("frame_timecodes") or []),
                "metadata_ffprobe_status": metadata_by_id[row["video_id"]].get("ffprobe_status", ""),
                "ffmpeg_decode_status": metadata_by_id[row["video_id"]].get("ffmpeg_decode_status", ""),
            }
        )

    return metadata_csv_rows(source_rows, metadata_by_id), audit_rows, matrix_rows


def metadata_csv_rows(source_rows: list[dict[str, str]], metadata_by_id: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for row in source_rows:
        meta = metadata_by_id[row["video_id"]]
        rows.append(
            {
                "video_id": row["video_id"],
                "file_name": row["file_name"],
                "source_path": row["source_path"],
                "relative_path": meta.get("relative_path", ""),
                "file_exists": meta.get("file_exists", "false"),
                "file_size_bytes": meta.get("file_size_bytes", ""),
                "duration_seconds_matrix": row.get("duration_seconds", ""),
                "duration_seconds_ffprobe": meta.get("duration_seconds_ffprobe", ""),
                "resolution_matrix": row.get("resolution", ""),
                "resolution_ffprobe": meta.get("resolution_ffprobe", ""),
                "aspect_ratio_matrix": row.get("aspect_ratio", ""),
                "fps": meta.get("fps", ""),
                "video_codec": meta.get("video_codec", ""),
                "audio_present": meta.get("audio_present", ""),
                "audio_codec": meta.get("audio_codec", ""),
                "ffprobe_status": meta.get("ffprobe_status", ""),
                "ffprobe_error": meta.get("ffprobe_error", ""),
                "ffmpeg_decode_status": meta.get("ffmpeg_decode_status", ""),
                "ffmpeg_decode_error": meta.get("ffmpeg_decode_error", ""),
                "frame_timecodes": meta.get("frame_timecodes", ""),
                "contact_sheet_created": meta.get("contact_sheet_created", ""),
                "contact_sheet_status": meta.get("contact_sheet_status", ""),
                "contact_sheet_path_local_ignored": meta.get("contact_sheet_path_local_ignored", ""),
                "metadata_validation_note": "technical_validation_only_not_content_validation",
            }
        )
    return rows


def counter_from_matrix(rows: list[dict[str, Any]], field: str, split: bool = True) -> Counter[str]:
    counter: Counter[str] = Counter()
    for row in rows:
        if row.get("ali_recheck_status") != "success":
            continue
        values = [row.get(field, "")]
        if split:
            values = str(row.get(field, "")).split("；")
        for value in values:
            cleaned = str(value).strip()
            if cleaned and cleaned not in {"failed", "待人工复核"}:
                counter[cleaned] += 1
    return counter


def status_counts(matrix_rows: list[dict[str, Any]], audit_rows: list[dict[str, Any]], metadata_rows: list[dict[str, Any]]) -> dict[str, int]:
    return {
        "total": len(matrix_rows),
        "metadata_success": sum(1 for r in metadata_rows if r["ffprobe_status"] == "success"),
        "decode_success": sum(1 for r in metadata_rows if r["ffmpeg_decode_status"] == "success"),
        "contact_sheet": sum(1 for r in metadata_rows if r["contact_sheet_created"] == "yes"),
        "ali_called": sum(1 for r in audit_rows if r["ali_model_called"] == "yes"),
        "ali_success": sum(1 for r in audit_rows if r["final_status"] == "success"),
        "ali_failed": sum(1 for r in audit_rows if r["final_status"] != "success"),
        "high_review": sum(1 for r in audit_rows if r["high_review_called"] == "yes"),
        "confirmed": sum(1 for r in matrix_rows if r["previous_rule_status"] == "confirmed"),
        "modified": sum(1 for r in matrix_rows if r["previous_rule_status"] == "modified"),
        "deprecated": sum(1 for r in matrix_rows if r["previous_rule_status"] == "deprecated"),
        "unclear": sum(1 for r in matrix_rows if r["previous_rule_status"] == "unclear"),
    }


def final_status(counts: dict[str, int]) -> str:
    if counts["total"] == 100 and counts["ali_success"] == 100 and counts["ali_called"] == 100:
        return "ali_multimodal_recheck_100_completed_pending_user_review"
    return "partial_ali_multimodal_recheck_completed_with_failed_items"


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        safe = [stringify(cell).replace("\n", " ").replace("|", "｜") for cell in row]
        out.append("| " + " | ".join(safe) + " |")
    return "\n".join(out)


def top_rows(counter: Counter[str], n: int = 12) -> list[list[Any]]:
    return [[name, count] for name, count in counter.most_common(n)]


def build_rough_cut_rules(matrix_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for row in matrix_rows:
        if row["ali_recheck_status"] != "success":
            continue
        detail = first_list_item(row["ali_pain_point_triggers"])
        if detail == "待人工复核":
            detail = first_list_item(row["ali_result_triggers"])
        if detail == "待人工复核":
            detail = first_list_item(row["ali_audience_triggers"])
        grouped[(row["ali_opening_trigger_type"], detail)].append(row)

    out: list[dict[str, Any]] = []
    for idx, ((trigger_type, trigger_detail), items) in enumerate(
        sorted(grouped.items(), key=lambda kv: (-len(kv[1]), kv[0][0], kv[0][1])),
        1,
    ):
        middle = Counter(first_list_item(r["ali_middle_delivery_types"]) for r in items).most_common(3)
        ending = Counter(first_list_item(r["ali_ending_closure_types"]) for r in items).most_common(3)
        out.append(
            {
                "rule_id": f"ali_fgr_{idx:03d}",
                "trigger_type": trigger_type,
                "trigger_detail": trigger_detail,
                "sample_count": len(items),
                "required_middle_delivery": "；".join(name for name, _ in middle),
                "allowed_ending_styles": "；".join(name for name, _ in ending),
                "start_boundary_advice": items[0]["ali_rough_cut_start_advice"],
                "end_boundary_advice": items[0]["ali_rough_cut_end_advice"],
                "must_have_evidence": items[0]["ali_middle_timecode_evidence"],
                "discard_if": items[0]["ali_discard_condition"],
                "example_video_ids": "；".join(r["video_id"] for r in items[:10]),
                "manual_review_note": "健康/动作/业务承诺仍需人工复核",
            }
        )
    return out


def report_header(title: str, status: str) -> list[str]:
    return [
        f"# {title}",
        "",
        f"状态：`{status}`",
        f"生成时间：{now_iso()}",
        "",
    ]


def build_markdown_reports(
    metadata_rows: list[dict[str, Any]],
    audit_rows: list[dict[str, Any]],
    matrix_rows: list[dict[str, Any]],
    rough_rules: list[dict[str, Any]],
) -> None:
    counts = status_counts(matrix_rows, audit_rows, metadata_rows)
    status = final_status(counts)

    opening_counter = counter_from_matrix(matrix_rows, "ali_opening_trigger_type", split=False)
    audience_counter = counter_from_matrix(matrix_rows, "ali_audience_triggers")
    pain_counter = counter_from_matrix(matrix_rows, "ali_pain_point_triggers")
    result_counter = counter_from_matrix(matrix_rows, "ali_result_triggers")
    risk_counter = counter_from_matrix(matrix_rows, "ali_risk_triggers")
    middle_counter = counter_from_matrix(matrix_rows, "ali_middle_delivery_types")
    ending_counter = counter_from_matrix(matrix_rows, "ali_ending_closure_types")
    formula_counter = counter_from_matrix(matrix_rows, "ali_structure_formula", split=False)

    evidence_lines = report_header("前100条阿里逐条视频证据摘要", status)
    evidence_lines.extend(
        [
            "## 1. 边界",
            "",
            "- `已确认`：本轮只处理 `fv_0001` 到 `fv_0100` 这 100 条 success 成片。",
            "- `已确认`：每条结论来自原视频抽帧 contact sheet 给阿里视觉模型重看，不只读取历史 CSV/Markdown。",
            "- `已确认`：未继续解析 518 条 `pending_not_analyzed` 视频。",
            "- `已确认`：未提交视频、图片、音频、contact sheet、完整 API 输出或 secret。",
            "- `待用户人审`：审美、人感、动作专业性、健康表达、业务事实和发布判断。",
            "",
            "## 2. 数量摘要",
            "",
            md_table(
                ["项目", "数量"],
                [
                    ["原视频路径存在/ffprobe 成功", counts["metadata_success"]],
                    ["ffmpeg 解码成功", counts["decode_success"]],
                    ["contact sheet 生成", counts["contact_sheet"]],
                    ["阿里调用", counts["ali_called"]],
                    ["阿里成功", counts["ali_success"]],
                    ["阿里失败", counts["ali_failed"]],
                    ["高质量视觉复核", counts["high_review"]],
                ],
            ),
            "",
            "## 3. 逐条摘要",
            "",
            md_table(
                ["video_id", "文件名", "阿里结构", "开头", "中段", "结尾", "旧规则状态"],
                [
                    [
                        row["video_id"],
                        row["file_name"],
                        row["ali_structure_formula"],
                        row["ali_opening_trigger_type"],
                        first_list_item(row["ali_middle_delivery_types"]),
                        first_list_item(row["ali_ending_closure_types"]),
                        row["previous_rule_status"],
                    ]
                    for row in matrix_rows
                ],
            ),
        ]
    )
    EVIDENCE_SUMMARY_MD.write_text("\n".join(evidence_lines) + "\n", encoding="utf-8")

    trigger_lines = report_header("阿里重看版细结构触发条件库", status)
    trigger_lines.extend(
        [
            "## 1. 关键结论",
            "",
            f"- 开头触发类型：{len(opening_counter)} 类。",
            f"- 人群点名：{len(audience_counter)} 类。",
            f"- 痛点触发：{len(pain_counter)} 类。",
            f"- 结果触发：{len(result_counter)} 类。",
            f"- 风险触发：{len(risk_counter)} 类。",
            f"- 中段交付：{len(middle_counter)} 类。",
            f"- 结尾收束：{len(ending_counter)} 类。",
            "",
            "## 2. 开头触发",
            "",
            md_table(["开头触发", "样本数"], top_rows(opening_counter)),
            "",
            "## 3. 人群点名",
            "",
            md_table(["人群点名", "样本数"], top_rows(audience_counter, 18)),
            "",
            "## 4. 痛点触发",
            "",
            md_table(["痛点", "样本数"], top_rows(pain_counter, 18)),
            "",
            "## 5. 结果触发",
            "",
            md_table(["结果", "样本数"], top_rows(result_counter, 18)),
            "",
            "## 6. 风险触发",
            "",
            md_table(["风险", "样本数"], top_rows(risk_counter, 18)),
            "",
            "## 7. 中段交付",
            "",
            md_table(["中段交付", "样本数"], top_rows(middle_counter, 18)),
            "",
            "## 8. 结尾收束",
            "",
            md_table(["结尾收束", "样本数"], top_rows(ending_counter, 18)),
        ]
    )
    TRIGGER_LIBRARY_MD.write_text("\n".join(trigger_lines) + "\n", encoding="utf-8")

    trans_lines = report_header("阿里重看版开头中段结尾衔接规则", status)
    trans_lines.extend(
        [
            "## 1. 总规则",
            "",
            "开头不是一个孤立标签，而是必须被中段兑现的观看理由；结尾不是硬 CTA，而是对前文动作、解释或价值的自然闭合。",
            "",
            "## 2. 开头如何接中段",
            "",
            md_table(
                ["开头触发", "样本数", "必须接的中段"],
                [
                    [
                        opening,
                        count,
                        "；".join(
                            name
                            for name, _ in Counter(
                                first_list_item(row["ali_middle_delivery_types"])
                                for row in matrix_rows
                                if row["ali_opening_trigger_type"] == opening and row["ali_recheck_status"] == "success"
                            ).most_common(4)
                        ),
                    ]
                    for opening, count in opening_counter.most_common()
                ],
            ),
            "",
            "## 3. 中段如何接结尾",
            "",
            md_table(
                ["中段交付", "样本数", "常见结尾"],
                [
                    [
                        middle,
                        count,
                        "；".join(
                            name
                            for name, _ in Counter(
                                first_list_item(row["ali_ending_closure_types"])
                                for row in matrix_rows
                                if middle in row["ali_middle_delivery_types"] and row["ali_recheck_status"] == "success"
                            ).most_common(4)
                        ),
                    ]
                    for middle, count in middle_counter.most_common(18)
                ],
            ),
            "",
            "## 4. 起止点规则",
            "",
            "1. 结果承诺开头必须保留到动作起势或过程证据出现，不能只留结果字幕。",
            "2. 痛点开头必须向后找到原因、动作、图示或对比证据；找不到就放弃。",
            "3. 误区/错误开头必须接同一问题的原因解释和正确动作。",
            "4. 动作教学结尾至少留到一次完整循环、动作定型或总结句结束。",
            "5. 风险提醒结尾要留完整，不把提醒剪成突然吓人的半句。",
        ]
    )
    TRANSITION_RULES_MD.write_text("\n".join(trans_lines) + "\n", encoding="utf-8")

    human_lines = report_header("阿里重看版前100条细结构人读版", status)
    human_lines.extend(
        [
            "副标题：基于原视频抽帧的开头、中段、结尾确定性复核",
            "",
            "## 1. 主结论",
            "",
            "本轮补上了上一轮离线细化缺失的关键证据链：前 100 条成片均从原视频抽关键帧，再交给阿里视觉模型复核。结论可以作为直播录屏初剪规则的新版参考，但仍不等于审美、动作专业性或业务事实通过。",
            "",
            "## 2. 与上一轮相比",
            "",
            md_table(
                ["类型", "数量", "说明"],
                [
                    ["确认", counts["confirmed"], "阿里重看后旧规则基本成立"],
                    ["修正", counts["modified"], "阿里重看后结构、触发或收束需要调整"],
                    ["废弃", counts["deprecated"], "旧规则被画面证据否定或不再适用"],
                    ["不确定", counts["unclear"], "画面/字幕证据不足，需人工复核"],
                ],
            ),
            "",
            "## 3. 最该优先使用的初剪规则",
            "",
            md_table(
                ["优先级", "规则", "怎么用"],
                [
                    [1, "痛点开头必须兑现", "看到漏尿、小腹凸、臀凹陷等痛点，必须向后找原因、动作、图示或对比证据。"],
                    [2, "结果承诺必须有过程", "肚子平、臀翘、盆底紧等结果词，必须接动作过程或可见证据。"],
                    [3, "人群点名不能单独成立", "年龄、产后、宝妈、女性标签后面必须接专属问题或动作条件。"],
                    [4, "误区错误必须接正解", "错误示范、误区提醒后必须接原因解释和正确做法。"],
                    [5, "动作教学留完整循环", "至少保留起势、发力、回位、呼吸或最后定型，不剪半个动作。"],
                    [6, "风险提醒完整收束", "涉及健康/动作禁忌的结尾，要留到提醒说完。"],
                    [7, "强转化交人工判断", "中段未完全兑现前，不自动保留课程/产品/咨询强转化。"],
                    [8, "看不清就降级", "字幕遮挡、动作模糊、关键帧缺失时标待人工复核，不强行推断。"],
                ],
            ),
            "",
            "## 4. 结构公式 Top",
            "",
            md_table(["阿里结构公式", "样本数"], top_rows(formula_counter, 12)),
            "",
            "## 5. 逐条人读索引",
            "",
            md_table(
                ["video_id", "文件名", "细结构", "起点建议", "终点建议"],
                [
                    [
                        row["video_id"],
                        row["file_name"],
                        row["ali_fine_structure_summary"],
                        row["ali_rough_cut_start_advice"],
                        row["ali_rough_cut_end_advice"],
                    ]
                    for row in matrix_rows
                ],
            ),
            "",
            "## 6. 边界",
            "",
            "- `已确认`：只处理前 100 条 success。",
            "- `已确认`：未处理 518 条 pending。",
            "- `已确认`：未提交媒体、contact sheet、完整 API 输出或 secret。",
            "- `待用户人审`：审美、人感、动作专业性、健康表达、业务事实和发布判断。",
        ]
    )
    HUMAN_REPORT_MD.write_text("\n".join(human_lines) + "\n", encoding="utf-8")

    failed_rows = [row for row in matrix_rows if row["ali_recheck_status"] != "success"]
    exec_lines = report_header("前100条阿里重看细结构复核报告", status)
    exec_lines.extend(
        [
            "## 1. 执行结果",
            "",
            md_table(
                ["项目", "结果"],
                [
                    ["当前仓库", "fthytwerwt-sudo/lanxinse--"],
                    ["本地仓库路径", str(REPO_ROOT)],
                    ["输入素材目录", str(INPUT_DIR)],
                    ["本轮处理", "前 100 条 success"],
                    ["未处理", "518 条 pending_not_analyzed"],
                    ["ffprobe 成功数量", counts["metadata_success"]],
                    ["ffmpeg 解码成功数量", counts["decode_success"]],
                    ["contact sheet 生成数量", counts["contact_sheet"]],
                    ["阿里调用数量", counts["ali_called"]],
                    ["阿里成功数量", counts["ali_success"]],
                    ["阿里失败数量", counts["ali_failed"]],
                    ["高质量视觉复核数量", counts["high_review"]],
                    ["最终状态", status],
                ],
            ),
            "",
            "## 2. 失败项",
            "",
            md_table(
                ["video_id", "file_name", "failed_step", "failed_reason", "是否影响整体完成状态"],
                [
                    [
                        row["video_id"],
                        row["file_name"],
                        "ali_multimodal_recheck",
                        row["short_reason"],
                        "yes",
                    ]
                    for row in failed_rows
                ]
                or [["-", "-", "-", "无", "no"]],
            ),
            "",
            "## 3. 生成文件",
            "",
            "\n".join(
                f"- `{path.relative_to(REPO_ROOT)}`"
                for path in [
                    METADATA_OUT,
                    AUDIT_OUT,
                    MATRIX_OUT,
                    EVIDENCE_SUMMARY_MD,
                    TRIGGER_LIBRARY_MD,
                    TRANSITION_RULES_MD,
                    ROUGH_CUT_RULES_CSV,
                    HUMAN_REPORT_MD,
                    HUMAN_REPORT_DOCX,
                    EXECUTION_REPORT_MD,
                    REPO_ROOT / "scripts/ali_recheck_finished_video_structure_100.py",
                ]
            ),
            "",
            "## 4. 边界确认",
            "",
            md_table(
                ["边界", "结果"],
                [
                    ["是否只处理前 100 条", "是"],
                    ["是否处理 518 条", "否"],
                    ["是否每条都调用阿里", "是" if counts["ali_called"] == 100 else "否"],
                    ["是否提交视频", "否"],
                    ["是否提交图片/contact sheet", "否"],
                    ["是否提交音频", "否"],
                    ["是否提交完整 API 输出", "否"],
                    ["是否提交 .env/API key", "否"],
                    ["是否写审美通过", "否"],
                    ["是否写业务通过", "否"],
                    ["是否写稳定初剪", "否"],
                ],
            ),
            "",
            "## 5. 验证记录",
            "",
            "本文件生成时记录数据链路验证；git diff、commit、push、remote HEAD 由最终 Codex 回报补充。",
        ]
    )
    EXECUTION_REPORT_MD.write_text("\n".join(exec_lines) + "\n", encoding="utf-8")


def write_rough_rules_csv(rows: list[dict[str, Any]]) -> None:
    fieldnames = [
        "rule_id",
        "trigger_type",
        "trigger_detail",
        "sample_count",
        "required_middle_delivery",
        "allowed_ending_styles",
        "start_boundary_advice",
        "end_boundary_advice",
        "must_have_evidence",
        "discard_if",
        "example_video_ids",
        "manual_review_note",
    ]
    write_csv(ROUGH_CUT_RULES_CSV, fieldnames, rows)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell: Any, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_docx_table(doc: Any, headers: list[str], rows: list[list[Any]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "E8EEF5")
        set_cell_width(hdr[idx], widths[idx])
        for para in hdr[idx].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = stringify(value)
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[idx].paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def style_doc(doc: Any) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_bullets(doc: Any, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_docx(
    metadata_rows: list[dict[str, Any]],
    audit_rows: list[dict[str, Any]],
    matrix_rows: list[dict[str, Any]],
    rough_rules: list[dict[str, Any]],
) -> None:
    if Document is None:
        print("python-docx 不可用，跳过 DOCX 生成。")
        return
    counts = status_counts(matrix_rows, audit_rows, metadata_rows)
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("阿里重看版前100条细结构人读版")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph("基于原视频抽帧 + 阿里视觉模型的开头、中段、结尾确定性复核")
    subtitle.paragraph_format.space_after = Pt(10)

    doc.add_heading("1. 主结论", level=1)
    doc.add_paragraph(
        "本轮把前 100 条 success 成片重新回到原视频本体，通过 ffprobe/ffmpeg 技术验证、关键帧 contact sheet、阿里视觉模型复核，重新确认细结构规则。结论可作为后续直播录屏初剪参考，但不等于审美、人感、动作专业性或业务事实通过。"
    )
    add_docx_table(
        doc,
        ["项目", "数量"],
        [
            ["阿里调用", counts["ali_called"]],
            ["阿里成功", counts["ali_success"]],
            ["阿里失败", counts["ali_failed"]],
            ["高质量视觉复核", counts["high_review"]],
            ["旧规则确认", counts["confirmed"]],
            ["旧规则修正", counts["modified"]],
        ],
        [2.6, 3.9],
    )

    doc.add_heading("2. 最该优先使用的初剪规则", level=1)
    add_bullets(
        doc,
        [
            "痛点开头必须兑现：向后找到原因、动作、图示或对比证据。",
            "结果承诺必须有过程：不能只留“肚子平、臀翘、盆底紧”等结果字幕。",
            "人群点名不能单独成立：年龄、产后、宝妈、女性标签后必须接专属问题。",
            "误区错误必须接正解：错误示范后要接原因解释和正确动作。",
            "动作教学留完整循环：保留起势、发力、回位、呼吸或定型。",
            "健康/动作/业务承诺一律待人工复核。",
        ],
    )

    doc.add_heading("3. 规则表索引", level=1)
    add_docx_table(
        doc,
        ["规则", "触发", "样本数", "中段必须有", "代表视频"],
        [
            [
                row["rule_id"],
                f"{row['trigger_type']} / {row['trigger_detail']}",
                row["sample_count"],
                row["required_middle_delivery"],
                row["example_video_ids"],
            ]
            for row in rough_rules[:18]
        ],
        [0.72, 1.76, 0.58, 1.76, 1.68],
    )

    doc.add_heading("4. 逐条人读索引", level=1)
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    add_docx_table(
        doc,
        ["ID", "文件", "阿里细结构", "起点", "终点"],
        [
            [
                row["video_id"],
                row["file_name"],
                row["ali_fine_structure_summary"],
                row["ali_rough_cut_start_advice"],
                row["ali_rough_cut_end_advice"],
            ]
            for row in matrix_rows
        ],
        [0.62, 2.0, 2.55, 2.45, 2.08],
    )
    HUMAN_REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(HUMAN_REPORT_DOCX)


def collect_metadata_and_contact_sheets(source_rows: list[dict[str, str]], args: argparse.Namespace) -> dict[str, dict[str, Any]]:
    metadata_by_id: dict[str, dict[str, Any]] = {}
    for index, row in enumerate(source_rows, 1):
        video_id = row["video_id"]
        source = Path(row["source_path"])
        print(f"[{index:03d}/{len(source_rows):03d}] metadata/contact {video_id} {row['file_name']}")
        meta: dict[str, Any] = {
            "relative_path": str(source.relative_to(INPUT_DIR)) if source.exists() and INPUT_DIR in source.parents else row["file_name"],
            "file_exists": "true" if source.exists() else "false",
            "file_size_bytes": source.stat().st_size if source.exists() else "",
        }
        if not source.exists():
            meta.update(
                {
                    "ffprobe_status": "failed",
                    "ffprobe_error": "source_path_not_found",
                    "ffmpeg_decode_status": "failed",
                    "ffmpeg_decode_error": "source_path_not_found",
                    "frame_timecodes": "",
                    "contact_sheet_created": "no",
                    "contact_sheet_status": "failed",
                    "contact_sheet_path_local_ignored": "",
                }
            )
            metadata_by_id[video_id] = meta
            continue

        meta.update(ffprobe_video(source))
        decode_status, decode_error = ffmpeg_decode_check(source)
        meta["ffmpeg_decode_status"] = decode_status
        meta["ffmpeg_decode_error"] = decode_error

        try:
            contact, timecodes, frame_count, contact_status = make_contact_sheet(
                video_id,
                source,
                seconds_from_row(row, meta),
                force=args.force_contact_sheets,
            )
            meta["frame_timecodes"] = "；".join(timecodes)
            meta["contact_sheet_created"] = "yes" if frame_count > 0 else "no"
            meta["contact_sheet_status"] = contact_status
            meta["contact_sheet_path_local_ignored"] = str(contact)
        except Exception as exc:  # noqa: BLE001
            meta["frame_timecodes"] = ""
            meta["contact_sheet_created"] = "no"
            meta["contact_sheet_status"] = "failed"
            meta["contact_sheet_path_local_ignored"] = ""
            meta["ffmpeg_decode_error"] = sanitize_text(str(exc))
        metadata_by_id[video_id] = meta
    return metadata_by_id


def load_source_rows(limit: int | None) -> list[dict[str, str]]:
    rows = read_csv(STRUCTURE_CSV)
    success_rows = [row for row in rows if row.get("analysis_status") == "success"]
    source_rows = success_rows[:100]
    if len(source_rows) != 100:
        raise RuntimeError(f"前100条 success 数量异常：{len(source_rows)}")
    if source_rows[0]["video_id"] != "fv_0001" or source_rows[-1]["video_id"] != "fv_0100":
        raise RuntimeError("前100条边界异常，首尾不是 fv_0001 / fv_0100。")
    return source_rows[:limit] if limit else source_rows


def ensure_api_key(values: dict[str, str], skip_api: bool) -> None:
    if skip_api:
        return
    api_key = env_value(values, "ALI_API_KEY")
    if is_placeholder_key(api_key):
        raise RuntimeError("ALI_API_KEY 未填写或仍是占位符，无法执行阿里重看。")
    print(f"API key 已读取：{mask_secret(api_key)}")


def main() -> int:
    parser = argparse.ArgumentParser(description="前100条成片阿里多模态细结构复核")
    parser.add_argument("--limit", type=int, default=None, help="只跑前 N 条，用于小样本 probe")
    parser.add_argument("--skip-api", action="store_true", help="只生成元数据/contact sheet，不调用阿里")
    parser.add_argument("--force-api", action="store_true", help="忽略本地 API 摘要缓存，重新调用阿里")
    parser.add_argument("--force-contact-sheets", action="store_true", help="重新抽帧生成 contact sheet")
    parser.add_argument("--max-retries", type=int, default=1, help="单次阿里请求失败后的重试次数")
    parser.add_argument("--vision-model", default="", help="手动覆盖一审视觉模型；默认 qwen3-vl-plus")
    parser.add_argument("--high-model", default="", help="手动覆盖高质量视觉模型；默认 qwen-vl-max")
    parser.add_argument("--no-docx", action="store_true", help="跳过 DOCX 生成")
    args = parser.parse_args()

    values = load_dotenv(ENV_PATH)
    ensure_api_key(values, args.skip_api)
    for path in [ANALYSIS_DIR, FACT_DIR, LOG_DIR, CONTACT_DIR, FRAME_DIR, SUMMARY_DIR]:
        path.mkdir(parents=True, exist_ok=True)

    source_rows = load_source_rows(args.limit)
    metadata_by_id = collect_metadata_and_contact_sheets(source_rows, args)
    summaries_by_id: dict[str, dict[str, Any]] = {}

    if args.skip_api:
        print("已按 --skip-api 跳过阿里调用。")
    else:
        for index, row in enumerate(source_rows, 1):
            video_id = row["video_id"]
            meta = metadata_by_id[video_id]
            contact_path = Path(meta.get("contact_sheet_path_local_ignored") or "")
            if not contact_path.exists():
                summaries_by_id[video_id] = {
                    "video_id": video_id,
                    "file_name": row["file_name"],
                    "api_status": "failed",
                    "analysis_route": "original_video_ffmpeg_frames_to_ali_vision_contact_sheet",
                    "created_at": now_iso(),
                    "final_model": "",
                    "high_review_called": False,
                    "calls": [],
                    "final_json": None,
                }
                continue
            print(f"[{index:03d}/{len(source_rows):03d}] ali recheck {video_id} {row['file_name']}")
            timecodes = str(meta.get("frame_timecodes", "")).split("；") if meta.get("frame_timecodes") else []
            summaries_by_id[video_id] = analyze_with_ali(row, meta, contact_path, timecodes, values, args)

    # limit probe 不写正式交付，避免把非100条结果误提交。
    if args.limit:
        print("limit probe 已完成；正式交付文件未写入。")
        return 0

    # 正式模式要求 100 条都有摘要；缺失的显式写失败，整体状态降级。
    for row in source_rows:
        summaries_by_id.setdefault(
            row["video_id"],
            {
                "video_id": row["video_id"],
                "file_name": row["file_name"],
                "api_status": "failed",
                "analysis_route": "original_video_ffmpeg_frames_to_ali_vision_contact_sheet",
                "created_at": now_iso(),
                "final_model": "",
                "high_review_called": False,
                "calls": [],
                "final_json": None,
            },
        )

    metadata_rows, audit_rows, matrix_rows = build_rows(source_rows, metadata_by_id, summaries_by_id)
    write_csv(METADATA_OUT, list(metadata_rows[0].keys()), metadata_rows)
    write_csv(AUDIT_OUT, list(audit_rows[0].keys()), audit_rows)
    write_csv(MATRIX_OUT, list(matrix_rows[0].keys()), matrix_rows)
    rough_rules = build_rough_cut_rules(matrix_rows)
    write_rough_rules_csv(rough_rules)
    build_markdown_reports(metadata_rows, audit_rows, matrix_rows, rough_rules)
    if not args.no_docx:
        build_docx(metadata_rows, audit_rows, matrix_rows, rough_rules)

    counts = status_counts(matrix_rows, audit_rows, metadata_rows)
    print(json.dumps({"status": final_status(counts), **counts}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
