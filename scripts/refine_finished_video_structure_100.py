#!/usr/bin/env python3
"""基于前100条已解析成片，生成细结构二次解析产物。"""

from __future__ import annotations

import csv
import os
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover - 验证环境缺 python-docx 时只生成非 DOCX 产物
    Document = None


REPO_ROOT = Path(__file__).resolve().parents[1]
MATRIX_IN = REPO_ROOT / "素材解析_pipeline_material_analysis/08_finished_video_analysis/成片结构矩阵_finished_video_structure_matrix.csv"
EVIDENCE_IN = REPO_ROOT / "素材解析_pipeline_material_analysis/08_finished_video_analysis/成片证据索引_finished_video_evidence_index.csv"
OUTPUT_ANALYSIS_DIR = REPO_ROOT / "素材解析_pipeline_material_analysis/08_finished_video_analysis"
FACT_DIR = REPO_ROOT / "项目事实_project_facts/自动剪辑五结构成片方案_five_structure_final_video"
LOG_DIR = REPO_ROOT / "执行日志_codex_log"

FINE_MATRIX_OUT = OUTPUT_ANALYSIS_DIR / "前100条细结构矩阵_fine_grained_structure_matrix.csv"
STATS_OUT = OUTPUT_ANALYSIS_DIR / "细结构触发条件统计_fine_structure_trigger_stats.csv"
TRIGGER_LIBRARY_OUT = FACT_DIR / "25_前100条细结构触发条件库_fine_grained_structure_trigger_library.md"
TRANSITION_RULES_OUT = FACT_DIR / "26_开头中段结尾衔接规则_opening_middle_ending_transition_rules.md"
ROUGH_CUT_RULES_OUT = FACT_DIR / "27_Codex初剪细结构规则表_codex_rough_cut_fine_structure_rules.csv"
HUMAN_REPORT_MD_OUT = FACT_DIR / "28_前100条细结构解析人读版_human_readable_fine_structure_report.md"
HUMAN_REPORT_DOCX_OUT = FACT_DIR / "28_前100条细结构解析人读版_human_readable_fine_structure_report.docx"
EXECUTION_REPORT_OUT = LOG_DIR / "110_前100条成片细结构二次解析报告_fine_grained_structure_reanalysis_report.md"


AUDIENCE_PATTERNS = [
    ("年龄点名", "30岁/30+女性", re.compile(r"30\+|30岁|30岁的")),
    ("年龄点名", "35岁以后女性", re.compile(r"35岁|35\+")),
    ("年龄点名", "40岁/40+女性", re.compile(r"40岁|40\+|40＋|四十")),
    ("年龄点名", "50岁女性", re.compile(r"50岁|五十")),
    ("身份点名", "产后妈妈/宝妈", re.compile(r"产后|宝妈|妈妈|生完|生了")),
    ("身份点名", "女性/姐妹", re.compile(r"女性|女生|姐妹|女人|高段位")),
    ("身体状态点名", "漏尿/漏水/漏气人群", re.compile(r"漏尿|漏水|漏气|跑跳|咳嗽|失控|尴尬")),
    ("身体状态点名", "盆底松弛/盆底无力", re.compile(r"盆底.*(松|无力)|松弛|盆盆松|收紧|紧致|J致|致")),
    ("身体状态点名", "盆底高张/疼痛不适", re.compile(r"高张|疼痛|不舒服")),
    ("身体状态点名", "小腹凸/悬垂腹/腹压问题", re.compile(r"小肚子|小腹|悬垂腹|腹内压|腹部|肚子凸|肚子平")),
    ("身体状态点名", "臀凹陷/妈妈臀/扁平臀", re.compile(r"妈妈臀|臀凹陷|扁平臀|蜜桃臀|臀饱满|臀翘|塌臀")),
    ("身体状态点名", "膨出/脱垂/下垂", re.compile(r"膨出|脱垂|下垂|宫下垂|Z宫|前壁|内脏下垂")),
    ("动作状态点名", "练错动作/误区人群", re.compile(r"错误|误区|别再|单纯|错|不管|如果已经")),
]

PAIN_PATTERNS = [
    ("漏尿漏水尴尬", re.compile(r"漏尿|漏水|漏气|跑跳|咳嗽|失控|尴尬")),
    ("盆底松弛无力", re.compile(r"盆底.*(松|无力)|盆盆松|松弛|收紧|紧致|J致|致")),
    ("盆底高张疼痛", re.compile(r"高张|疼痛|不舒服")),
    ("小腹凸/腹压/悬垂腹", re.compile(r"小肚子|小腹|悬垂腹|腹内压|腹部|肚子凸")),
    ("妈妈臀/臀凹陷/臀扁平", re.compile(r"妈妈臀|臀凹陷|扁平臀|臀饱满|臀翘|塌臀")),
    ("膨出脱垂/下垂", re.compile(r"膨出|脱垂|下垂|宫下垂|Z宫|前壁|内脏下垂")),
    ("假胯宽/骨盆外翻/体态问题", re.compile(r"假跨|假胯|骨盆|体态|背变厚|腿细|胯")),
    ("练错动作/误区", re.compile(r"错误|误区|别再|单纯|错|疯狂练")),
    ("抗衰/好状态焦虑", re.compile(r"抗衰|好状态|变丑|年轻|好体态")),
]

RESULT_PATTERNS = [
    ("小腹平/肚子平", re.compile(r"小肚子平|肚子平|小腹.*不见|腹.*变小|少女腹|马甲线")),
    ("臀翘/臀饱满/蜜桃臀", re.compile(r"臀翘|臀饱满|蜜桃臀|妈妈臀.*改变|会跳舞的臀|扁平臀.*改变")),
    ("盆底紧致/八爪鱼/抓握力", re.compile(r"盆底.*(紧|致|J)|八爪鱼|抓握力|包裹感")),
    ("腰细/胯小/腿细", re.compile(r"腰细|胯.*小|腿细|收腰|收胯")),
    ("漏尿漏水改善", re.compile(r"告别.*漏|漏.*改善|漏.*改变|尴尬.*没有|改变.*漏")),
    ("体态改善/抗衰状态", re.compile(r"好体态|抗衰|身体好状态|背.*薄|高段位")),
]

RISK_PATTERNS = [
    ("高发误区/比例刺激", re.compile(r"70%|中招|误区")),
    ("越练越糟风险", re.compile(r"越练越漏|越练越严重|疯狂练|别再")),
    ("单一方法过度承诺", re.compile(r"单纯|只需|一个动作|一个小球|搞定|不见了")),
    ("医学/健康边界风险", re.compile(r"脱垂|膨出|漏尿|高张|疼痛|宫下垂|前壁|内脏")),
    ("效果承诺风险", re.compile(r"平了|翘了|紧了|致了|瘦|改善|告别|修复")),
]


@dataclass
class FineRow:
    values: dict[str, str]


def clean_header(row: dict[str, str]) -> dict[str, str]:
    return {k.replace("\ufeff", ""): (v or "").strip() for k, v in row.items()}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as f:
        return [clean_header(row) for row in csv.DictReader(f)]


def write_csv(path: Path, fieldnames: list[str], rows: Iterable[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


def text_blob(row: dict[str, str]) -> str:
    fields = [
        "file_name",
        "structure_formula",
        "opening_type",
        "opening_0_3s_function",
        "opening_reason",
        "middle_delivery_type",
        "middle_delivery_evidence",
        "ending_closure_type",
        "ending_next_action",
        "manual_review_items",
        "evidence_frame_summaries",
        "notes",
    ]
    return "｜".join(row.get(f, "") for f in fields)


def find_matches(patterns: list[tuple[str, re.Pattern[str]]], text: str) -> list[str]:
    return [label for label, pattern in patterns if pattern.search(text)]


def classify_audience(text: str) -> tuple[str, str, str, str, str]:
    matches = [(kind, detail) for kind, detail, pattern in AUDIENCE_PATTERNS if pattern.search(text)]
    if not matches:
        return "泛人群/场景点名", "未显式细分_基于画面或标题待复核", "无显式年龄", "无显式身份", "无显式身体状态"

    kind_counts = Counter(kind for kind, _ in matches)
    primary = kind_counts.most_common(1)[0][0]
    details = []
    for _, detail in matches:
        if detail not in details:
            details.append(detail)

    age = "；".join(d for k, d in matches if k == "年龄点名") or "无显式年龄"
    identity = "；".join(d for k, d in matches if k == "身份点名") or "无显式身份"
    state = "；".join(d for k, d in matches if k in {"身体状态点名", "动作状态点名"}) or "无显式身体状态"
    return primary, "；".join(details), age, identity, state


def classify_opening(row: dict[str, str], pain: str, result: str, risk: str) -> tuple[str, str, str]:
    opening = row["opening_type"] + "｜" + row["opening_0_3s_function"] + "｜" + row["opening_reason"]
    formula = row["structure_formula"]
    if "错误" in formula or "误区" in opening:
        style = "误区/错误先抛型"
        condition = "适合已有错误动作、错误认知或反面案例的素材，开头后必须马上解释为什么错。"
        abrupt = "如果只保留错误/恐吓画面，后面不接同一问题的原因和正确做法，会突兀。"
    elif "结果前置" in formula or "结果" in opening or result != "无明确结果触发":
        style = "结果承诺前置型"
        condition = f"适合结果很明确且中段能展示过程的素材；本条结果锚点：{result}。"
        abrupt = "如果只有结果字幕，没有动作过程或证据，开头会像硬广告。"
    elif "痛点" in formula or "痛点" in opening or pain != "无明确痛点触发":
        style = "痛点可视化/问题先抛型"
        condition = f"适合能把身体困扰具体化的素材；本条痛点锚点：{pain}。"
        abrupt = "如果痛点太大但中段只给无关动作，开头会显得焦虑化且不兑现。"
    elif "人群" in formula:
        style = "人群身份点名型"
        condition = "适合年龄、身份、产后状态或身体目标明确的素材，开头后要接该人群的真实问题。"
        abrupt = "如果只喊人群标签，后文没有该人群专属问题或动作条件，会突兀。"
    else:
        style = "主题/场景进入型"
        condition = "适合画面本身已经清楚说明主题的素材，需保留标题或第一句解释。"
        abrupt = "如果开头缺主语或缺问题来源，观众需要脑补。"

    if risk != "无明确风险触发":
        abrupt += f" 风险词需要克制处理：{risk}。"
    return style, condition, abrupt


def classify_middle(row: dict[str, str]) -> tuple[str, str, str]:
    middle = row["middle_delivery_type"] + "｜" + row["middle_delivery_evidence"] + "｜" + row["evidence_frame_summaries"]
    types = []
    evidence = []

    if re.search(r"错误|正确|对比|反面", middle):
        types.append("错误-正确对比拆解")
        evidence.append("对比画面/反例证据")
    if re.search(r"分步|步骤|拆解|阶段|口令|吸气|呼气|次数|100次|计数", middle):
        types.append("分步口令/计数教学")
        evidence.append("字幕口令/次数证据")
    if re.search(r"单动作|循环|同一动作|重复", middle):
        types.append("单动作循环演示")
        evidence.append("连续动作帧证据")
    if re.search(r"多个动作|多动作|四个动作|三个动作|两个动作|轮播", row["file_name"] + middle):
        types.append("多动作组合演示")
        evidence.append("多动作段落证据")
    if re.search(r"图解|图示|解剖图|手绘|标注", middle):
        types.append("图解/解剖辅助说明")
        evidence.append("图示/标注证据")
    if re.search(r"原因|原理|归因|解释|分层", middle):
        types.append("原因解释/归因讲解")
        evidence.append("口播或字幕解释证据")
    if re.search(r"效果|前后|结果|改善", middle):
        evidence.append("效果对比或结果提示证据")

    if not types:
        types.append(row["middle_delivery_type"] or "中段交付待复核")
    if not evidence:
        evidence.append("关键帧摘要/字幕证据")

    scores = [int(row.get(k, "0") or 0) for k in ("source_integrity_score", "visual_speech_continuity_score", "editing_flow_score")]
    if min(scores) >= 5 and row["jump_cut_risk"] == "low":
        level = "高：开头问题、中段动作/解释、结尾收束基本闭环"
    elif min(scores) >= 3:
        level = "中：结构可用，但画面/说话/桥接仍需人工复核"
    else:
        level = "低：只可作参考，不应直接进入初剪规则"
    return "；".join(dict.fromkeys(types)), "；".join(dict.fromkeys(evidence)), level


def classify_ending(row: dict[str, str]) -> tuple[str, str, str]:
    ending = row["ending_closure_type"] + "｜" + row["ending_next_action"] + "｜" + row["why_complete_or_incomplete"]
    if re.search(r"风险|警示|提醒|注意", ending):
        style = "风险提醒收束"
        condition = "适合误区、禁忌、健康边界较强的内容，结尾要留到提醒说完。"
        risk = "如果前文没有解释风险来源，结尾突然提醒会像吓人。"
    elif re.search(r"行动|号召|收藏|关注|课程|产品|跟练|指令", ending):
        style = "轻行动/跟练指令收束"
        condition = "适合中段已经交付动作或原因后，用低压跟练、收藏、咨询入口收住。"
        risk = "如果中段未兑现，结尾直接行动号召会像硬转化。"
    elif re.search(r"结果|价值|强化|升华|承诺", ending):
        style = "结果/价值重申收束"
        condition = "适合结果前置或痛点解决类视频，用结果回扣前文价值。"
        risk = "如果结果承诺没有过程证据，结尾会显得夸大。"
    elif re.search(r"动作|定格|定型|视觉|收束|微笑|手势", ending):
        style = "动作完成/视觉定型收束"
        condition = "适合动作教学类素材，保留动作完成位或定型姿态。"
        risk = "如果动作还没完成就切掉，结尾会断。"
    elif re.search(r"坚持|重复|每天|分钟|100次|计数", ending):
        style = "坚持建议/频次暗示收束"
        condition = "适合低门槛跟练素材，结尾回到频次或坚持建议。"
        risk = "如果频次没有专业依据，必须标人工复核，不能写业务通过。"
    else:
        style = row["ending_closure_type"] or "结尾待复核"
        condition = "适合当前样本中的自然收束方式，后续初剪需人工复核。"
        risk = "如果缺最后一句总结或动作完成位，结尾可能突兀。"
    return style, condition, risk


def transition_rule(row: dict[str, str], opening_style: str, middle_type: str, ending_style: str) -> str:
    if "错误" in opening_style:
        return "错误/误区开头后必须接同一错误的原因解释，再接正确动作；结尾用定型或安全提醒。"
    if "结果" in opening_style:
        return "结果前置后必须接过程证据或动作演示，结尾只能轻重申结果，不能再叠加新承诺。"
    if "痛点" in opening_style:
        return "痛点开头后必须接原因、动作或可见方案，结尾用跟练/注意事项收住。"
    if "人群" in opening_style:
        return "人群点名后必须接该人群的身体状态或动作条件，中段不能换成泛动作。"
    if "多动作" in middle_type:
        return "多动作中段要按同一目标递进，结尾回到总目标或每日练习，不另开新主题。"
    return f"{opening_style}要由{middle_type}兑现，最后用{ending_style}闭合。"


def boundary_advice(row: dict[str, str], opening_style: str, ending_style: str) -> tuple[str, str, str]:
    start = "从标题/第一句完整主语开始，保留0-3秒观看理由。"
    if "痛点" in opening_style:
        start = "向前保留痛点出现前的主语和第一帧痛点画面，不能只截结果词。"
    elif "人群" in opening_style:
        start = "从年龄/身份/身体状态点名完整出现处开始，至少保留一帧人物或动作预告。"
    elif "错误" in opening_style:
        start = "从错误动作或误区字幕完整出现处开始，保留观众能识别错误的画面。"
    elif "结果" in opening_style:
        start = "从结果承诺和动作起势同时出现处开始，若前面有适用人群要一并保留。"

    end = "留到动作完成位、总结句或低压行动说完后再切。"
    if "风险" in ending_style:
        end = "必须留到风险提醒/注意事项完整说完，必要时向后多留3-10秒。"
    elif "行动" in ending_style:
        end = "留到跟练/收藏/咨询等轻行动完整出现；若出现强转化，交人工决定是否删减。"
    elif "动作" in ending_style:
        end = "留到动作定型或最后一次完整循环结束，避免切在半个动作。"
    elif "结果" in ending_style:
        end = "留到结果回扣前文价值，不要额外拼接新结果。"

    discard = "如果开头痛点/人群、中段交付、结尾收束不属于同一问题，放弃该片段。"
    if "错误" in row["structure_formula"]:
        discard = "看不出错误和正确差异，或错误后没有解释/正确动作时放弃。"
    elif "痛点" in row["structure_formula"]:
        discard = "痛点没有被原因、动作或证据兑现时放弃，不能靠字幕硬接。"
    elif "低门槛动作" in row["structure_formula"]:
        discard = "动作被遮挡、缺发力点/方向/次数，或只剩结果口号时放弃。"
    return start, end, discard


def make_fine_row(row: dict[str, str], evidence_by_video: dict[str, dict[str, str]]) -> FineRow:
    blob = text_blob(row)
    audience_type, audience_detail, age, identity, state = classify_audience(blob)
    pains = find_matches(PAIN_PATTERNS, blob)
    results = find_matches(RESULT_PATTERNS, blob)
    risks = find_matches(RISK_PATTERNS, blob)
    pain = "；".join(pains) if pains else "无明确痛点触发"
    result = "；".join(results) if results else "无明确结果触发"
    risk = "；".join(risks) if risks else "无明确风险触发"
    opening_style, opening_condition, opening_risk = classify_opening(row, pain, result, risk)
    middle_type, middle_evidence, middle_level = classify_middle(row)
    ending_style, ending_condition, ending_risk = classify_ending(row)
    trans = transition_rule(row, opening_style, middle_type, ending_style)
    start_advice, end_advice, discard = boundary_advice(row, opening_style, ending_style)
    ev = evidence_by_video.get(row["video_id"], {})
    scores = [int(row.get(k, "0") or 0) for k in ("source_integrity_score", "visual_speech_continuity_score", "editing_flow_score")]
    confidence = "high_existing_analysis" if min(scores) >= 5 and row["jump_cut_risk"] == "low" else "medium_existing_analysis"
    manual_review = row["manual_review_items"] or "待用户人审"
    evidence_text = "；".join(
        part
        for part in [
            f"时间码：{row.get('evidence_timecodes', '')}",
            f"关键帧：{row.get('evidence_frame_summaries', '')[:180]}",
            f"证据索引：{ev.get('evidence_summary', '')[:120]}",
        ]
        if part.strip("；：")
    )

    return FineRow(
        {
            "video_id": row["video_id"],
            "file_name": row["file_name"],
            "original_structure_formula": row["structure_formula"],
            "video_type": row["video_type_primary"],
            "audience_trigger_type": audience_type,
            "audience_detail": audience_detail,
            "age_trigger": age,
            "identity_trigger": identity,
            "body_state_trigger": state,
            "pain_point_trigger": pain,
            "result_trigger": result,
            "risk_trigger": risk,
            "opening_style": opening_style,
            "opening_trigger_condition": opening_condition,
            "opening_abrupt_risk": opening_risk,
            "middle_delivery_type": middle_type,
            "middle_evidence_type": middle_evidence,
            "middle_fulfillment_level": middle_level,
            "ending_style": ending_style,
            "ending_trigger_condition": ending_condition,
            "ending_abrupt_risk": ending_risk,
            "transition_rule": trans,
            "rough_cut_start_advice": start_advice,
            "rough_cut_end_advice": end_advice,
            "discard_condition": discard,
            "manual_review_items": manual_review,
            "confidence": confidence,
            "evidence": evidence_text,
        }
    )


def stat_rows(fine_rows: list[FineRow]) -> list[dict[str, str]]:
    specs = [
        ("人群触发", "audience_trigger_type", "audience_detail"),
        ("痛点触发", "pain_point_trigger", "pain_point_trigger"),
        ("结果触发", "result_trigger", "result_trigger"),
        ("风险触发", "risk_trigger", "risk_trigger"),
        ("中段交付", "middle_delivery_type", "middle_delivery_type"),
        ("结尾收束", "ending_style", "ending_style"),
    ]
    out: list[dict[str, str]] = []
    for category, group_field, detail_field in specs:
        grouped: dict[str, list[FineRow]] = defaultdict(list)
        for fr in fine_rows:
            detail = fr.values[detail_field]
            if "；" in detail and category in {"痛点触发", "结果触发", "风险触发", "中段交付"}:
                for item in detail.split("；"):
                    if item and not item.startswith("无明确"):
                        grouped[item].append(fr)
            else:
                grouped[detail].append(fr)
        for detail, items in sorted(grouped.items(), key=lambda kv: (-len(kv[1]), kv[0])):
            examples = "；".join(fr.values["video_id"] for fr in items[:8])
            out.append(
                {
                    "category": category,
                    "trigger_type": items[0].values[group_field] if group_field in items[0].values else category,
                    "trigger_detail": detail,
                    "count": str(len(items)),
                    "example_video_ids": examples,
                    "status_note": "基于前100条已解析样本归纳，待用户人审；不外推到518条pending样本",
                }
            )
    return out


def rough_cut_rules(fine_rows: list[FineRow]) -> list[dict[str, str]]:
    grouped: dict[tuple[str, str], list[FineRow]] = defaultdict(list)
    for fr in fine_rows:
        key = (fr.values["audience_trigger_type"], first_non_empty_trigger(fr))
        grouped[key].append(fr)

    rows = []
    for idx, ((trigger_type, trigger_detail), items) in enumerate(
        sorted(grouped.items(), key=lambda kv: (-len(kv[1]), kv[0][0], kv[0][1])),
        1,
    ):
        middle = Counter(fr.values["middle_delivery_type"] for fr in items).most_common(2)
        endings = Counter(fr.values["ending_style"] for fr in items).most_common(3)
        risks = Counter(fr.values["opening_abrupt_risk"] for fr in items).most_common(1)
        examples = "；".join(fr.values["video_id"] for fr in items[:10])
        rows.append(
            {
                "rule_id": f"fgr_{idx:03d}",
                "trigger_type": trigger_type,
                "trigger_detail": trigger_detail,
                "required_middle_delivery": "；".join(x for x, _ in middle),
                "allowed_ending_styles": "；".join(x for x, _ in endings),
                "forbidden_ending_styles": "强卖式结尾；新痛点结尾；未兑现结果承诺；半个动作截断",
                "start_boundary_advice": items[0].values["rough_cut_start_advice"],
                "end_boundary_advice": items[0].values["rough_cut_end_advice"],
                "must_have_evidence": items[0].values["middle_evidence_type"],
                "discard_if": items[0].values["discard_condition"],
                "common_abrupt_risk": risks[0][0] if risks else "待复核",
                "example_video_ids": examples,
            }
        )
    return rows


def first_non_empty_trigger(fr: FineRow) -> str:
    for field in ("pain_point_trigger", "result_trigger", "risk_trigger", "audience_detail"):
        value = fr.values[field]
        if value and not value.startswith("无明确"):
            return value.split("；")[0]
    return "泛主题待复核"


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        safe = [str(cell).replace("\n", " ").replace("|", "｜") for cell in row]
        out.append("| " + " | ".join(safe) + " |")
    return "\n".join(out)


def summarize_counts(fine_rows: list[FineRow]) -> dict[str, int]:
    def count_values(field: str) -> int:
        values = set()
        for fr in fine_rows:
            for item in fr.values[field].split("；"):
                if item and not item.startswith("无明确"):
                    values.add(item)
        return len(values)

    return {
        "audience": count_values("audience_detail"),
        "pain": count_values("pain_point_trigger"),
        "result": count_values("result_trigger"),
        "risk": count_values("risk_trigger"),
        "middle": count_values("middle_delivery_type"),
        "ending": count_values("ending_style"),
    }


def top_counter(fine_rows: list[FineRow], field: str, split: bool = True, n: int = 10) -> list[tuple[str, int]]:
    c: Counter[str] = Counter()
    for fr in fine_rows:
        parts = fr.values[field].split("；") if split else [fr.values[field]]
        for item in parts:
            if item and not item.startswith("无明确"):
                c[item] += 1
    return c.most_common(n)


def build_trigger_library(fine_rows: list[FineRow], rules: list[dict[str, str]]) -> str:
    counts = summarize_counts(fine_rows)
    lines = [
        "# 前100条细结构触发条件库",
        "",
        "状态：`fine_grained_structure_reanalysis_completed_pending_user_review`",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 1. 边界",
        "",
        "- `已确认`：本库只基于 `成片结构矩阵_finished_video_structure_matrix.csv` 中 100 条 `analysis_status=success` 样本。",
        "- `已确认`：剩余 518 条仍是 `pending_not_analyzed`，本轮没有继续解析。",
        "- `已确认`：本轮未重新调用阿里模型，未读取/修改原始视频，属于已有结构结果的离线细化。",
        "- `待验证`：细结构规则能否改善直播录屏初剪，仍需要后续 6 条候选或人工复核验证。",
        "",
        "## 2. 数量摘要",
        "",
        md_table(
            ["细结构类型", "数量"],
            [
                ["人群触发细分", counts["audience"]],
                ["痛点触发细分", counts["pain"]],
                ["结果触发细分", counts["result"]],
                ["风险触发细分", counts["risk"]],
                ["中段交付细分", counts["middle"]],
                ["结尾收束细分", counts["ending"]],
            ],
        ),
        "",
        "## 3. 人群点名细分类",
        "",
        trigger_section(
            top_counter(fine_rows, "audience_detail", n=14),
            "人群点名不是只写“女性”，必须落到年龄、身份、身体状态或动作状态。",
        ),
        "",
        "## 4. 痛点触发细分类",
        "",
        trigger_section(
            top_counter(fine_rows, "pain_point_trigger", n=12),
            "痛点开头必须在中段被原因、动作或证据兑现；否则会变成焦虑开头。",
        ),
        "",
        "## 5. 结果触发细分类",
        "",
        trigger_section(
            top_counter(fine_rows, "result_trigger", n=12),
            "结果可以吸引点击，但必须有过程证据，且不能写成业务事实通过。",
        ),
        "",
        "## 6. 风险触发细分类",
        "",
        trigger_section(
            top_counter(fine_rows, "risk_trigger", n=12),
            "风险词只能作为观看理由，不能替代解释和正确做法。",
        ),
        "",
        "## 7. 动作与证据触发",
        "",
        trigger_section(
            top_counter(fine_rows, "middle_delivery_type", n=12),
            "动作和证据决定中段是否接得住开头。只有结果字幕，没有动作过程，不能进入初剪。",
        ),
        "",
        "## 8. 代表规则入口",
        "",
        md_table(
            ["规则", "触发", "中段必须有", "结尾允许", "代表视频"],
            [
                [
                    r["rule_id"],
                    f"{r['trigger_type']} / {r['trigger_detail']}",
                    r["required_middle_delivery"],
                    r["allowed_ending_styles"],
                    r["example_video_ids"],
                ]
                for r in rules[:12]
            ],
        ),
        "",
        "## 9. 禁用情况",
        "",
        "- 开头只剩年龄、人群或结果词，但中段没有同一问题的动作、原因或证据。",
        "- 痛点太大，中段只给无关动作，观众需要脑补因果。",
        "- 错误/反面案例后不解释原因，也不接正确做法。",
        "- 结尾突然强卖课程、产品或新承诺，前文没有建立价值。",
        "- 动作教学被截在半个动作、字幕挡住发力点、或缺次数/方向/呼吸等必要条件。",
        "- 医学、健康、体型、盆底、漏尿、疼痛等表达没有人工复核时，不得写审美或业务通过。",
    ]
    return "\n".join(lines) + "\n"


def trigger_section(items: list[tuple[str, int]], note: str) -> str:
    rows = []
    for label, count in items:
        rows.append(
            [
                label,
                str(count),
                "适合素材：开头能直接点中该人群/问题，且中段能同题兑现。",
                "禁用：只有标签或结果，没有原因/动作/证据。",
                "起止点：起点保留完整主语，终点留到动作或总结完整。",
            ]
        )
    return note + "\n\n" + md_table(["细分", "样本数", "适用素材", "禁用情况", "初剪建议"], rows)


def build_transition_rules(fine_rows: list[FineRow]) -> str:
    opening_counts = top_counter(fine_rows, "opening_style", split=False, n=20)
    ending_counts = top_counter(fine_rows, "ending_style", split=False, n=20)
    lines = [
        "# 开头-中段-结尾衔接规则",
        "",
        "状态：`fine_grained_structure_reanalysis_completed_pending_user_review`",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 1. 核心结论",
        "",
        "开头自然不自然，不取决于钩子够不够猛，而取决于中段有没有兑现同一个问题。结尾自然不自然，不取决于有没有 CTA，而取决于前文价值有没有被完整收住。",
        "",
        "## 2. 开头类型与必须承接",
        "",
        md_table(
            ["开头类型", "样本数", "中段必须接什么", "常见突兀风险"],
            [
                [
                    label,
                    str(count),
                    opening_required_middle(label),
                    opening_common_risk(label),
                ]
                for label, count in opening_counts
            ],
        ),
        "",
        "## 3. 哪些开头必须接原因解释",
        "",
        "- 误区/错误先抛型：必须解释错在哪里，再给正确动作。",
        "- 痛点可视化/问题先抛型：必须解释原因或至少给动作为什么能接住痛点。",
        "- 风险比例型，如 `70%女性中招`：必须给来源或机制解释，不能只制造危机感。",
        "",
        "## 4. 哪些开头必须接动作演示",
        "",
        "- 人群身份点名型：点名后要马上给该人群能做的动作或判断条件。",
        "- 结果承诺前置型：结果后必须有操作过程，不能只保留结果字样。",
        "- 单动作解决型：至少保留一次完整动作循环，最好含起势、发力、回位。",
        "",
        "## 5. 哪些开头必须接对比证据",
        "",
        "- 错误/正确结构：错误与正确必须是同一个动作问题。",
        "- 痛点可视化 + 效果对比：前后对比之间必须有过程桥接。",
        "- 结果强承诺：必须有动作过程、图示、口令或关键帧证据支撑。",
        "",
        "## 6. 结尾收束规则",
        "",
        md_table(
            ["结尾类型", "样本数", "适合接在什么中段后", "突兀风险"],
            [
                [
                    label,
                    str(count),
                    ending_allowed_context(label),
                    ending_common_risk(label),
                ]
                for label, count in ending_counts
            ],
        ),
        "",
        "## 7. 向后多留 3-10 秒的情况",
        "",
        "- 结尾正在说风险提醒、注意事项、禁忌或适用人群边界。",
        "- 动作还没有完成一次完整循环，尤其是呼吸、提臀、夹球、回位等动作。",
        "- 结尾刚从动作转到总结句，需要保留总结句才能自然闭合。",
        "- 出现轻 CTA 但语气仍在承接前文价值，可以留到低压行动完整说完。",
        "",
        "## 8. 应该放弃片段的情况",
        "",
        "- 开头和中段不是同一个问题。",
        "- 中段只剩动作次数，没有发力点、方向、呼吸、工具或安全边界。",
        "- 反面案例后没有原因解释和正确做法。",
        "- 结尾突然切课程/产品/强转化，前文没有建立解决方案价值。",
        "- 需要凭剪辑者脑补才能成立的桥接，一律不作为自动初剪候选。",
    ]
    return "\n".join(lines) + "\n"


def opening_required_middle(label: str) -> str:
    if "错误" in label:
        return "同一错误的原因解释 + 正确动作"
    if "痛点" in label:
        return "痛点原因、动作方案或可见证据"
    if "结果" in label:
        return "操作过程、动作演示、图示或对比证据"
    if "人群" in label:
        return "该人群专属身体状态、动作条件或低门槛动作"
    return "主题解释或完整动作链"


def opening_common_risk(label: str) -> str:
    if "错误" in label:
        return "只吓人、不解释；错误和正确不是同一问题"
    if "痛点" in label:
        return "痛点大，中段轻；焦虑开头没有兑现"
    if "结果" in label:
        return "结果像硬广，缺过程证据"
    if "人群" in label:
        return "只喊人群标签，后文泛化"
    return "缺主语或缺第一句解释"


def ending_allowed_context(label: str) -> str:
    if "风险" in label:
        return "误区纠错、健康边界、动作禁忌"
    if "行动" in label:
        return "中段已完成动作/原因/方案交付"
    if "结果" in label:
        return "前文有过程或证据支撑"
    if "动作" in label:
        return "动作教学、纠错、单动作循环"
    return "完整表达已闭合后"


def ending_common_risk(label: str) -> str:
    if "行动" in label:
        return "中段未兑现时突然强卖"
    if "结果" in label:
        return "重复承诺但没有证据"
    if "风险" in label:
        return "前面没有解释风险来源"
    if "动作" in label:
        return "动作未完成就切断"
    return "缺总结句或新开主题"


def build_human_report(fine_rows: list[FineRow], rules: list[dict[str, str]]) -> str:
    counts = summarize_counts(fine_rows)
    top_opening_risks = top_counter(fine_rows, "opening_abrupt_risk", split=False, n=5)
    top_ending_risks = top_counter(fine_rows, "ending_abrupt_risk", split=False, n=5)
    lines = [
        "# 前100条细结构解析人读版报告",
        "",
        "副标题：为什么初剪开头/结尾会突兀，以及后续 Codex 应该怎么取起止点",
        "",
        "状态：`fine_grained_structure_reanalysis_completed_pending_user_review`",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 1. 主结论",
        "",
        "本轮把前 100 条已解析成片，从“大结构公式”继续拆到“触发条件 + 中段兑现 + 结尾收束 + 初剪起止点”。",
        "",
        "这次没有继续解析 518 条 pending 视频，也没有重新调用阿里模型。所有细化都来自已有结构矩阵、证据索引和人读版报告，因此结论是 `部分成立，待用户人审`，不是审美通过，也不是业务通过。",
        "",
        "## 2. 细结构数量",
        "",
        md_table(
            ["项目", "数量"],
            [
                ["人群触发细分", counts["audience"]],
                ["痛点触发细分", counts["pain"]],
                ["结果触发细分", counts["result"]],
                ["风险触发细分", counts["risk"]],
                ["中段交付细分", counts["middle"]],
                ["结尾收束细分", counts["ending"]],
            ],
        ),
        "",
        "## 3. 为什么之前初剪会开头突兀",
        "",
        "开头突兀通常不是因为开头不够刺激，而是因为只截到了“标签”。比如只截到 `40+女性`、`产后妈妈`、`小腹凸`、`漏尿`，但没有保留它后面对应的问题解释或动作预告，观众就会觉得前后断开。",
        "",
        "本轮细化后，开头必须同时满足三件事：",
        "",
        "1. 观众知道这条和谁有关。",
        "2. 观众知道具体问题是什么。",
        "3. 中段马上兑现这个问题，而不是换成另一个动作或另一个结果。",
        "",
        md_table(
            ["常见开头突兀风险", "样本数"],
            [[label, str(count)] for label, count in top_opening_risks],
        ),
        "",
        "## 4. 为什么之前初剪会结尾突兀",
        "",
        "结尾突兀通常是因为剪在了半个动作、半句总结、或突然接了一个强行动指令。人工剪辑看起来顺，是因为它会留到动作完成位、总结句、轻行动或风险提醒说完。",
        "",
        "本轮细化后，结尾至少要满足一条：动作完成、价值回扣、轻行动完整、风险提醒完整。没有这些，宁可向后多留 3-10 秒，也不要硬切。",
        "",
        md_table(
            ["常见结尾突兀风险", "样本数"],
            [[label, str(count)] for label, count in top_ending_risks],
        ),
        "",
        "## 5. 后续直播录屏初剪最该先用的 5 条规则",
        "",
        md_table(
            ["优先级", "规则", "怎么用"],
            [
                ["1", "痛点开头必须兑现", "看到漏尿、小腹凸、臀凹陷等痛点，必须向后找到原因、动作或图示证据；找不到就放弃。"],
                ["2", "人群点名不能单独成立", "年龄/身份/产后标签后面必须接该人群专属问题或动作条件。"],
                ["3", "动作教学留完整循环", "至少保留起势、发力、回位或一次完整呼吸，不剪半个动作。"],
                ["4", "错误示范必须接正解", "错误动作、误区、反面案例后必须有原因解释和正确动作。"],
                ["5", "强结果必须有证据", "肚子平、臀翘、盆底紧等结果词只能在中段有过程证据时保留。"],
            ],
        ),
        "",
        "## 6. Codex 后续怎么用",
        "",
        "先读 `27_Codex初剪细结构规则表_codex_rough_cut_fine_structure_rules.csv`，用触发类型决定是否保留片段；再读本报告判断为什么开头/结尾自然；最后用 `前100条细结构矩阵_fine_grained_structure_matrix.csv` 找相似样本。",
        "",
        "自动初剪时，不要为了凑结构硬拼。宁可少选，也不要把开头痛点、中段动作、结尾行动拼成三个互不相干的段落。",
        "",
        "## 7. 代表样本索引",
        "",
        md_table(
            ["视频", "细人群/问题", "中段兑现", "结尾收束", "起止点"],
            [
                [
                    fr.values["video_id"],
                    first_non_empty_trigger(fr),
                    fr.values["middle_delivery_type"],
                    fr.values["ending_style"],
                    fr.values["rough_cut_start_advice"] + " / " + fr.values["rough_cut_end_advice"],
                ]
                for fr in fine_rows[:20]
            ],
        ),
        "",
        "完整 100 条逐条细结构见 CSV 矩阵。本报告只保留人读版摘要，避免把业务读者淹没在字段里。",
        "",
        "## 8. 边界",
        "",
        "- `已确认`：只处理前 100 条 success。",
        "- `已确认`：未处理 518 条 pending。",
        "- `已确认`：未提交媒体、API 原始输出、secret 或缓存。",
        "- `待验证`：细结构对直播录屏初剪的改善效果。",
        "- `待用户人审`：审美、人感、动作专业性和业务事实。",
    ]
    return "\n".join(lines) + "\n"


def build_execution_report(
    fine_rows: list[FineRow],
    stats: list[dict[str, str]],
    rules: list[dict[str, str]],
    existing_video_count: int,
) -> str:
    counts = summarize_counts(fine_rows)
    generated = [
        str(EXECUTION_REPORT_OUT.relative_to(REPO_ROOT)),
        str(TRIGGER_LIBRARY_OUT.relative_to(REPO_ROOT)),
        str(TRANSITION_RULES_OUT.relative_to(REPO_ROOT)),
        str(ROUGH_CUT_RULES_OUT.relative_to(REPO_ROOT)),
        str(HUMAN_REPORT_MD_OUT.relative_to(REPO_ROOT)),
        str(HUMAN_REPORT_DOCX_OUT.relative_to(REPO_ROOT)),
        str(FINE_MATRIX_OUT.relative_to(REPO_ROOT)),
        str(STATS_OUT.relative_to(REPO_ROOT)),
        "scripts/refine_finished_video_structure_100.py",
    ]
    lines = [
        "# 前100条成片细结构二次解析报告",
        "",
        "状态：`fine_grained_structure_reanalysis_completed_pending_user_review`",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "任务类型：`finished_video_fine_grained_structure_reanalysis`",
        "",
        "## 1. 执行结果",
        "",
        md_table(
            ["项目", "结果"],
            [
                ["当前仓库", "`fthytwerwt-sudo/lanxinse--`"],
                ["本地仓库路径", f"`{REPO_ROOT}`"],
                ["输入素材目录", "`/Volumes/WD_BLACK/澜心社剪辑/剪辑解析数据/AI需要的成片`"],
                ["结构矩阵总行数", "618"],
                ["本轮处理", "100 条 `success`"],
                ["未处理", "518 条 `pending_not_analyzed`"],
                ["原视频路径存在数", str(existing_video_count)],
                ["是否重新调用阿里", "否"],
                ["是否抽新帧/读原视频内容", "否"],
                ["是否提交媒体", "否"],
                ["是否提交 secret", "否"],
                ["是否提交 API 原始输出", "否"],
                ["状态边界", "`部分成立，待用户人审；不写审美/业务通过`"],
            ],
        ),
        "",
        "## 2. 读取文件",
        "",
        "- `AGENTS.md`",
        "- `GPT项目资料同步包_gpt_project_mechanism_sync/00_GPT_Project上传说明_readme.md`",
        "- `GPT项目资料同步包_gpt_project_mechanism_sync/01_三层架构与事实源边界_three_layer_source_boundary.md`",
        "- `GPT项目资料同步包_gpt_project_mechanism_sync/02_P0-P1-P2锚点与抗漂移机制_anchor_priority_anti_drift.md`",
        "- `GPT项目资料同步包_gpt_project_mechanism_sync/04_Codex执行落库机制_codex_execution_to_repo_protocol.md`",
        "- `GPT项目资料同步包_gpt_project_mechanism_sync/07_输出硬规则与中文语义对齐_output_hard_rules.md`",
        "- `GPT项目资料同步包_gpt_project_mechanism_sync/23_六层需求确认与实现设计闸门机制_six_layer_requirement_implementation_gate.md`",
        "- `执行日志_codex_log/106_阿里模型重连验证报告_ali_model_reconnect_after_env_update_report.md`",
        "- `执行日志_codex_log/107_AI需要的成片全量结构解析报告_finished_video_full_analysis_report.md`",
        "- `执行日志_codex_log/108_前100条成片样本人读版DOCX生成报告_human_readable_docx_report.md`",
        "- `项目事实_project_facts/自动剪辑五结构成片方案_five_structure_final_video/20_视频结构公式库_video_structure_formula_library.md`",
        "- `项目事实_project_facts/自动剪辑五结构成片方案_five_structure_final_video/21_初剪完整性与素材连续性判断标准_rough_cut_integrity_continuity_standard.md`",
        "- `项目事实_project_facts/自动剪辑五结构成片方案_five_structure_final_video/22_成片样本类型与结构总表_finished_video_type_structure_inventory.md`",
        "- `项目事实_project_facts/自动剪辑五结构成片方案_five_structure_final_video/23_Codex成片解析字段规范_finished_video_analysis_schema.md`",
        "- `项目事实_project_facts/自动剪辑五结构成片方案_five_structure_final_video/24_前100条成片样本结构解析人读版_human_readable_report.md`",
        "- `素材解析_pipeline_material_analysis/08_finished_video_analysis/成片结构矩阵_finished_video_structure_matrix.csv`",
        "- `素材解析_pipeline_material_analysis/08_finished_video_analysis/成片证据索引_finished_video_evidence_index.csv`",
        "",
        "## 3. 生成文件",
        "",
        *[f"- `{p}`" for p in generated],
        "",
        "## 4. 细结构结果摘要",
        "",
        md_table(
            ["细结构", "数量"],
            [
                ["人群触发", counts["audience"]],
                ["痛点触发", counts["pain"]],
                ["结果触发", counts["result"]],
                ["风险触发", counts["risk"]],
                ["中段交付", counts["middle"]],
                ["结尾收束", counts["ending"]],
                ["规则表行数", len(rules)],
                ["统计表行数", len(stats)],
            ],
        ),
        "",
        "最容易造成开头突兀：只截取年龄/身份/痛点/结果标签，没有保留后续同题兑现段。",
        "",
        "最容易造成结尾突兀：动作未完成、总结句未说完、或中段未兑现时突然接强行动/强结果。",
        "",
        "## 5. 验证记录",
        "",
        "本报告生成时，文件生成链路已完成；最终验证、commit、push、remote HEAD readback 由本轮 Codex 最终回报补充。",
        "",
        "计划执行命令：",
        "",
        "- `python3 scripts/check_ali_config_safety.py`",
        "- `python3 -m py_compile scripts/refine_finished_video_structure_100.py`",
        "- `ls -lh` 检查 DOCX 与 Markdown/CSV 产物",
        "- CSV 行数检查：矩阵 100 行，统计表非空",
        "- DOCX render PNG 视觉 QA",
        "- `git diff --check`",
        "- path-limited `git add` / `git commit` / `git push` / remote HEAD readback",
        "",
        "## 6. 边界确认",
        "",
        md_table(
            ["边界", "结果"],
            [
                ["是否解析 518 条", "否"],
                ["是否提交视频", "否"],
                ["是否提交图片", "否"],
                ["是否提交音频", "否"],
                ["是否提交 `.env`", "否"],
                ["是否提交 API key/token/cookie", "否"],
                ["是否提交完整 API 输出", "否"],
                ["是否写审美通过", "否"],
                ["是否写业务通过", "否"],
                ["是否写稳定初剪", "否"],
            ],
        ),
        "",
        "## 7. 下一步建议",
        "",
        "1. 用户先复核细结构人读版 DOCX。",
        "2. 如果细结构对，再回到直播录屏，重新剪 6 条初剪候选。",
        "3. 如果开头规则不对，继续修开头触发条件。",
        "4. 如果结尾规则不对，继续修结尾收束规则。",
        "5. 如果仍然太粗，再新增起止点微调规则。",
    ]
    return "\n".join(lines) + "\n"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E2EC")
        borders.append(tag)
    tbl_pr.append(borders)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_docx_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_p(doc, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def build_docx(fine_rows: list[FineRow], rules: list[dict[str, str]], path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is unavailable")

    counts = summarize_counts(fine_rows)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color

    header = section.header.paragraphs[0]
    header.text = "澜心社自动剪辑 | 前100条细结构二次解析"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(88, 96, 105)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("前100条成片细结构解析人读版")
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("为什么初剪开头/结尾会突兀，以及 Codex 后续怎么取起止点")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(70, 70, 70)

    add_docx_table(
        doc,
        ["项目", "状态"],
        [
            ["处理范围", "前100条 success 成片样本"],
            ["未处理范围", "518条 pending_not_analyzed"],
            ["模型/API", "未重新调用阿里，基于已有CSV/Markdown离线细化"],
            ["验收边界", "待用户人审，不写审美/业务通过"],
        ],
        [1.7, 4.6],
    )

    doc.add_heading("1. 主结论", level=1)
    add_p(doc, "这次把已有大结构继续拆成触发条件、中段兑现、结尾收束和初剪起止点。核心判断是：开头自然不自然，看中段是否兑现；结尾自然不自然，看动作、总结、轻行动或风险提醒是否完整闭合。")
    add_p(doc, "本报告是给人复核的摘要。完整逐条字段在前100条细结构矩阵 CSV 中。")

    doc.add_heading("2. 细结构数量", level=1)
    add_docx_table(
        doc,
        ["细结构类型", "数量"],
        [
            ["人群触发细分", str(counts["audience"])],
            ["痛点触发细分", str(counts["pain"])],
            ["结果触发细分", str(counts["result"])],
            ["风险触发细分", str(counts["risk"])],
            ["中段交付细分", str(counts["middle"])],
            ["结尾收束细分", str(counts["ending"])],
        ],
        [4.6, 1.4],
    )

    doc.add_heading("3. 开头为什么会突兀", level=1)
    add_p(doc, "只截到“40+女性”“产后妈妈”“小腹凸”“漏尿”这类标签，但没保留后面的问题解释或动作预告，观众就会觉得前后断开。开头必须让人知道：和谁有关、问题是什么、中段马上怎么接。")

    doc.add_heading("4. 结尾为什么会突兀", level=1)
    add_p(doc, "结尾常见问题是剪在半个动作、半句总结，或中段还没兑现就突然接强行动。后续初剪应留到动作完成位、价值回扣、轻行动或风险提醒完整出现。")

    doc.add_heading("5. 后续初剪优先规则", level=1)
    add_docx_table(
        doc,
        ["优先级", "规则", "使用方式"],
        [
            ["1", "痛点开头必须兑现", "向后找原因、动作或图示证据；找不到就放弃。"],
            ["2", "人群点名不能单独成立", "年龄/身份/产后标签后面必须接专属问题或动作条件。"],
            ["3", "动作教学留完整循环", "保留起势、发力、回位或完整呼吸，不剪半个动作。"],
            ["4", "错误示范必须接正解", "错误、误区、反面案例后必须有原因解释和正确动作。"],
            ["5", "强结果必须有证据", "肚子平、臀翘、盆底紧等结果词必须有过程证据。"],
        ],
        [0.7, 1.7, 3.8],
    )

    doc.add_heading("6. 代表规则", level=1)
    add_docx_table(
        doc,
        ["规则", "触发", "中段必须有", "代表视频"],
        [
            [
                r["rule_id"],
                f"{r['trigger_type']} / {r['trigger_detail']}",
                r["required_middle_delivery"],
                r["example_video_ids"],
            ]
            for r in rules[:8]
        ],
        [0.7, 2.0, 2.0, 1.6],
    )

    doc.add_heading("7. 边界", level=1)
    for item in [
        "只处理前100条 success，不处理518条 pending。",
        "未重新调用阿里模型，未抽新帧，未读取或修改原视频。",
        "不提交媒体、secret、API原始输出或缓存。",
        "审美、人感、动作专业性、健康/业务事实均待用户和业务复核。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # 附录使用横向页，承载前20条代表样本索引，避免正文过密。
    doc.add_page_break()
    sec = doc.add_section(WD_ORIENT.LANDSCAPE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    doc.add_heading("附录：前4条代表样本索引", level=1)
    add_docx_table(
        doc,
        ["视频", "细人群/问题", "中段兑现", "结尾", "起止点"],
        [
            [
                fr.values["video_id"],
                first_non_empty_trigger(fr),
                fr.values["middle_delivery_type"],
                fr.values["ending_style"],
                "起：" + fr.values["rough_cut_start_advice"] + "\n止：" + fr.values["rough_cut_end_advice"],
            ]
            for fr in fine_rows[:4]
        ],
        [0.7, 1.7, 2.0, 1.3, 3.2],
    )

    doc.save(path)


def main() -> None:
    rows = read_csv(MATRIX_IN)
    evidence_rows = read_csv(EVIDENCE_IN)
    status_counts = Counter(row["analysis_status"] for row in rows)
    if status_counts["success"] != 100 or status_counts["pending_not_analyzed"] != 518:
        raise SystemExit(f"无法确认100/518边界：{status_counts}")

    success_rows = [row for row in rows if row["analysis_status"] == "success"]
    evidence_by_video = {row["video_id"]: row for row in evidence_rows if row.get("analysis_status") == "success"}
    fine_rows = [make_fine_row(row, evidence_by_video) for row in success_rows]
    existing_video_count = sum(1 for row in success_rows if os.path.exists(row["source_path"]))

    fieldnames = list(fine_rows[0].values.keys())
    write_csv(FINE_MATRIX_OUT, fieldnames, [fr.values for fr in fine_rows])

    stats = stat_rows(fine_rows)
    write_csv(
        STATS_OUT,
        ["category", "trigger_type", "trigger_detail", "count", "example_video_ids", "status_note"],
        stats,
    )

    rules = rough_cut_rules(fine_rows)
    write_csv(
        ROUGH_CUT_RULES_OUT,
        [
            "rule_id",
            "trigger_type",
            "trigger_detail",
            "required_middle_delivery",
            "allowed_ending_styles",
            "forbidden_ending_styles",
            "start_boundary_advice",
            "end_boundary_advice",
            "must_have_evidence",
            "discard_if",
            "common_abrupt_risk",
            "example_video_ids",
        ],
        rules,
    )

    TRIGGER_LIBRARY_OUT.write_text(build_trigger_library(fine_rows, rules), encoding="utf-8")
    TRANSITION_RULES_OUT.write_text(build_transition_rules(fine_rows), encoding="utf-8")
    HUMAN_REPORT_MD_OUT.write_text(build_human_report(fine_rows, rules), encoding="utf-8")
    build_docx(fine_rows, rules, HUMAN_REPORT_DOCX_OUT)
    EXECUTION_REPORT_OUT.write_text(
        build_execution_report(fine_rows, stats, rules, existing_video_count),
        encoding="utf-8",
    )

    print(f"success_rows={len(fine_rows)}")
    print(f"pending_not_analyzed={status_counts['pending_not_analyzed']}")
    print(f"existing_source_paths={existing_video_count}")
    print(f"stats_rows={len(stats)}")
    print(f"rules_rows={len(rules)}")
    print(f"docx={HUMAN_REPORT_DOCX_OUT}")


if __name__ == "__main__":
    main()
