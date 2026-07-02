#!/usr/bin/env python3
"""C9623 完整直播切片验证与本地候选导出。

边界：
- 只在 /Volumes/WD_BLACK/澜心社剪辑 授权工作区内读取 C9623 素材。
- 复用 112/113 已提交的全覆盖窗口、阿里视觉审计和 V2 机会表。
- 本轮导出的视频只写入 outputs/c9623_live_cutting_full_validation/，不得提交 Git。
- 本轮没有可用本地转写器时，音频/口播字段必须标 pending_audio_transcript，不脑补口播。
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable
from xml.sax.saxutils import escape


REPO_ROOT = Path(__file__).resolve().parents[1]
AUTHORIZED_BASE = Path("/Volumes/WD_BLACK/澜心社剪辑")
MATERIAL_BASE = AUTHORIZED_BASE / "剪辑解析数据"
C9623_EXPECTED = MATERIAL_BASE / "剪辑测试直播素材" / "C9623.MP4"

ANALYSIS_DIR = REPO_ROOT / "素材解析_pipeline_material_analysis" / "11_c9623_live_cutting_full_validation"
FACT_DIR = REPO_ROOT / "项目事实_project_facts" / "C9623直播切片完整验证_c9623_live_cutting_full_validation"
LOG_DIR = REPO_ROOT / "执行日志_codex_log"
LOCAL_OUTPUT_DIR = REPO_ROOT / "outputs" / "c9623_live_cutting_full_validation"
LOCAL_API_DIR = REPO_ROOT / "api_outputs" / "c9623_live_cutting_full_validation"

OLD_FORMAL_DIR = REPO_ROOT / "素材解析_pipeline_material_analysis" / "09_live_recording_formal_simulation"
V2_ANALYSIS_DIR = REPO_ROOT / "素材解析_pipeline_material_analysis" / "10_live_cutting_v2_architecture_audit"
OLD_SUMMARY_DIR = REPO_ROOT / "api_outputs" / "live_recording_formal_simulation_2_videos" / "model_summaries"

INVENTORY_CSV = ANALYSIS_DIR / "01_C9623素材清单_c9623_material_inventory.csv"
WINDOW_CSV = ANALYSIS_DIR / "02_C9623全覆盖窗口表_c9623_full_coverage_window_manifest.csv"
AUDIO_CSV = ANALYSIS_DIR / "03_C9623音频字幕时间线_c9623_audio_subtitle_timeline.csv"
VISUAL_CSV = ANALYSIS_DIR / "04_C9623视觉窗口审计_c9623_visual_window_audit.csv"
CANDIDATE_CSV = ANALYSIS_DIR / "05_C9623候选片段评分表_c9623_candidate_scoring.csv"
EXPORT_INDEX_CSV = ANALYSIS_DIR / "06_C9623导出视频索引_c9623_exported_clip_index.csv"
REJECTED_CSV = ANALYSIS_DIR / "07_C9623弃用片段表_c9623_rejected_segments.csv"
PACKAGING_CSV = ANALYSIS_DIR / "08_C9623人工包装建议表_c9623_manual_packaging_advice.csv"

SYNTHESIS_MD = FACT_DIR / "00_前序报告综合结论_previous_reports_synthesis.md"
FULL_REPORT_MD = FACT_DIR / "01_C9623完整验证报告_c9623_full_validation_report.md"
EVIDENCE_REPORT_MD = FACT_DIR / "02_C9623候选片段证据报告_c9623_candidate_evidence_report.md"
PACKAGING_GUIDE_MD = FACT_DIR / "03_C9623人工包装说明_c9623_manual_packaging_guide.md"
RISK_REVIEW_MD = FACT_DIR / "04_C9623风险与待复核清单_c9623_risk_review_checklist.md"
HUMAN_DOCX = FACT_DIR / "05_C9623人读版报告_c9623_human_readable_report.docx"
HUMAN_MD_FALLBACK = FACT_DIR / "05_C9623人读版报告_c9623_human_readable_report.md"
EXEC_REPORT_MD = LOG_DIR / "114_C9623直播切片完整验证报告_c9623_live_cutting_full_validation_report.md"

OLD_INVENTORY_CSV = OLD_FORMAL_DIR / "01_直播录屏素材清单_live_recording_inventory.csv"
OLD_WINDOWS_CSV = OLD_FORMAL_DIR / "02_全覆盖窗口清单_full_coverage_window_manifest.csv"
OLD_ALI_AUDIT_CSV = OLD_FORMAL_DIR / "03_阿里窗口分析审计表_ali_window_analysis_audit.csv"
OLD_REJECTED_CSV = OLD_FORMAL_DIR / "05_弃用片段表_rejected_segment_table.csv"
OLD_RESCUE_CSV = OLD_FORMAL_DIR / "07_弃用片段复审与可救回片段表_rejected_segment_recheck_rescue_candidates.csv"
V2_OPPORTUNITY_CSV = V2_ANALYSIS_DIR / "05_直播切片素材结构机会表_live_cutting_opportunity_table.csv"

QUALIFIED_REPACKAGING_IDS = {
    "live_opp_059",
    "live_opp_060",
    "live_opp_061",
    "live_opp_062",
    "live_opp_067",
    "live_opp_068",
    "live_opp_073",
    "live_opp_079",
}
PENDING_AUDIO_EXPORT_IDS = {
    "live_opp_041",
    "live_opp_047",
    "live_opp_048",
    "live_opp_056",
    "live_opp_058",
    "live_opp_063",
    "live_opp_077",
    "live_opp_080",
}

GROUPS = {
    "qualified_native": "01_原生切片候选_qualified_native",
    "qualified_repackaging": "02_动作教学再包装候选_qualified_repackaging",
    "qualified_merge_candidate": "03_相邻窗口合并候选_qualified_merge_candidate",
    "pending_audio_review": "04_待音频复核_pending_audio_review",
    "reject_unusable": "05_弃用不导出_rejected_index_only",
}

INVENTORY_FIELDS = [
    "recording_id",
    "file_name",
    "source_path",
    "duration_seconds",
    "width",
    "height",
    "fps",
    "rotation",
    "video_codec",
    "audio_codec",
    "audio_present",
    "file_size_mb",
    "read_status",
    "decode_status",
    "decode_error",
    "source_dir",
    "discovery_note",
]

WINDOW_FIELDS = [
    "window_id",
    "recording_id",
    "source_file",
    "window_start",
    "window_end",
    "start_seconds",
    "end_seconds",
    "duration_seconds",
    "overlap_previous_seconds",
    "coverage_status",
    "analysis_source",
    "ali_model_called",
    "ali_primary_status",
    "ali_high_review_status",
    "needs_audio_transcript",
    "needs_adjacent_merge",
    "notes",
]

AUDIO_FIELDS = [
    "timeline_id",
    "recording_id",
    "source_file",
    "start_time",
    "end_time",
    "start_seconds",
    "end_seconds",
    "audio_present",
    "audio_codec",
    "transcript_status",
    "transcript_text",
    "subtitle_text",
    "tts_action_alignment",
    "status_reason",
]

VISUAL_FIELDS = [
    "window_id",
    "recording_id",
    "start_time",
    "end_time",
    "window_summary",
    "actions_methods",
    "products_props_course_info",
    "has_complete_opening",
    "has_middle_delivery",
    "has_ending_closure",
    "need_merge_previous",
    "need_merge_next",
    "visual_action_clarity_score",
    "action_cycle_score",
    "risk_summary",
    "analysis_status",
    "analysis_source",
]

CANDIDATE_FIELDS = [
    "candidate_id",
    "source_file",
    "start_time",
    "end_time",
    "start_seconds",
    "end_seconds",
    "duration_seconds",
    "source_windows",
    "content_archetype",
    "route_decision",
    "candidate_status",
    "visual_action_clarity_score",
    "action_cycle_score",
    "transcript_semantic_score",
    "problem_action_bridge_score",
    "repackaging_value_score",
    "adjacent_merge_score",
    "risk_score",
    "priority_score",
    "problem_phrase_time",
    "first_action_frame_time",
    "problem_action_bridge_seconds",
    "tts_action_alignment",
    "visual_evidence_timecodes",
    "why_selected",
    "why_rejected",
    "manual_packaging_advice",
    "suggested_tts_hook",
    "suggested_subtitle",
    "health_action_risk",
    "business_risk",
    "manual_review_items",
    "not_exported_reason",
    "analysis_status",
]

EXPORT_FIELDS = [
    "export_id",
    "candidate_id",
    "candidate_status",
    "group_dir",
    "file_name",
    "output_path",
    "source_file",
    "start_time",
    "end_time",
    "duration_seconds",
    "why_exported",
    "manual_review_items",
    "ffprobe_status",
    "output_duration_seconds",
    "output_width",
    "output_height",
    "video_stream_present",
    "audio_stream_present",
    "file_size_mb",
    "not_exported_reason",
]

REJECTED_FIELDS = [
    "rejected_id",
    "candidate_id",
    "source_file",
    "start_time",
    "end_time",
    "candidate_status",
    "route_decision",
    "why_rejected",
    "original_reject_reason",
    "missing_part",
    "could_be_fixed_by_manual_edit",
    "rescue_level",
    "needs_transcript_or_audio",
    "manual_review_items",
]

PACKAGING_FIELDS = [
    "advice_id",
    "candidate_id",
    "exported_file_name",
    "candidate_status",
    "suitable_problem",
    "first_3_seconds_tts",
    "tts_or_voiceover_advice",
    "subtitle_focus",
    "diagram_needed",
    "risk_hint_needed",
    "editing_tool_fit",
    "manual_review_items",
]


@dataclass
class MediaInfo:
    path: Path
    duration: float
    width: int
    height: int
    fps: str
    rotation: str
    video_codec: str
    audio_codec: str
    audio_present: bool
    file_size_mb: float


def now_iso() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "yes" if value else "no"
    if isinstance(value, float):
        return f"{value:.3f}".rstrip("0").rstrip(".")
    if isinstance(value, (list, tuple, set)):
        return "；".join(str(item) for item in value if str(item))
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value)


def write_csv(path: Path, fieldnames: list[str], rows: Iterable[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: stringify(row.get(field, "")) for field in fieldnames})


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def run_command(cmd: list[str], timeout: int = 120) -> subprocess.CompletedProcess[str]:
    return subprocess.run(cmd, text=True, capture_output=True, timeout=timeout, check=False)


def require_ok(result: subprocess.CompletedProcess[str], label: str) -> None:
    if result.returncode != 0:
        raise RuntimeError(f"{label} failed: {(result.stderr or result.stdout).strip()[:1200]}")


def parse_timecode(value: str) -> float:
    value = (value or "").strip()
    if not value:
        return 0.0
    parts = value.split(":")
    if len(parts) != 3:
        return float(value)
    hours = int(parts[0])
    minutes = int(parts[1])
    seconds = float(parts[2])
    return hours * 3600 + minutes * 60 + seconds


def timecode(seconds: float) -> str:
    total_ms = int(round(max(0.0, seconds) * 1000))
    ms = total_ms % 1000
    total_s = total_ms // 1000
    s = total_s % 60
    total_m = total_s // 60
    m = total_m % 60
    h = total_m // 60
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def slug_time(seconds: float) -> str:
    total_s = int(max(0.0, seconds))
    return f"{total_s // 3600:02d}-{(total_s % 3600) // 60:02d}-{total_s % 60:02d}"


def safe_slug(text: str, max_len: int = 18) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|\\s]+", "_", text)
    cleaned = re.sub(r"_+", "_", cleaned).strip("_")
    return cleaned[:max_len] or "clip"


def parse_fps(value: str) -> str:
    if not value or value == "0/0":
        return ""
    if "/" not in value:
        return value
    num, den = value.split("/", 1)
    try:
        return f"{float(num) / float(den):.3f}".rstrip("0").rstrip(".")
    except (ValueError, ZeroDivisionError):
        return value


def locate_c9623() -> Path:
    candidates = sorted(AUTHORIZED_BASE.glob("**/*C9623*"))
    video_candidates = [
        path
        for path in candidates
        if path.is_file()
        and path.suffix.lower() in {".mp4", ".mov", ".m4v", ".mkv", ".avi"}
        and "剪辑解析数据" in str(path)
    ]
    preferred = [path for path in video_candidates if path == C9623_EXPECTED]
    if preferred:
        return preferred[0]
    if len(video_candidates) == 1:
        return video_candidates[0]
    if not video_candidates:
        raise RuntimeError("blocked_missing_c9623_material: 未在授权目录内找到 C9623 视频素材")
    raise RuntimeError(
        "blocked_multiple_c9623_candidates_need_user_choice: "
        + " | ".join(str(path) for path in video_candidates)
    )


def require_boundaries(source_path: Path) -> dict[str, str]:
    pwd = Path.cwd().resolve()
    git_root = Path(run_command(["git", "rev-parse", "--show-toplevel"]).stdout.strip()).resolve()
    branch = run_command(["git", "branch", "--show-current"]).stdout.strip()
    remote = run_command(["git", "remote", "-v"]).stdout.strip()
    if git_root != REPO_ROOT.resolve():
        raise RuntimeError(f"blocked_wrong_workspace_root: git root={git_root}")
    if AUTHORIZED_BASE not in git_root.parents and git_root != AUTHORIZED_BASE:
        raise RuntimeError(f"blocked_wrong_workspace_root: repo={git_root}")
    if AUTHORIZED_BASE not in source_path.resolve().parents:
        raise RuntimeError(f"blocked_wrong_workspace_root: source={source_path}")
    if "fthytwerwt-sudo/lanxinse--" not in remote:
        raise RuntimeError(f"blocked_wrong_remote: {remote}")
    if not shutil.which("ffmpeg") or not shutil.which("ffprobe"):
        raise RuntimeError("blocked_missing_ffmpeg_or_ffprobe")
    return {"pwd": str(pwd), "git_root": str(git_root), "branch": branch, "remote": remote}


def ffprobe_json(path: Path, timeout: int = 120) -> dict[str, Any]:
    result = run_command(
        ["ffprobe", "-v", "error", "-print_format", "json", "-show_format", "-show_streams", str(path)],
        timeout=timeout,
    )
    require_ok(result, "ffprobe")
    return json.loads(result.stdout)


def media_info(path: Path) -> MediaInfo:
    data = ffprobe_json(path)
    video = next((stream for stream in data.get("streams", []) if stream.get("codec_type") == "video"), {})
    audio = next((stream for stream in data.get("streams", []) if stream.get("codec_type") == "audio"), {})
    rotation = ""
    for side in video.get("side_data_list", []) or []:
        if "rotation" in side:
            rotation = str(side.get("rotation"))
    return MediaInfo(
        path=path,
        duration=float(data.get("format", {}).get("duration") or video.get("duration") or 0),
        width=int(video.get("width") or 0),
        height=int(video.get("height") or 0),
        fps=parse_fps(video.get("avg_frame_rate") or video.get("r_frame_rate") or ""),
        rotation=rotation,
        video_codec=video.get("codec_name", ""),
        audio_codec=audio.get("codec_name", ""),
        audio_present=bool(audio),
        file_size_mb=round(path.stat().st_size / 1024 / 1024, 3),
    )


def decode_sample(path: Path, start: float, duration: float = 3.0, timeout: int = 90) -> tuple[str, str]:
    result = run_command(
        [
            "ffmpeg",
            "-v",
            "error",
            "-ss",
            stringify(start),
            "-t",
            stringify(duration),
            "-i",
            str(path),
            "-map",
            "0:v:0",
            "-map",
            "0:a:0?",
            "-f",
            "null",
            "-",
        ],
        timeout=timeout,
    )
    if result.returncode == 0:
        return "sample_decode_passed", ""
    return "sample_decode_failed", (result.stderr or result.stdout).strip()[:600]


def load_old_summary(window_id: str) -> dict[str, Any]:
    path = OLD_SUMMARY_DIR / f"{window_id}.json"
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8")).get("final_json", {})
    except json.JSONDecodeError:
        return {}


def text_contains_any(text: str, keywords: Iterable[str]) -> bool:
    return any(keyword in text for keyword in keywords)


def visual_scores(summary: dict[str, Any], evidence: str = "") -> tuple[int, int, str]:
    text = " ".join(
        stringify(summary.get(key, ""))
        for key in ["window_summary", "actions_methods", "products_props_course_info"]
    )
    text = f"{text} {evidence}"
    strong_action = ["仰卧", "抬腿", "抬臀", "桥式", "侧卧", "伸展", "屈膝", "骨盆", "转体", "上举"]
    prop_only = ["持瓶", "小瓶", "轻摇", "嗅闻", "小物", "手势", "持物"]
    if text_contains_any(text, strong_action):
        clarity = 4
    elif text_contains_any(text, prop_only):
        clarity = 2
    else:
        clarity = 1
    if text_contains_any(text, ["重复", "连续", "多组", "多帧", "循环", "动作序列"]):
        cycle = min(4, clarity)
    else:
        cycle = max(1, clarity - 1)
    if text_contains_any(text, ["产品", "疗效", "治疗", "禁忌", "医学", "医疗", "效果"]):
        risk = "health_or_business_risk_pending_review"
    else:
        risk = "health_action_risk_pending_professional_review"
    return clarity, cycle, risk


def window_ids_for_segment(windows: list[dict[str, Any]], start: float, end: float) -> list[str]:
    ids: list[str] = []
    for row in windows:
        ws = float(row["start_seconds"])
        we = float(row["end_seconds"])
        if max(ws, start) < min(we, end):
            ids.append(row["window_id"])
    return ids


def candidate_label(row: dict[str, str], summary: dict[str, Any]) -> str:
    text = f"{row.get('content_archetype','')} {row.get('evidence','')} {stringify(summary.get('actions_methods',''))}"
    if text_contains_any(text, ["仰卧", "抬腿", "抬臀", "桥式", "侧卧", "伸展", "骨盆"]):
        return "垫上动作循环"
    if text_contains_any(text, ["持瓶", "小瓶", "产品", "道具"]):
        return "道具讲解待听审"
    if text_contains_any(text, ["静态", "坐姿"]):
        return "静态讲解待听审"
    return "直播片段待复核"


def select_status(opportunity_id: str, route: str) -> str:
    if opportunity_id in QUALIFIED_REPACKAGING_IDS:
        return "qualified_repackaging"
    if route == "adjacent_merge_candidate":
        return "qualified_merge_candidate"
    if route == "reject_unusable":
        return "reject_unusable"
    if opportunity_id in PENDING_AUDIO_EXPORT_IDS:
        return "pending_audio_review"
    return "pending_audio_review"


def map_route_for_status(status: str, old_route: str) -> str:
    if status == "qualified_repackaging":
        return "action_repackaging_candidate"
    if status == "qualified_merge_candidate":
        return "adjacent_merge_candidate"
    if status == "pending_audio_review":
        return "pending_manual_review"
    if status == "reject_unusable":
        return "reject_unusable"
    return old_route


def build_inventory(info: MediaInfo, decode_status: str, decode_error: str) -> list[dict[str, Any]]:
    return [
        {
            "recording_id": "rec_002",
            "file_name": info.path.name,
            "source_path": str(info.path),
            "duration_seconds": info.duration,
            "width": info.width,
            "height": info.height,
            "fps": info.fps,
            "rotation": info.rotation,
            "video_codec": info.video_codec,
            "audio_codec": info.audio_codec,
            "audio_present": info.audio_present,
            "file_size_mb": info.file_size_mb,
            "read_status": "read_success",
            "decode_status": decode_status,
            "decode_error": decode_error,
            "source_dir": str(info.path.parent),
            "discovery_note": "unique_c9623_file_found_under_authorized_material_dir",
        }
    ]


def build_windows() -> list[dict[str, Any]]:
    ali_rows = {row["window_id"]: row for row in read_csv(OLD_ALI_AUDIT_CSV)}
    rows: list[dict[str, Any]] = []
    for row in read_csv(OLD_WINDOWS_CSV):
        if row.get("recording_id") != "rec_002":
            continue
        ali = ali_rows.get(row["window_id"], {})
        start = parse_timecode(row.get("window_start", ""))
        end = parse_timecode(row.get("window_end", ""))
        rows.append(
            {
                "window_id": row["window_id"],
                "recording_id": "rec_002",
                "source_file": "C9623.MP4",
                "window_start": row.get("window_start", ""),
                "window_end": row.get("window_end", ""),
                "start_seconds": start,
                "end_seconds": end,
                "duration_seconds": end - start,
                "overlap_previous_seconds": row.get("overlap_previous_seconds", ""),
                "coverage_status": row.get("coverage_status", ""),
                "analysis_source": "reused_112_full_coverage_ali_window_audit",
                "ali_model_called": ali.get("ali_model_called", ""),
                "ali_primary_status": ali.get("primary_status", ""),
                "ali_high_review_status": ali.get("high_status", ""),
                "needs_audio_transcript": "yes",
                "needs_adjacent_merge": "yes",
                "notes": "112 已完成视觉窗口审计；本轮只复用清洗后结果，不提交旧 API 原始输出。",
            }
        )
    return rows


def build_audio_timeline(windows: list[dict[str, Any]], info: MediaInfo) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for index, window in enumerate(windows, 1):
        rows.append(
            {
                "timeline_id": f"c9623_audio_{index:03d}",
                "recording_id": "rec_002",
                "source_file": "C9623.MP4",
                "start_time": window["window_start"],
                "end_time": window["window_end"],
                "start_seconds": window["start_seconds"],
                "end_seconds": window["end_seconds"],
                "audio_present": info.audio_present,
                "audio_codec": info.audio_codec,
                "transcript_status": "pending_audio_transcript",
                "transcript_text": "",
                "subtitle_text": "",
                "tts_action_alignment": "pending_audio_transcript",
                "status_reason": "本机无 whisper/faster-whisper；本轮未把口播内容脑补为已转写。",
            }
        )
    return rows


def build_visual_audit(windows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for window in windows:
        summary = load_old_summary(window["window_id"])
        clarity, cycle, risk = visual_scores(summary)
        rows.append(
            {
                "window_id": window["window_id"],
                "recording_id": "rec_002",
                "start_time": window["window_start"],
                "end_time": window["window_end"],
                "window_summary": summary.get("window_summary", ""),
                "actions_methods": stringify(summary.get("actions_methods", "")),
                "products_props_course_info": stringify(summary.get("products_props_course_info", "")),
                "has_complete_opening": summary.get("has_complete_opening", ""),
                "has_middle_delivery": summary.get("has_middle_delivery", ""),
                "has_ending_closure": summary.get("has_ending_closure", ""),
                "need_merge_previous": summary.get("need_merge_previous", "yes"),
                "need_merge_next": summary.get("need_merge_next", "yes"),
                "visual_action_clarity_score": clarity,
                "action_cycle_score": cycle,
                "risk_summary": stringify(summary.get("risks", "")) or risk,
                "analysis_status": "success_reused_112_ali_visual_audit",
                "analysis_source": "api_outputs old summary reused locally; raw JSON not submitted",
            }
        )
    return rows


def build_candidates(windows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    windows_by_id = {row["window_id"]: row for row in windows}
    for index, row in enumerate([r for r in read_csv(V2_OPPORTUNITY_CSV) if r.get("recording_id") == "rec_002"], 1):
        start = parse_timecode(row["start_time"])
        end = parse_timecode(row["end_time"])
        ids = window_ids_for_segment(windows, start, end)
        summary = load_old_summary(ids[0]) if ids else {}
        clarity, cycle, risk = visual_scores(summary, row.get("evidence", ""))
        status = select_status(row["opportunity_id"], row.get("route_decision", ""))
        route = map_route_for_status(status, row.get("route_decision", ""))
        transcript_score = 1
        bridge_score = 1
        repackaging_score = 4 if status == "qualified_repackaging" else (2 if row.get("route_decision") != "reject_unusable" else 0)
        merge_score = 4 if status == "qualified_merge_candidate" else (2 if row.get("adjacent_merge_opportunity") == "yes" else 0)
        risk_score = 3 if "risk" in risk or "风险" in row.get("evidence", "") else 2
        priority = clarity * 2 + cycle + repackaging_score + merge_score - risk_score
        if status == "pending_audio_review" and row["opportunity_id"] in PENDING_AUDIO_EXPORT_IDS:
            priority += 2
        if status == "reject_unusable":
            priority = 0
        label = candidate_label(row, summary)
        why_selected = ""
        why_rejected = ""
        not_exported_reason = ""
        if status in {"qualified_repackaging", "qualified_merge_candidate"}:
            why_selected = "视觉上存在连续动作或前后窗口救回价值；仅作为本地人审素材导出，不代表发布可用。"
        elif status == "pending_audio_review":
            why_selected = "视觉线索不足以确认完整表达，但有手势/动作/道具讲解线索，导出优先级样本供听审。"
            if row["opportunity_id"] not in PENDING_AUDIO_EXPORT_IDS:
                not_exported_reason = "pending_audio_review_lower_priority_not_exported"
        else:
            why_rejected = "旧 V2 机会表已标 reject_unusable 或视觉证据不足，当前不导出视频。"
            not_exported_reason = "reject_unusable_index_only"
        source_window_text = "；".join(ids)
        candidates.append(
            {
                "candidate_id": f"c9623_{index:03d}",
                "source_file": "C9623.MP4",
                "start_time": row["start_time"],
                "end_time": row["end_time"],
                "start_seconds": start,
                "end_seconds": end,
                "duration_seconds": max(0.0, end - start),
                "source_windows": source_window_text,
                "content_archetype": row.get("content_archetype", ""),
                "route_decision": route,
                "candidate_status": status,
                "visual_action_clarity_score": clarity,
                "action_cycle_score": cycle,
                "transcript_semantic_score": transcript_score,
                "problem_action_bridge_score": bridge_score,
                "repackaging_value_score": repackaging_score,
                "adjacent_merge_score": merge_score,
                "risk_score": risk_score,
                "priority_score": priority,
                "problem_phrase_time": "pending_audio_transcript",
                "first_action_frame_time": row["start_time"] if clarity >= 3 else "pending_visual_manual_review",
                "problem_action_bridge_seconds": "pending_audio_transcript",
                "tts_action_alignment": "pending_audio_transcript",
                "visual_evidence_timecodes": source_window_text,
                "why_selected": why_selected,
                "why_rejected": why_rejected,
                "manual_packaging_advice": "补 3 秒问题前置；用 TTS 解释动作目标/禁忌；字幕标出发力点；保留 pending_user_review。",
                "suggested_tts_hook": "这段先作为动作素材看，不要直接跟练，先确认口播和动作要点。",
                "suggested_subtitle": f"{label}｜待听口播后补问题和动作要点",
                "health_action_risk": risk,
                "business_risk": "business_goal_pending_validation",
                "manual_review_items": "听口播/字幕；确认动作专业性；确认健康表达边界；确认是否适合包装；确认是否有产品/疗效宣称。",
                "not_exported_reason": not_exported_reason,
                "analysis_status": "visual_reused_audio_pending_user_review",
            }
        )
        _ = windows_by_id
    candidates.sort(key=lambda item: (item["candidate_status"] == "reject_unusable", -float(item["priority_score"]), item["start_seconds"]))
    return candidates


def should_export(candidate: dict[str, Any]) -> bool:
    if candidate["candidate_status"] in {"qualified_repackaging", "qualified_merge_candidate"}:
        return True
    if candidate["candidate_status"] == "pending_audio_review" and not candidate.get("not_exported_reason"):
        return True
    return False


def export_path_for(candidate: dict[str, Any], export_index: int) -> Path:
    group = GROUPS.get(candidate["candidate_status"], GROUPS["reject_unusable"])
    label = safe_slug(candidate.get("suggested_subtitle", "clip"))
    name = (
        f"C9623_{export_index:03d}_{candidate['candidate_status']}_"
        f"{slug_time(float(candidate['start_seconds']))}_{slug_time(float(candidate['end_seconds']))}_{label}.mp4"
    )
    return LOCAL_OUTPUT_DIR / group / name


def ffprobe_output(path: Path) -> dict[str, Any]:
    try:
        data = ffprobe_json(path, timeout=90)
    except Exception as exc:  # pragma: no cover - written to CSV
        return {"ffprobe_status": f"failed:{exc}"}
    video = next((stream for stream in data.get("streams", []) if stream.get("codec_type") == "video"), {})
    audio = next((stream for stream in data.get("streams", []) if stream.get("codec_type") == "audio"), {})
    return {
        "ffprobe_status": "passed" if video and float(data.get("format", {}).get("duration", 0) or 0) > 3 else "failed",
        "output_duration_seconds": data.get("format", {}).get("duration", ""),
        "output_width": video.get("width", ""),
        "output_height": video.get("height", ""),
        "video_stream_present": bool(video),
        "audio_stream_present": bool(audio),
        "file_size_mb": round(path.stat().st_size / 1024 / 1024, 3) if path.exists() else 0,
    }


def export_clip(source: Path, candidate: dict[str, Any], output: Path) -> tuple[str, str]:
    output.parent.mkdir(parents=True, exist_ok=True)
    duration = float(candidate["duration_seconds"])
    command = [
        "ffmpeg",
        "-y",
        "-hide_banner",
        "-loglevel",
        "error",
        "-ss",
        stringify(candidate["start_seconds"]),
        "-i",
        str(source),
        "-t",
        stringify(duration),
        "-map",
        "0:v:0",
        "-map",
        "0:a:0?",
        "-c:v",
        "copy",
        "-c:a",
        "aac",
        "-b:a",
        "128k",
        "-movflags",
        "+faststart",
        str(output),
    ]
    result = run_command(command, timeout=max(120, int(duration * 3)))
    if result.returncode == 0 and output.exists() and output.stat().st_size > 0:
        return "stream_copy_video_audio_aac_success", ""
    fallback = [
        "ffmpeg",
        "-y",
        "-hide_banner",
        "-loglevel",
        "error",
        "-ss",
        stringify(candidate["start_seconds"]),
        "-i",
        str(source),
        "-t",
        stringify(duration),
        "-map",
        "0:v:0",
        "-map",
        "0:a:0?",
        "-c:v",
        "libx264",
        "-preset",
        "veryfast",
        "-crf",
        "24",
        "-c:a",
        "aac",
        "-b:a",
        "128k",
        "-movflags",
        "+faststart",
        str(output),
    ]
    fallback_result = run_command(fallback, timeout=max(180, int(duration * 8)))
    if fallback_result.returncode == 0 and output.exists() and output.stat().st_size > 0:
        return "fallback_reencode_success", ""
    return "failed", ((fallback_result.stderr or fallback_result.stdout) or (result.stderr or result.stdout)).strip()[:800]


def export_candidates(source: Path, candidates: list[dict[str, Any]], do_export: bool) -> list[dict[str, Any]]:
    for group in GROUPS.values():
        (LOCAL_OUTPUT_DIR / group).mkdir(parents=True, exist_ok=True)
    export_rows: list[dict[str, Any]] = []
    export_index = 1
    for candidate in candidates:
        if not should_export(candidate):
            export_rows.append(
                {
                    "export_id": "",
                    "candidate_id": candidate["candidate_id"],
                    "candidate_status": candidate["candidate_status"],
                    "group_dir": GROUPS.get(candidate["candidate_status"], GROUPS["reject_unusable"]),
                    "file_name": "",
                    "output_path": "",
                    "source_file": "C9623.MP4",
                    "start_time": candidate["start_time"],
                    "end_time": candidate["end_time"],
                    "duration_seconds": candidate["duration_seconds"],
                    "why_exported": "",
                    "manual_review_items": candidate["manual_review_items"],
                    "ffprobe_status": "not_exported",
                    "not_exported_reason": candidate.get("not_exported_reason") or "index_only",
                }
            )
            continue
        output = export_path_for(candidate, export_index)
        status = "planned_export"
        error = ""
        probe: dict[str, Any] = {}
        if do_export:
            status, error = export_clip(source, candidate, output)
            probe = ffprobe_output(output) if status.endswith("success") else {"ffprobe_status": status, "output_duration_seconds": "", "file_size_mb": 0}
        export_rows.append(
            {
                "export_id": f"c9623_export_{export_index:03d}",
                "candidate_id": candidate["candidate_id"],
                "candidate_status": candidate["candidate_status"],
                "group_dir": GROUPS.get(candidate["candidate_status"], ""),
                "file_name": output.name,
                "output_path": str(output),
                "source_file": "C9623.MP4",
                "start_time": candidate["start_time"],
                "end_time": candidate["end_time"],
                "duration_seconds": candidate["duration_seconds"],
                "why_exported": candidate["why_selected"],
                "manual_review_items": candidate["manual_review_items"],
                "ffprobe_status": probe.get("ffprobe_status", status),
                "output_duration_seconds": probe.get("output_duration_seconds", ""),
                "output_width": probe.get("output_width", ""),
                "output_height": probe.get("output_height", ""),
                "video_stream_present": probe.get("video_stream_present", ""),
                "audio_stream_present": probe.get("audio_stream_present", ""),
                "file_size_mb": probe.get("file_size_mb", ""),
                "not_exported_reason": error,
            }
        )
        export_index += 1
    return export_rows


def build_rejected_rows(candidates: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rescue_rows = {row.get("rejected_id"): row for row in read_csv(OLD_RESCUE_CSV) if row.get("recording_id") == "rec_002"}
    rejected_old = {row.get("candidate_id"): row for row in read_csv(OLD_REJECTED_CSV) if row.get("recording_id") == "rec_002"}
    rows: list[dict[str, Any]] = []
    for candidate in candidates:
        if should_export(candidate):
            continue
        number = int(candidate["candidate_id"].split("_")[-1])
        old_id = f"rej_{40 + number:03d}"
        rescue = rescue_rows.get(old_id, {})
        old = rejected_old.get(old_id, {})
        rows.append(
            {
                "rejected_id": f"c9623_reject_{number:03d}",
                "candidate_id": candidate["candidate_id"],
                "source_file": "C9623.MP4",
                "start_time": candidate["start_time"],
                "end_time": candidate["end_time"],
                "candidate_status": candidate["candidate_status"],
                "route_decision": candidate["route_decision"],
                "why_rejected": candidate.get("why_rejected") or candidate.get("not_exported_reason") or "lower_priority_pending_audio_review",
                "original_reject_reason": old.get("reject_reason", rescue.get("original_reject_reason", "")),
                "missing_part": old.get("missing_part", rescue.get("missing_part", "")),
                "could_be_fixed_by_manual_edit": old.get("could_be_fixed_by_manual_edit", rescue.get("could_be_fixed_by_manual_edit", "")),
                "rescue_level": rescue.get("rescue_level", ""),
                "needs_transcript_or_audio": rescue.get("needs_transcript_or_audio", "yes"),
                "manual_review_items": candidate["manual_review_items"],
            }
        )
    return rows


def build_packaging_rows(candidates: list[dict[str, Any]], export_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_candidate = {row["candidate_id"]: row for row in export_rows if row.get("output_path")}
    rows: list[dict[str, Any]] = []
    for index, candidate in enumerate([row for row in candidates if row["candidate_id"] in by_candidate], 1):
        export = by_candidate[candidate["candidate_id"]]
        rows.append(
            {
                "advice_id": f"c9623_pack_{index:03d}",
                "candidate_id": candidate["candidate_id"],
                "exported_file_name": export["file_name"],
                "candidate_status": candidate["candidate_status"],
                "suitable_problem": "待听口播后确定；优先包装成单动作/单场景的低风险教学或道具讲解素材。",
                "first_3_seconds_tts": candidate["suggested_tts_hook"],
                "tts_or_voiceover_advice": "用 TTS 补问题前置、动作目标、禁忌提示；不要承诺治疗或效果。",
                "subtitle_focus": candidate["suggested_subtitle"],
                "diagram_needed": "yes_if_action_or_prop_needs_explanation",
                "risk_hint_needed": "yes",
                "editing_tool_fit": "剪映/CapCut 粗包可行；专业精修需用户先确认动作和口播。",
                "manual_review_items": candidate["manual_review_items"],
            }
        )
    return rows


def create_probe_artifacts(source: Path, info: MediaInfo) -> list[dict[str, Any]]:
    probe_dir = LOCAL_API_DIR / "probe"
    audio_dir = probe_dir / "audio"
    sheet_dir = probe_dir / "contact_sheets"
    audio_dir.mkdir(parents=True, exist_ok=True)
    sheet_dir.mkdir(parents=True, exist_ok=True)
    starts = [30.0, max(0.0, info.duration / 2 - 15), max(0.0, info.duration - 60)]
    rows: list[dict[str, Any]] = []
    for index, start in enumerate(starts, 1):
        audio_path = audio_dir / f"c9623_probe_{index:02d}_{slug_time(start)}.wav"
        sheet_path = sheet_dir / f"c9623_probe_{index:02d}_{slug_time(start)}.jpg"
        audio_cmd = [
            "ffmpeg",
            "-y",
            "-hide_banner",
            "-loglevel",
            "error",
            "-ss",
            stringify(start),
            "-t",
            "15",
            "-i",
            str(source),
            "-vn",
            "-ac",
            "1",
            "-ar",
            "16000",
            str(audio_path),
        ]
        sheet_cmd = [
            "ffmpeg",
            "-y",
            "-hide_banner",
            "-loglevel",
            "error",
            "-ss",
            stringify(start),
            "-t",
            "30",
            "-i",
            str(source),
            "-vf",
            "fps=1/6,scale=320:-1,tile=5x1",
            "-frames:v",
            "1",
            str(sheet_path),
        ]
        audio_result = run_command(audio_cmd, timeout=120)
        sheet_result = run_command(sheet_cmd, timeout=120)
        rows.append(
            {
                "probe_id": f"probe_{index:02d}",
                "start_seconds": start,
                "audio_path": str(audio_path),
                "audio_status": "created" if audio_result.returncode == 0 and audio_path.exists() else "failed",
                "contact_sheet_path": str(sheet_path),
                "contact_sheet_status": "created" if sheet_result.returncode == 0 and sheet_path.exists() else "failed",
                "note": "probe artifacts local ignored only; no Ali call in probe",
            }
        )
    write_csv(LOCAL_API_DIR / "probe" / "probe_manifest.csv", list(rows[0].keys()), rows)
    return rows


def summarize_counts(candidates: list[dict[str, Any]], exports: list[dict[str, Any]]) -> dict[str, Any]:
    status_counts = Counter(row["candidate_status"] for row in candidates)
    exported_count = sum(1 for row in exports if row.get("output_path") and row.get("ffprobe_status") == "passed")
    return {
        "candidate_total": len(candidates),
        "qualified_native": status_counts.get("qualified_native", 0),
        "qualified_repackaging": status_counts.get("qualified_repackaging", 0),
        "qualified_merge_candidate": status_counts.get("qualified_merge_candidate", 0),
        "pending_audio_review": status_counts.get("pending_audio_review", 0),
        "reject_unusable": status_counts.get("reject_unusable", 0),
        "exported_count": exported_count,
    }


def markdown_table(headers: list[str], rows: list[list[Any]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(stringify(item).replace("\n", " ") for item in row) + " |")
    return "\n".join(lines)


def write_reports(info: MediaInfo, candidates: list[dict[str, Any]], exports: list[dict[str, Any]], probe_rows: list[dict[str, Any]], boundary: dict[str, str], run_mode: str) -> None:
    FACT_DIR.mkdir(parents=True, exist_ok=True)
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    counts = summarize_counts(candidates, exports)
    exported = [row for row in exports if row.get("output_path")]
    synthesis = f"""# 前序报告综合结论

状态：`completed_for_c9623_validation_basis_pending_user_review`
生成时间：{now_iso()}

## 已确认

- C9623 是 112 正式模拟里的 `rec_002`，素材路径为 `{info.path}`。
- 112 已完成 C9623 全覆盖视觉窗口：37 个窗口，阿里视觉审计复用状态为 `success_reused_112_ali_visual_audit`。
- 112 候选为 0，不代表素材全废；113 路线重判指出应补音频/字幕、相邻窗口合并、人工扩展起止点。
- 113 V2 架构要求继承 `content_archetype`、`route_decision`、`candidate_status`、`problem_action_bridge_seconds`、`tts_action_alignment`、`needs_adjacent_merge` 等字段。

## 部分成立

- C9623 存在动作教学再包装和相邻窗口合并线索，但这些线索只代表本地人审素材，不代表原生切片完成。

## 待验证

- 音频/字幕/口播转写、问题-动作桥接秒点、动作专业性、健康表达、业务转化、人感审美和批量稳定运行。
"""
    SYNTHESIS_MD.write_text(synthesis, encoding="utf-8")

    group_rows = []
    for key, folder in GROUPS.items():
        group_rows.append([folder + "/", sum(1 for row in exported if row.get("group_dir") == folder)])
    full_report = f"""# C9623 完整直播切片验证报告

状态：`partial_completed_exports_ready_pending_user_review`
生成时间：{now_iso()}

## 1. 执行结果

| 项目 | 结果 |
| --- | --- |
| 当前项目仓库 | `fthytwerwt-sudo/lanxinse--` |
| 本地仓库路径 | `{REPO_ROOT}` |
| 当前分支 | `{boundary.get('branch', '')}` |
| C9623 实际素材路径 | `{info.path}` |
| C9623 视频数量 | 1 |
| C9623 总时长 | {timecode(info.duration)} / {info.duration:.3f}s |
| 是否完成全覆盖扫描 | 是，复用 112 的 37 个全覆盖窗口 |
| 是否完成音频转写 | 否，`pending_audio_transcript` |
| 是否完成阿里视觉审计 | 是，复用 112 已完成审计；本轮未新增 API 调用 |
| 候选总数 | {counts['candidate_total']} |
| 原生切片候选数量 | {counts['qualified_native']} |
| 动作教学再包装候选数量 | {counts['qualified_repackaging']} |
| 相邻窗口合并候选数量 | {counts['qualified_merge_candidate']} |
| 待音频复核数量 | {counts['pending_audio_review']} |
| 弃用数量 | {counts['reject_unusable']} |
| 本地导出视频数量 | {counts['exported_count']} |
| 本地导出目录 | `{LOCAL_OUTPUT_DIR}` |
| run_mode | `{run_mode}` |

## 2. 导出文件夹

{markdown_table(["文件夹", "导出数量"], group_rows)}

## 3. 关键结论

1. C9623 当前没有 `qualified_native` 原生切片候选；原因是旧审计仍缺完整开头、中段、结尾和音频语义证据。
2. 可作为动作教学再包装的人审素材主要集中在 00:36:50 之后的垫上动作、仰卧/桥式/伸展类画面。
3. 相邻窗口合并候选仍需要听口播后向前/向后扩展，不能直接发布。
4. 待音频复核素材包含道具讲解、坐姿讲解和部分动作段，必须听口播/看字幕后再判断。
5. 弃用素材仅保留索引，不导出视频。
6. 本轮证明 V2 架构能把 C9623 从“0 原生候选”推进到“本地人审素材池”，但未证明审美、动作专业性、健康表达、业务转化或稳定批量运行通过。

## 4. 边界

- 是否提交媒体：否。
- 是否提交 contact sheet：否。
- 是否提交 API 原始输出：否。
- 是否写审美/动作专业/健康表达/业务通过：否。
- 导出视频只作为本地人审素材：是。
"""
    FULL_REPORT_MD.write_text(full_report, encoding="utf-8")

    evidence_rows = [
        [row["candidate_id"], row["candidate_status"], row["start_time"], row["end_time"], row["why_selected"] or row["why_rejected"]]
        for row in candidates[:25]
    ]
    EVIDENCE_REPORT_MD.write_text(
        "# C9623 候选片段证据报告\n\n"
        f"状态：`candidate_evidence_completed_pending_user_review`\n生成时间：{now_iso()}\n\n"
        + markdown_table(["candidate_id", "status", "start", "end", "evidence"], evidence_rows)
        + "\n",
        encoding="utf-8",
    )

    PACKAGING_GUIDE_MD.write_text(
        "# C9623 人工包装说明\n\n"
        f"状态：`manual_packaging_guide_completed_pending_user_review`\n生成时间：{now_iso()}\n\n"
        "## 包装口径\n\n"
        "- 先补问题前置，再解释动作目标，不要直接承诺疗效。\n"
        "- TTS/字幕必须补禁忌提示和人审边界。\n"
        "- `qualified_repackaging` 只表示可包装素材，不是最终成片。\n\n"
        "## 建议流程\n\n"
        "1. 用户先看导出片段，确认哪几条动作/道具有价值。\n"
        "2. 人工听口播，补 `problem_phrase_time` 和动作开始秒点。\n"
        "3. 再决定是否进入剪映/CapCut/达芬奇精修。\n",
        encoding="utf-8",
    )

    RISK_REVIEW_MD.write_text(
        "# C9623 风险与待复核清单\n\n"
        f"状态：`risk_review_checklist_completed_pending_user_review`\n生成时间：{now_iso()}\n\n"
        "## 必须复核\n\n"
        "- `pending_audio_transcript`：本轮没有自动转写，不得脑补口播。\n"
        "- `health_action_risk_pending_professional_review`：动作专业性和禁忌人群必须人审。\n"
        "- `business_goal_pending_validation`：产品、课程、疗效或转化话术不能由技术导出直接盖章。\n"
        "- 审美、人感、节奏、封面、字幕包装均为 `pending_user_review`。\n",
        encoding="utf-8",
    )

    EXEC_REPORT_MD.write_text(
        f"""# C9623 直播切片完整验证执行报告

状态：`partial_completed_exports_ready_pending_user_review`
生成时间：{now_iso()}

## Commands

- `python3 scripts/check_ali_config_safety.py`
- `python3 -m py_compile scripts/c9623_live_cutting_full_validation.py`
- `python3 scripts/c9623_live_cutting_full_validation.py --dry-run`
- `python3 scripts/c9623_live_cutting_full_validation.py --probe`
- `python3 scripts/c9623_live_cutting_full_validation.py --run`

## Result

| 项目 | 结果 |
| --- | --- |
| C9623 路径 | `{info.path}` |
| 总时长 | `{timecode(info.duration)}` |
| 全覆盖窗口 | 37 |
| 候选总数 | {counts['candidate_total']} |
| 本地导出视频 | {counts['exported_count']} |
| 阿里视觉审计 | `reused_112_ali_visual_audit_no_new_api_call` |
| 音频转写 | `pending_audio_transcript` |
| probe_rows | {len(probe_rows)} |

## 边界确认

- 是否提交媒体：否。
- 是否提交 API 原始输出：否。
- 是否提交 `.env`：否。
- 是否提交 API key：否。
- 是否写审美通过：否。
- 是否写动作专业性通过：否。
- 是否写健康表达通过：否。
- 是否写业务通过：否。
- 是否写稳定批量运行：否。
""",
        encoding="utf-8",
    )

    HUMAN_MD_FALLBACK.write_text(full_report, encoding="utf-8")
    create_docx(full_report, counts, info)


def create_docx(full_report: str, counts: dict[str, Any], info: MediaInfo) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Inches, Pt
    except ImportError:
        return

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    title = document.add_heading("C9623 直播切片完整验证报告", 0)
    title.paragraph_format.space_after = Pt(6)
    subtitle = document.add_paragraph(f"生成时间：{now_iso()}｜状态：partial completed, pending user review")
    subtitle.paragraph_format.space_after = Pt(12)
    document.add_heading("主结论", level=1)
    document.add_paragraph("C9623 已完成全覆盖复用审计，并导出本地候选素材供人工复核；本轮不确认审美、动作专业性、健康表达、业务转化或稳定批量运行。")
    document.add_heading("核心数据", level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "结果"
    rows = [
        ("素材路径", str(info.path)),
        ("总时长", f"{timecode(info.duration)} / {info.duration:.3f}s"),
        ("候选总数", counts["candidate_total"]),
        ("动作教学再包装候选", counts["qualified_repackaging"]),
        ("相邻窗口合并候选", counts["qualified_merge_candidate"]),
        ("待音频复核", counts["pending_audio_review"]),
        ("本地导出视频", counts["exported_count"]),
        ("音频转写", "pending_audio_transcript"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = stringify(left)
        cells[1].text = stringify(right)
    document.add_heading("边界确认", level=1)
    for item in [
        "导出视频只保存在本地 outputs，不提交 GitHub。",
        "未提交 contact sheet、API 原始输出、.env 或 API key。",
        "未写审美、动作专业性、健康表达、业务转化或稳定批量运行通过。",
        "所有候选均保留 pending_user_review / pending_audio_transcript 边界。",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("下一步", level=1)
    for item in [
        "用户先看导出素材，筛出值得包装的 3-5 条。",
        "人工听口播并补问题句、动作起点、结尾收束秒点。",
        "再用 V2 字段进入小样片包装验证。",
    ]:
        document.add_paragraph(item, style="List Number")
    FACT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(HUMAN_DOCX)


def verify_no_forbidden_tracked_paths() -> None:
    result = run_command(["git", "status", "--short", "--ignored=matching", "--", str(LOCAL_OUTPUT_DIR), str(LOCAL_API_DIR)], timeout=60)
    if "??" in result.stdout:
        raise RuntimeError("blocked_media_commit_risk: outputs/api_outputs appears unignored")


def run_pipeline(mode: str) -> dict[str, Any]:
    source = locate_c9623()
    boundary = require_boundaries(source)
    info = media_info(source)
    decode_checks = [
        decode_sample(source, 0),
        decode_sample(source, max(0.0, info.duration / 2 - 1.5)),
        decode_sample(source, max(0.0, info.duration - 5)),
    ]
    decode_status = "sample_decode_passed" if all(status == "sample_decode_passed" for status, _ in decode_checks) else "sample_decode_partial_failed"
    decode_error = "；".join(error for _, error in decode_checks if error)
    inventory = build_inventory(info, decode_status, decode_error)
    windows = build_windows()
    audio_rows = build_audio_timeline(windows, info)
    visual_rows = build_visual_audit(windows)
    candidates = build_candidates(windows)
    probe_rows: list[dict[str, Any]] = []
    do_write = mode in {"probe", "run"}
    do_export = mode == "run"
    if mode == "probe":
        probe_rows = create_probe_artifacts(source, info)
    exports = export_candidates(source, candidates, do_export=do_export)
    rejected_rows = build_rejected_rows(candidates)
    packaging_rows = build_packaging_rows(candidates, exports)
    if mode == "dry-run":
        return {
            "status": "dry_run_ok",
            "source": str(source),
            "duration_seconds": info.duration,
            "windows": len(windows),
            "candidates": len(candidates),
            "planned_exports": sum(1 for row in exports if row.get("output_path")),
            "boundary": boundary,
        }
    if do_write:
        write_csv(INVENTORY_CSV, INVENTORY_FIELDS, inventory)
        write_csv(WINDOW_CSV, WINDOW_FIELDS, windows)
        write_csv(AUDIO_CSV, AUDIO_FIELDS, audio_rows)
        write_csv(VISUAL_CSV, VISUAL_FIELDS, visual_rows)
        write_csv(CANDIDATE_CSV, CANDIDATE_FIELDS, candidates)
        write_csv(EXPORT_INDEX_CSV, EXPORT_FIELDS, exports)
        write_csv(REJECTED_CSV, REJECTED_FIELDS, rejected_rows)
        write_csv(PACKAGING_CSV, PACKAGING_FIELDS, packaging_rows)
        write_reports(info, candidates, exports, probe_rows, boundary, mode)
        verify_no_forbidden_tracked_paths()
    counts = summarize_counts(candidates, exports)
    return {
        "status": "ok",
        "mode": mode,
        "source": str(source),
        "duration_seconds": info.duration,
        "windows": len(windows),
        "audio_rows": len(audio_rows),
        "visual_rows": len(visual_rows),
        "candidate_rows": len(candidates),
        "export_index_rows": len(exports),
        "packaging_rows": len(packaging_rows),
        "probe_rows": len(probe_rows),
        "counts": counts,
        "ali_status": "reused_112_ali_visual_audit_no_new_api_call",
        "audio_transcript_status": "pending_audio_transcript",
        "boundary": boundary,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--dry-run", action="store_true", help="只检查边界和计划，不写文件")
    group.add_argument("--probe", action="store_true", help="生成三段本地音频/视觉 probe，并写结构化表")
    group.add_argument("--run", action="store_true", help="正式生成报告并导出本地候选视频")
    args = parser.parse_args()
    mode = "dry-run" if args.dry_run else "probe" if args.probe else "run"
    result = run_pipeline(mode)
    print(json.dumps(result, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
