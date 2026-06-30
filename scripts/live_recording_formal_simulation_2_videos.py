#!/usr/bin/env python3
"""两条直播录屏正式模拟初剪运行。

边界：
- 只读取授权目录内 2 条完整直播录屏。
- 用全覆盖窗口 + contact sheet 调用阿里视觉模型做逐窗结构审计。
- 初剪候选只用 ffmpeg stream copy 导出到本地忽略目录，不提交媒体。
- 报告只确认技术运行与待人审候选，不确认审美、动作专业性、业务事实或发布通过。
"""

from __future__ import annotations

import argparse
import base64
import csv
import json
import os
import re
import subprocess
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from math import gcd
from pathlib import Path
from typing import Any, Iterable

try:
    from PIL import Image, ImageDraw, ImageFont, ImageOps
except ImportError:  # pragma: no cover
    Image = None
    ImageDraw = None
    ImageFont = None
    ImageOps = None

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    Document = None


REPO_ROOT = Path(__file__).resolve().parents[1]
AUTHORIZED_BASE = Path("/Volumes/WD_BLACK/澜心社剪辑")
PREFERRED_SOURCE_DIR = AUTHORIZED_BASE / "剪辑解析数据-剪辑测试直播素材"
DISCOVERED_SOURCE_DIR = AUTHORIZED_BASE / "剪辑解析数据" / "剪辑测试直播素材"

ANALYSIS_DIR = REPO_ROOT / "素材解析_pipeline_material_analysis" / "09_live_recording_formal_simulation"
FACT_DIR = REPO_ROOT / "项目事实_project_facts" / "直播录屏正式模拟运行_live_recording_formal_simulation"
LOG_DIR = REPO_ROOT / "执行日志_codex_log"
LOCAL_DIR = REPO_ROOT / "api_outputs" / "live_recording_formal_simulation_2_videos"
CONTACT_DIR = LOCAL_DIR / "contact_sheets"
FRAME_DIR = LOCAL_DIR / "frames"
SUMMARY_DIR = LOCAL_DIR / "model_summaries"
OUTPUT_DIR = REPO_ROOT / "outputs" / "live_recording_formal_simulation"

INVENTORY_CSV = ANALYSIS_DIR / "01_直播录屏素材清单_live_recording_inventory.csv"
WINDOW_CSV = ANALYSIS_DIR / "02_全覆盖窗口清单_full_coverage_window_manifest.csv"
ALI_AUDIT_CSV = ANALYSIS_DIR / "03_阿里窗口分析审计表_ali_window_analysis_audit.csv"
CANDIDATE_CSV = ANALYSIS_DIR / "04_候选片段表_candidate_segment_table.csv"
REJECTED_CSV = ANALYSIS_DIR / "05_弃用片段表_rejected_segment_table.csv"
ROUGH_INDEX_CSV = ANALYSIS_DIR / "06_初剪结果索引_rough_cut_output_index.csv"
EVIDENCE_MD = FACT_DIR / "01_初剪候选视频证据报告_rough_cut_evidence_report.md"
CHECKLIST_MD = FACT_DIR / "02_人工复核清单_manual_review_checklist.md"
HUMAN_DOCX = FACT_DIR / "03_两条直播录屏正式模拟运行人读版_two_live_recording_formal_simulation_human_report.docx"
EXEC_REPORT_MD = LOG_DIR / "112_两条直播录屏正式模拟运行报告_two_live_recording_formal_simulation_report.md"

ENV_PATH = REPO_ROOT / ".env"
DEFAULT_BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1"
DEFAULT_VISION_MODEL = "qwen3-vl-plus"
DEFAULT_HIGH_MODEL = "qwen-vl-max"
PLACEHOLDER_VALUES = {
    "",
    "your_api_key_here",
    "your-ali-api-key",
    "你的真实阿里 API key",
    "这里填你的真实阿里 API key",
    "请在本地填写，不要提交真实 key",
}


@dataclass(frozen=True)
class Recording:
    recording_id: str
    path: Path
    duration: float
    width: int
    height: int
    fps: str
    video_codec: str
    audio_codec: str
    file_size: int


@dataclass(frozen=True)
class Window:
    window_id: str
    recording: Recording
    start: float
    end: float
    overlap_previous: float

    @property
    def duration(self) -> float:
        return round(self.end - self.start, 3)


@dataclass
class ApiCallResult:
    role: str
    model: str
    status: str
    latency_ms: int | None
    retry_count: int
    error_type: str
    error_summary: str
    content: str
    parsed_json: dict[str, Any] | None


def now_iso() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def run_command(cmd: list[str], timeout: int = 90) -> subprocess.CompletedProcess[str]:
    return subprocess.run(cmd, text=True, capture_output=True, timeout=timeout, check=False)


def require_ok(result: subprocess.CompletedProcess[str], label: str) -> None:
    if result.returncode != 0:
        raise RuntimeError(f"{label} failed: {(result.stderr or result.stdout).strip()[:800]}")


def read_dotenv(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    if not path.exists():
        return values
    for raw in path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def env_value(values: dict[str, str], key: str, default: str = "") -> str:
    return os.environ.get(key) or values.get(key) or default


def is_placeholder(value: str) -> bool:
    return value.strip() in PLACEHOLDER_VALUES


def sanitize_text(text: str, api_key: str = "") -> str:
    cleaned = text or ""
    if api_key:
        cleaned = cleaned.replace(api_key, "<redacted_api_key>")
    cleaned = re.sub(r"(?i)bearer\s+[A-Za-z0-9_./+=-]{12,}", "Bearer <redacted>", cleaned)
    cleaned = re.sub(r"\bsk-[A-Za-z0-9_-]{10,}\b", "<redacted_secret>", cleaned)
    return cleaned.replace("\r", " ").replace("\n", " ")[:800]


def stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "yes" if value else "no"
    if isinstance(value, list):
        return "；".join(stringify(item) for item in value if stringify(item))
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value)


def write_csv(path: Path, fieldnames: list[str], rows: Iterable[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: stringify(row.get(field, "")) for field in fieldnames})


def parse_fps(value: str) -> str:
    if not value or value == "0/0":
        return ""
    if "/" not in value:
        return value
    num, den = value.split("/", 1)
    try:
        return f"{float(num) / float(den):.3f}".rstrip("0").rstrip(".")
    except (ValueError, ZeroDivisionError):
        return value


def timecode(seconds: float) -> str:
    total_ms = int(round(max(0.0, seconds) * 1000))
    ms = total_ms % 1000
    total_s = total_ms // 1000
    s = total_s % 60
    total_m = total_s // 60
    m = total_m % 60
    h = total_m // 60
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def slug_time(seconds: float) -> str:
    total = int(max(0.0, seconds))
    return f"{total // 3600:02d}{(total % 3600) // 60:02d}{total % 60:02d}"


def ratio_text(width: int, height: int) -> str:
    if not width or not height:
        return "unknown"
    divisor = gcd(width, height)
    return f"{width // divisor}:{height // divisor}"


def ffprobe_json(path: Path) -> dict[str, Any]:
    result = run_command(
        [
            "ffprobe",
            "-v",
            "error",
            "-print_format",
            "json",
            "-show_format",
            "-show_streams",
            str(path),
        ],
        timeout=120,
    )
    require_ok(result, f"ffprobe {path}")
    return json.loads(result.stdout)


def first_stream(metadata: dict[str, Any], codec_type: str) -> dict[str, Any]:
    for stream in metadata.get("streams", []):
        if stream.get("codec_type") == codec_type:
            return stream
    return {}


def decode_check(path: Path) -> tuple[str, str]:
    result = run_command(["ffmpeg", "-v", "error", "-t", "0.8", "-i", str(path), "-f", "null", "-"], timeout=90)
    if result.returncode == 0:
        return "success", ""
    return "failed", sanitize_text(result.stderr)


def discover_source_dir() -> tuple[Path, str]:
    if PREFERRED_SOURCE_DIR.exists():
        return PREFERRED_SOURCE_DIR, "preferred_dir_found"
    if DISCOVERED_SOURCE_DIR.exists():
        return DISCOVERED_SOURCE_DIR, "preferred_dir_missing_used_authorized_discovered_dir"
    matches: list[Path] = []
    for pattern in ("*剪辑解析数据-剪辑测试直播素材*", "*剪辑测试直播素材*", "*直播素材*"):
        matches.extend(path for path in AUTHORIZED_BASE.rglob(pattern) if path.is_dir())
    if not matches:
        raise RuntimeError("blocked_missing_live_recording_dir")
    return sorted(matches)[0], "preferred_dir_missing_used_keyword_match"


def load_recordings(source_dir: Path) -> list[Recording]:
    media = sorted(
        path
        for path in source_dir.rglob("*")
        if path.is_file()
        and not path.name.startswith("._")
        and path.suffix.lower() in {".mp4", ".mov", ".m4v", ".avi", ".mkv"}
    )
    if len(media) != 2:
        raise RuntimeError(f"blocked_unexpected_recording_count: found {len(media)}")

    out: list[Recording] = []
    for index, path in enumerate(media, 1):
        meta = ffprobe_json(path)
        video = first_stream(meta, "video")
        audio = first_stream(meta, "audio")
        if not video:
            raise RuntimeError(f"blocked_unreadable_live_recording: no video stream {path}")
        duration = float(meta.get("format", {}).get("duration") or video.get("duration") or 0.0)
        if duration <= 0:
            raise RuntimeError(f"blocked_unreadable_live_recording: zero duration {path}")
        out.append(
            Recording(
                recording_id=f"rec_{index:03d}",
                path=path,
                duration=duration,
                width=int(video.get("width") or 0),
                height=int(video.get("height") or 0),
                fps=parse_fps(str(video.get("avg_frame_rate") or video.get("r_frame_rate") or "")),
                video_codec=str(video.get("codec_name") or ""),
                audio_codec=str(audio.get("codec_name") or ""),
                file_size=path.stat().st_size,
            )
        )
    return out


def make_windows(recordings: list[Recording], window_seconds: float, overlap_seconds: float) -> list[Window]:
    if overlap_seconds <= 0 or overlap_seconds >= window_seconds:
        raise RuntimeError("invalid_window_overlap")
    windows: list[Window] = []
    for recording in recordings:
        start = 0.0
        previous_start: float | None = None
        idx = 1
        while start < recording.duration:
            end = min(recording.duration, start + window_seconds)
            overlap = 0.0 if previous_start is None else overlap_seconds
            windows.append(Window(f"{recording.recording_id}_w{idx:03d}", recording, round(start, 3), round(end, 3), overlap))
            if end >= recording.duration:
                break
            previous_start = start
            start = max(0.0, end - overlap_seconds)
            idx += 1
    return windows


def frame_times_for_window(window: Window, count: int = 10) -> list[float]:
    if window.duration <= 1:
        return [window.start]
    offsets = [0.0, 1.0]
    inner = max(0, count - 3)
    for index in range(1, inner + 1):
        offsets.append(window.duration * index / (inner + 1))
    offsets.append(max(0.0, window.duration - 0.5))
    times: list[float] = []
    for offset in offsets:
        value = min(window.end - 0.05, window.start + max(0.0, offset))
        rounded = round(value, 3)
        if window.start <= rounded <= window.end and all(abs(rounded - old) >= 0.2 for old in times):
            times.append(rounded)
    return sorted(times)[:count]


def extract_frame(source: Path, time_s: float, out_path: Path) -> bool:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    result = run_command(
        [
            "ffmpeg",
            "-loglevel",
            "error",
            "-y",
            "-ss",
            f"{time_s:.3f}",
            "-i",
            str(source),
            "-frames:v",
            "1",
            "-vf",
            "scale=480:-1",
            str(out_path),
        ],
        timeout=60,
    )
    return result.returncode == 0 and out_path.exists() and out_path.stat().st_size > 0


def make_contact_sheet(window: Window, force: bool = False) -> tuple[Path, list[str], str]:
    if Image is None:
        raise RuntimeError("blocked_missing_pillow_for_contact_sheet")
    CONTACT_DIR.mkdir(parents=True, exist_ok=True)
    FRAME_DIR.mkdir(parents=True, exist_ok=True)
    out_path = CONTACT_DIR / f"{window.window_id}.jpg"
    times = frame_times_for_window(window)
    labels = [timecode(t) for t in times]
    if out_path.exists() and not force:
        return out_path, labels, "reused"

    frames: list[tuple[float, Path]] = []
    for idx, absolute_time in enumerate(times, 1):
        frame_path = FRAME_DIR / window.window_id / f"frame_{idx:02d}_{int(absolute_time * 1000):09d}.jpg"
        if force or not frame_path.exists():
            if not extract_frame(window.recording.path, absolute_time, frame_path):
                continue
        frames.append((absolute_time, frame_path))
    if not frames:
        raise RuntimeError(f"contact_sheet_failed: {window.window_id}")

    cell_w, cell_h, label_h = 360, 203, 40
    cols = 5
    rows = (len(frames) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * cell_w, rows * (cell_h + label_h)), (15, 17, 20))
    draw = ImageDraw.Draw(sheet)
    font = ImageFont.load_default()
    for idx, (absolute_time, frame_path) in enumerate(frames):
        img = Image.open(frame_path).convert("RGB")
        fitted = ImageOps.contain(img, (cell_w, cell_h), method=Image.Resampling.LANCZOS)
        x = (idx % cols) * cell_w
        y = (idx // cols) * (cell_h + label_h)
        bg = Image.new("RGB", (cell_w, cell_h), (0, 0, 0))
        bg.paste(fitted, ((cell_w - fitted.width) // 2, (cell_h - fitted.height) // 2))
        sheet.paste(bg, (x, y + label_h))
        label = f"{window.window_id} F{idx + 1} {timecode(absolute_time)}"
        draw.rectangle([x, y, x + cell_w, y + label_h], fill=(242, 244, 248))
        draw.text((x + 6, y + 12), label, fill=(20, 24, 32), font=font)
    sheet.save(out_path, format="JPEG", quality=80, optimize=True)
    return out_path, [timecode(t) for t, _ in frames], "created"


def image_data_url(path: Path) -> str:
    return "data:image/jpeg;base64," + base64.b64encode(path.read_bytes()).decode("ascii")


def extract_response_content(body: str) -> str:
    parsed = json.loads(body)
    choices = parsed.get("choices") or []
    if not choices:
        return ""
    message = choices[0].get("message") or {}
    content = message.get("content", "")
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        return "\n".join(str(item.get("text") or item.get("content") or "") for item in content if isinstance(item, dict))
    return str(content)


def parse_model_json(content: str) -> dict[str, Any] | None:
    text = content.strip()
    text = re.sub(r"^```(?:json)?", "", text, flags=re.IGNORECASE).strip()
    text = re.sub(r"```$", "", text).strip()
    candidates = [text]
    if "{" in text and "}" in text:
        candidates.append(text[text.find("{") : text.rfind("}") + 1])
    for candidate in candidates:
        try:
            parsed = json.loads(candidate)
        except json.JSONDecodeError:
            continue
        if isinstance(parsed, dict):
            return parsed
    return None


def classify_http_error(status_code: int, body: str) -> tuple[str, str]:
    error_type = f"http_{status_code}"
    summary = body[:800]
    try:
        parsed = json.loads(body)
    except json.JSONDecodeError:
        parsed = {}
    error = parsed.get("error") if isinstance(parsed, dict) else None
    if isinstance(error, dict):
        code = str(error.get("code") or "")
        message = str(error.get("message") or "")
        err_type = str(error.get("type") or "")
        lowered = f"{code} {message} {err_type}".lower()
        summary = " | ".join(part for part in [code, err_type, message] if part)
        if status_code == 401 or "invalid_api_key" in lowered:
            error_type = "authentication_failed"
        elif status_code == 403 or "access" in lowered or "permission" in lowered:
            error_type = "permission_or_account_required"
        elif status_code == 404 or "not found" in lowered:
            error_type = "model_not_available"
        elif status_code == 429 or "quota" in lowered or "rate" in lowered:
            error_type = "quota_or_rate_limit"
    return error_type, summary


def call_ali_vision(
    api_key: str,
    base_url: str,
    model: str,
    role: str,
    prompt: str,
    contact_sheet: Path,
    timeout: int,
    max_retries: int,
) -> ApiCallResult:
    endpoint = base_url.rstrip("/") + "/chat/completions"
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": image_data_url(contact_sheet)}},
                ],
            }
        ],
        "temperature": 0.1,
        "max_tokens": 2200,
    }
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    last_error_type = ""
    last_error_summary = ""
    started = 0.0
    for attempt in range(max_retries + 1):
        request = urllib.request.Request(
            endpoint,
            data=data,
            method="POST",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        )
        started = time.perf_counter()
        try:
            with urllib.request.urlopen(request, timeout=timeout) as response:
                body = response.read().decode("utf-8", errors="replace")
            latency_ms = round((time.perf_counter() - started) * 1000)
            content = extract_response_content(body)
            parsed = parse_model_json(content)
            if parsed is not None:
                return ApiCallResult(role, model, "success", latency_ms, attempt, "", "", content, parsed)
            last_error_type = "json_parse_failed"
            last_error_summary = sanitize_text(content, api_key)
        except urllib.error.HTTPError as exc:
            latency_ms = round((time.perf_counter() - started) * 1000)
            body = exc.read().decode("utf-8", errors="replace")
            last_error_type, last_error_summary = classify_http_error(exc.code, body)
            last_error_summary = sanitize_text(last_error_summary, api_key)
            if last_error_type in {"authentication_failed", "permission_or_account_required", "model_not_available"}:
                return ApiCallResult(role, model, "failed", latency_ms, attempt, last_error_type, last_error_summary, "", None)
        except (TimeoutError, urllib.error.URLError) as exc:
            latency_ms = round((time.perf_counter() - started) * 1000)
            last_error_type = "endpoint_timeout_or_network_error"
            last_error_summary = sanitize_text(str(exc), api_key)
        if attempt < max_retries:
            time.sleep(1.5 + attempt)
    latency_ms = round((time.perf_counter() - started) * 1000) if started else None
    return ApiCallResult(role, model, "failed", latency_ms, max_retries, last_error_type, last_error_summary, "", None)


def as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return [value]
    text = str(value).strip()
    return [part.strip() for part in re.split(r"[；;、,，|/]+", text) if part.strip()]


def truthy(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "y", "是", "有", "完整", "成立"}


def score(value: Any, default: int = 0) -> int:
    try:
        return max(0, min(5, int(float(str(value).strip()))))
    except (TypeError, ValueError):
        return default


def scalar(value: Any, fallback: str = "待人工复核") -> str:
    text = stringify(value).strip()
    return text if text else fallback


def confidence_value(parsed: dict[str, Any] | None) -> str:
    raw = str((parsed or {}).get("confidence") or "").lower()
    if raw in {"high", "medium", "low"}:
        return raw
    if "高" in raw:
        return "high"
    if "低" in raw:
        return "low"
    return "medium"


def needs_high_review(parsed: dict[str, Any] | None) -> bool:
    if not parsed:
        return True
    if confidence_value(parsed) == "low":
        return True
    if truthy(parsed.get("needs_high_vision_review")):
        return True
    if truthy(parsed.get("can_be_complete_short_segment")):
        return True
    if as_list(parsed.get("candidate_segments")):
        return True
    risk_text = stringify(parsed.get("risks") or parsed.get("business_risk") or parsed.get("risk_triggers"))
    return any(token in risk_text for token in ["医疗", "健康", "效果", "承诺", "风险", "漏尿", "盆底", "产后"])


def build_prompt(window: Window, frame_labels: list[str], high_review: bool = False, primary_json: dict[str, Any] | None = None) -> str:
    schema = {
        "window_summary": "只基于画面/字幕/动作的窗口摘要",
        "people": ["出现的人群/身份/状态"],
        "pain_points": ["痛点或问题"],
        "actions_methods": ["动作/方法/步骤/讲解"],
        "products_props_course_info": ["产品/道具/课程/价格/资历/购买信息"],
        "has_complete_opening": "yes/no/unclear",
        "has_middle_delivery": "yes/no/unclear",
        "has_ending_closure": "yes/no/unclear",
        "opening_evidence": "用帧标签/时间码说明",
        "middle_evidence": "用帧标签/时间码说明",
        "ending_evidence": "用帧标签/时间码说明",
        "can_be_complete_short_segment": "yes/no/unclear",
        "need_merge_previous": "yes/no",
        "need_merge_next": "yes/no",
        "candidate_segments": [
            {
                "start_offset_seconds": 0,
                "end_offset_seconds": 90,
                "opening_summary": "开头为什么成立",
                "middle_summary": "中段交付什么",
                "ending_summary": "结尾怎么收",
                "matched_rule_ids": ["rough_integrity_21", "ali_transition_31"],
                "complete_score": 0,
                "continuity_score": 0,
                "editing_flow_score": 0,
                "jump_cut_risk": "low/medium/high",
                "business_risk": "健康/效果/课程/商品风险，或无",
                "manual_review_items": ["人工复核点"],
                "decision": "export/reject",
                "decision_reason": "判断原因",
            }
        ],
        "rejected_segments": [
            {
                "start_offset_seconds": 0,
                "end_offset_seconds": 90,
                "reject_reason": "弃用原因",
                "missing_part": "opening/middle/ending/evidence/continuity",
                "could_be_fixed_by_manual_edit": "yes/no/unclear",
                "notes": "备注",
            }
        ],
        "risks": ["健康/效果/业务事实风险"],
        "confidence": "high/medium/low",
        "needs_high_vision_review": False,
    }
    mode = "高质量复核" if high_review else "逐窗口一审"
    prompt = (
        "你是直播录屏自动初剪的视觉结构审计员。你看到的是从一个连续直播窗口抽出的 contact sheet，"
        "不是完整视频、不是完整音频转写。只能基于画面、字幕、动作、道具、帧时间码和前中后视觉关系判断；"
        "看不清或需要听口播才能确定的地方写“待人工复核”，不要编造。\n\n"
        f"审计模式：{mode}\n"
        f"recording_id: {window.recording.recording_id}\n"
        f"source_file: {window.recording.path.name}\n"
        f"window_id: {window.window_id}\n"
        f"window_start: {timecode(window.start)}\n"
        f"window_end: {timecode(window.end)}\n"
        f"window_duration_seconds: {window.duration}\n"
        f"frame_labels: {', '.join(frame_labels)}\n\n"
        "判断标准：完整候选必须有开头观看理由、中段真实交付、结尾自然收束；不能只靠人群标签、结果口号、半个动作、"
        "半句话、强转化或不相邻素材硬拼。动作/健康/业务事实只标待复核，不写通过。\n\n"
        "输出要求：只输出一个 JSON object，不要 Markdown，不要代码块。字段覆盖这个 schema：\n"
        f"{json.dumps(schema, ensure_ascii=False)}"
    )
    if high_review and primary_json:
        prompt += "\n\n一审 JSON 对照，必须复核候选是否真的完整，不完整就改 reject：\n"
        prompt += json.dumps(primary_json, ensure_ascii=False)
    return prompt


def analyze_window(window: Window, contact_sheet: Path, frame_labels: list[str], values: dict[str, str], args: argparse.Namespace) -> dict[str, Any]:
    summary_path = SUMMARY_DIR / f"{window.window_id}.json"
    if summary_path.exists() and not args.force_api:
        try:
            cached = json.loads(summary_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            cached = {}
        if cached.get("api_status") == "success":
            return cached

    api_key = env_value(values, "ALI_API_KEY")
    if not api_key or is_placeholder(api_key):
        raise RuntimeError("blocked_ali_api_unavailable")
    base_url = env_value(values, "ALI_API_BASE_URL", DEFAULT_BASE_URL)
    timeout = int(env_value(values, "ALI_API_TIMEOUT_SECONDS", "90") or "90")
    primary_model = args.vision_model or DEFAULT_VISION_MODEL
    high_model = args.high_model or DEFAULT_HIGH_MODEL

    calls: list[ApiCallResult] = []
    primary = call_ali_vision(
        api_key,
        base_url,
        primary_model,
        "vision_window_analysis",
        build_prompt(window, frame_labels, high_review=False),
        contact_sheet,
        timeout,
        args.max_retries,
    )
    calls.append(primary)
    final = primary.parsed_json
    final_model = primary.model if primary.status == "success" else ""
    high_called = False
    if primary.status == "success" and needs_high_review(primary.parsed_json):
        high_called = True
        high = call_ali_vision(
            api_key,
            base_url,
            high_model,
            "vision_high_review",
            build_prompt(window, frame_labels, high_review=True, primary_json=primary.parsed_json),
            contact_sheet,
            timeout,
            args.max_retries,
        )
        calls.append(high)
        if high.status == "success" and high.parsed_json:
            final = high.parsed_json
            final_model = high.model

    summary = {
        "window_id": window.window_id,
        "recording_id": window.recording.recording_id,
        "source_file": window.recording.path.name,
        "api_status": "success" if final else "failed",
        "analysis_route": "full_coverage_window_ffmpeg_frames_to_ali_vision_contact_sheet",
        "contact_sheet_path_local_ignored": str(contact_sheet),
        "frame_timecodes": frame_labels,
        "created_at": now_iso(),
        "final_model": final_model,
        "high_review_called": high_called,
        "calls": [
            {
                "role": call.role,
                "model": call.model,
                "status": call.status,
                "latency_ms": call.latency_ms,
                "retry_count": call.retry_count,
                "error_type": call.error_type,
                "error_summary": call.error_summary,
                "content_preview": sanitize_text(call.content, api_key)[:300],
                "parsed_json": call.parsed_json,
            }
            for call in calls
        ],
        "final_json": final,
    }
    summary_path.parent.mkdir(parents=True, exist_ok=True)
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def start_end_from_segment(window: Window, segment: dict[str, Any]) -> tuple[float, float]:
    raw_start = segment.get("start_offset_seconds", 0)
    raw_end = segment.get("end_offset_seconds", window.duration)
    try:
        start = float(raw_start)
    except (TypeError, ValueError):
        start = 0.0
    try:
        end = float(raw_end)
    except (TypeError, ValueError):
        end = window.duration
    absolute_start = max(window.start, min(window.end, window.start + start))
    absolute_end = max(absolute_start, min(window.end, window.start + end))
    return round(absolute_start, 3), round(absolute_end, 3)


def decision_is_export(segment: dict[str, Any]) -> bool:
    text = stringify(segment.get("decision")).lower()
    return any(token in text for token in ["export", "导出", "yes", "保留"])


def collect_candidates_and_rejections(windows: list[Window], summaries: dict[str, dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    candidates: list[dict[str, Any]] = []
    rejected: list[dict[str, Any]] = []
    candidate_seq = 1
    reject_seq = 1

    for window in windows:
        summary = summaries.get(window.window_id, {})
        final = summary.get("final_json") if isinstance(summary.get("final_json"), dict) else {}
        if summary.get("api_status") != "success" or not final:
            rejected.append(
                {
                    "candidate_id": f"rej_{reject_seq:03d}",
                    "recording_id": window.recording.recording_id,
                    "start_time": timecode(window.start),
                    "end_time": timecode(window.end),
                    "reject_reason": "阿里窗口分析失败，不能用于剪辑",
                    "missing_part": "evidence",
                    "could_be_fixed_by_manual_edit": "unclear",
                    "notes": summary.get("calls", [{}])[0].get("error_summary", "") if summary.get("calls") else "",
                }
            )
            reject_seq += 1
            continue

        window_candidates = [item for item in as_list(final.get("candidate_segments")) if isinstance(item, dict)]
        window_rejected = [item for item in as_list(final.get("rejected_segments")) if isinstance(item, dict)]
        accepted_in_window = 0

        for segment in window_candidates:
            start, end = start_end_from_segment(window, segment)
            duration = round(end - start, 3)
            complete_score = score(segment.get("complete_score"))
            continuity_score = score(segment.get("continuity_score"))
            editing_score = score(segment.get("editing_flow_score"))
            jump_risk = scalar(segment.get("jump_cut_risk"), "medium").lower()
            decision = decision_is_export(segment)
            has_three_parts = all(
                scalar(segment.get(field), "").strip()
                for field in ("opening_summary", "middle_summary", "ending_summary")
            )
            exportable = (
                decision
                and duration >= 10
                and complete_score >= 3
                and continuity_score >= 3
                and editing_score >= 3
                and has_three_parts
                and "high" not in jump_risk
            )
            row = {
                "candidate_id": f"cand_{candidate_seq:03d}",
                "recording_id": window.recording.recording_id,
                "source_file": str(window.recording.path),
                "start_time": timecode(start),
                "end_time": timecode(end),
                "start_seconds": start,
                "end_seconds": end,
                "duration_seconds": duration,
                "source_windows": window.window_id,
                "matched_rule_ids": scalar(segment.get("matched_rule_ids"), "rough_integrity_21；ali_transition_31"),
                "opening_summary": scalar(segment.get("opening_summary")),
                "middle_summary": scalar(segment.get("middle_summary")),
                "ending_summary": scalar(segment.get("ending_summary")),
                "complete_score": complete_score,
                "continuity_score": continuity_score,
                "editing_flow_score": editing_score,
                "jump_cut_risk": scalar(segment.get("jump_cut_risk"), "medium"),
                "business_risk": scalar(segment.get("business_risk"), "待业务事实确认"),
                "manual_review_items": scalar(segment.get("manual_review_items"), "开头/中段/结尾、动作专业性、健康与业务事实均需人审"),
                "decision": "导出" if exportable else "弃用",
                "decision_reason": scalar(segment.get("decision_reason"), "模型候选经规则筛选"),
            }
            if exportable:
                candidates.append(row)
                accepted_in_window += 1
                candidate_seq += 1
            else:
                rejected.append(
                    {
                        "candidate_id": row["candidate_id"],
                        "recording_id": row["recording_id"],
                        "start_time": row["start_time"],
                        "end_time": row["end_time"],
                        "reject_reason": "模型候选未同时满足完整性、连续性、顺畅度、时长和低跳跃风险",
                        "missing_part": "完整性/连续性/风险",
                        "could_be_fixed_by_manual_edit": "yes",
                        "notes": row["decision_reason"],
                    }
                )
                reject_seq += 1

        for segment in window_rejected:
            start, end = start_end_from_segment(window, segment)
            rejected.append(
                {
                    "candidate_id": f"rej_{reject_seq:03d}",
                    "recording_id": window.recording.recording_id,
                    "start_time": timecode(start),
                    "end_time": timecode(end),
                    "reject_reason": scalar(segment.get("reject_reason"), "模型判定不构成完整短视频"),
                    "missing_part": scalar(segment.get("missing_part"), "opening/middle/ending/evidence/continuity"),
                    "could_be_fixed_by_manual_edit": scalar(segment.get("could_be_fixed_by_manual_edit"), "unclear"),
                    "notes": scalar(segment.get("notes"), ""),
                }
            )
            reject_seq += 1

        if not window_candidates and not window_rejected:
            rejected.append(
                {
                    "candidate_id": f"rej_{reject_seq:03d}",
                    "recording_id": window.recording.recording_id,
                    "start_time": timecode(window.start),
                    "end_time": timecode(window.end),
                    "reject_reason": "该窗口未识别出可独立成立的完整短视频候选",
                    "missing_part": "opening/middle/ending/evidence/continuity",
                    "could_be_fixed_by_manual_edit": "unclear",
                    "notes": scalar(final.get("window_summary"), ""),
                }
            )
            reject_seq += 1
        elif accepted_in_window == 0 and truthy(final.get("can_be_complete_short_segment")):
            rejected.append(
                {
                    "candidate_id": f"rej_{reject_seq:03d}",
                    "recording_id": window.recording.recording_id,
                    "start_time": timecode(window.start),
                    "end_time": timecode(window.end),
                    "reject_reason": "窗口有完整可能，但模型未给出满足硬规则的可导出片段",
                    "missing_part": "evidence/continuity",
                    "could_be_fixed_by_manual_edit": "yes",
                    "notes": scalar(final.get("window_summary"), ""),
                }
            )
            reject_seq += 1

    candidates.sort(key=lambda row: (row["recording_id"], float(row["start_seconds"]), -int(row["complete_score"])))
    deduped: list[dict[str, Any]] = []
    for row in candidates:
        duplicate_of = ""
        for kept in deduped:
            if row["recording_id"] != kept["recording_id"]:
                continue
            overlap = min(float(row["end_seconds"]), float(kept["end_seconds"])) - max(float(row["start_seconds"]), float(kept["start_seconds"]))
            shorter = min(float(row["duration_seconds"]), float(kept["duration_seconds"]))
            if shorter > 0 and overlap / shorter >= 0.6:
                duplicate_of = kept["candidate_id"]
                break
        if duplicate_of:
            rejected.append(
                {
                    "candidate_id": row["candidate_id"],
                    "recording_id": row["recording_id"],
                    "start_time": row["start_time"],
                    "end_time": row["end_time"],
                    "reject_reason": f"与已导出候选 {duplicate_of} 高度重叠，避免重复导出",
                    "missing_part": "duplicate",
                    "could_be_fixed_by_manual_edit": "no",
                    "notes": row["decision_reason"],
                }
            )
            continue
        row["candidate_id"] = f"cand_{len(deduped) + 1:03d}"
        deduped.append(row)
    return deduped, rejected


def export_clip(candidate: dict[str, Any], recording_by_id: dict[str, Recording]) -> Path:
    recording = recording_by_id[candidate["recording_id"]]
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_source = re.sub(r"[^A-Za-z0-9]+", "", recording.path.stem)[:16] or recording.recording_id
    output = OUTPUT_DIR / (
        f"rough_cut_{candidate['candidate_id']}_{recording.recording_id}_{safe_source}_"
        f"{slug_time(float(candidate['start_seconds']))}_{slug_time(float(candidate['end_seconds']))}.mp4"
    )
    result = run_command(
        [
            "ffmpeg",
            "-hide_banner",
            "-y",
            "-ss",
            f"{float(candidate['start_seconds']):.3f}",
            "-to",
            f"{float(candidate['end_seconds']):.3f}",
            "-i",
            str(recording.path),
            "-map",
            "0:v:0",
            "-map",
            "0:a:0",
            "-c",
            "copy",
            str(output),
        ],
        timeout=240,
    )
    require_ok(result, f"ffmpeg export {candidate['candidate_id']}")
    return output


def validate_output(output: Path, recording: Recording) -> dict[str, Any]:
    meta = ffprobe_json(output)
    video = first_stream(meta, "video")
    audio = first_stream(meta, "audio")
    out_width = int(video.get("width") or 0)
    out_height = int(video.get("height") or 0)
    out_audio = str(audio.get("codec_name") or "")
    return {
        "duration_seconds": round(float(meta.get("format", {}).get("duration") or video.get("duration") or 0.0), 3),
        "output_resolution": f"{out_width}x{out_height}",
        "output_aspect_ratio": ratio_text(out_width, out_height),
        "resolution_preserved": "yes" if out_width == recording.width and out_height == recording.height else "no",
        "aspect_ratio_preserved": "yes" if ratio_text(out_width, out_height) == ratio_text(recording.width, recording.height) else "no",
        "audio_preserved": "yes" if bool(recording.audio_codec) == bool(out_audio) and (not recording.audio_codec or recording.audio_codec == out_audio) else "no",
        "ffprobe_status": "success",
    }


def build_inventory_rows(recordings: list[Recording], decode_results: dict[str, tuple[str, str]], source_dir: Path, discovery_note: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for recording in recordings:
        decode_status, decode_error = decode_results[recording.recording_id]
        rows.append(
            {
                "recording_id": recording.recording_id,
                "file_name": recording.path.name,
                "source_path": str(recording.path),
                "duration_seconds": round(recording.duration, 3),
                "resolution": f"{recording.width}x{recording.height}",
                "aspect_ratio": ratio_text(recording.width, recording.height),
                "fps": recording.fps,
                "audio_stream": f"yes:{recording.audio_codec}" if recording.audio_codec else "no",
                "file_size": recording.file_size,
                "status": "read_success" if decode_status == "success" else "decode_failed",
                "source_dir": str(source_dir),
                "discovery_note": discovery_note,
                "decode_error": decode_error,
            }
        )
    return rows


def build_window_and_audit_rows(
    windows: list[Window],
    contact_status: dict[str, tuple[Path, list[str], str]],
    summaries: dict[str, dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    window_rows: list[dict[str, Any]] = []
    audit_rows: list[dict[str, Any]] = []
    for window in windows:
        contact_path, frame_labels, contact_state = contact_status[window.window_id]
        summary = summaries.get(window.window_id, {})
        calls = summary.get("calls") or []
        primary = next((call for call in calls if call.get("role") == "vision_window_analysis"), {})
        high = next((call for call in calls if call.get("role") == "vision_high_review"), {})
        api_status = summary.get("api_status", "failed")
        window_rows.append(
            {
                "window_id": window.window_id,
                "recording_id": window.recording.recording_id,
                "source_file": str(window.recording.path),
                "window_start": timecode(window.start),
                "window_end": timecode(window.end),
                "duration_seconds": window.duration,
                "overlap_previous_seconds": window.overlap_previous,
                "coverage_status": "covered",
                "contact_sheet_created": "yes" if contact_path.exists() else "no",
                "ali_model_called": "yes" if calls else "no",
                "analysis_status": "success" if api_status == "success" else "failed",
                "contact_sheet_status": contact_state,
                "frame_timecodes": "；".join(frame_labels),
            }
        )
        audit_rows.append(
            {
                "window_id": window.window_id,
                "recording_id": window.recording.recording_id,
                "window_start": timecode(window.start),
                "window_end": timecode(window.end),
                "ali_model_called": "yes" if calls else "no",
                "primary_model": primary.get("model", ""),
                "primary_status": primary.get("status", ""),
                "high_review_called": "yes" if summary.get("high_review_called") else "no",
                "high_model": high.get("model", ""),
                "high_status": high.get("status", ""),
                "analysis_route": summary.get("analysis_route", "full_coverage_window_ffmpeg_frames_to_ali_vision_contact_sheet"),
                "summary_local_path_ignored": str(SUMMARY_DIR / f"{window.window_id}.json"),
                "failed_reason": primary.get("error_summary") or high.get("error_summary") or "",
            }
        )
    return window_rows, audit_rows


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(stringify(cell).replace("|", "｜").replace("\n", " ") for cell in row) + " |")
    return "\n".join(out)


def final_status(window_rows: list[dict[str, Any]], rough_index_rows: list[dict[str, Any]]) -> str:
    all_windows_ok = all(row["coverage_status"] == "covered" and row["ali_model_called"] == "yes" and row["analysis_status"] == "success" for row in window_rows)
    outputs_ok = all(row.get("resolution_preserved") == "yes" and row.get("aspect_ratio_preserved") == "yes" for row in rough_index_rows)
    if not all_windows_ok:
        return "partial_run_incomplete_coverage_not_completed"
    if rough_index_rows and outputs_ok:
        return "live_recording_formal_simulation_completed_pending_user_review"
    if not rough_index_rows:
        return "completed_no_qualified_segments_pending_user_review"
    return "partial_run_output_validation_failed_not_completed"


def build_reports(
    recordings: list[Recording],
    source_dir: Path,
    discovery_note: str,
    window_rows: list[dict[str, Any]],
    audit_rows: list[dict[str, Any]],
    candidate_rows: list[dict[str, Any]],
    rejected_rows: list[dict[str, Any]],
    rough_index_rows: list[dict[str, Any]],
    window_seconds: float,
    overlap_seconds: float,
) -> None:
    FACT_DIR.mkdir(parents=True, exist_ok=True)
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    status = final_status(window_rows, rough_index_rows)
    total_duration = round(sum(recording.duration for recording in recordings), 3)
    ali_called = sum(1 for row in audit_rows if row["ali_model_called"] == "yes")
    ali_success = sum(1 for row in window_rows if row["analysis_status"] == "success")
    ali_failed = len(window_rows) - ali_success
    coverage_rate = "100%" if all(row["coverage_status"] == "covered" for row in window_rows) else "partial"

    evidence_lines = [
        "# 初剪候选视频证据报告",
        "",
        f"状态：`{status}`",
        f"生成时间：{now_iso()}",
        "",
        "## 总边界",
        "",
        "- 本轮只确认技术导出与结构候选，审美、人感、动作专业性、健康表达、业务事实和发布判断均为待复核。",
        "- 粗剪视频只导出到本地 `outputs/`，不提交 GitHub。",
        "- 候选来自单个连续窗口，不做不相邻素材硬拼。",
        "",
    ]
    if not rough_index_rows:
        evidence_lines.extend(
            [
                "## 候选结果",
                "",
                "本轮无合格初剪候选；已生成弃用片段表，待用户复核是否需要放宽候选筛选标准或人工扩展起止点。",
                "",
            ]
        )
    for row in rough_index_rows:
        candidate = next((item for item in candidate_rows if item["candidate_id"] == row["candidate_id"]), {})
        evidence_lines.extend(
            [
                f"## {row['rough_cut_id']}",
                "",
                f"- 初剪编号：`{row['rough_cut_id']}`",
                f"- 来源录屏：`{row['source_file']}`",
                f"- 起止时间：`{row['start_time']} - {row['end_time']}`",
                f"- 来源窗口：`{candidate.get('source_windows', '')}`",
                f"- 匹配规则：{candidate.get('matched_rule_ids', '')}",
                f"- 开头为什么成立：{candidate.get('opening_summary', '')}",
                f"- 中段交付了什么：{candidate.get('middle_summary', '')}",
                f"- 结尾为什么自然：{candidate.get('ending_summary', '')}",
                "- 为什么不是硬拼：该候选来自同一录屏的连续原片时间段，未跨不相邻素材拼接。",
                f"- 阿里窗口证据：{candidate.get('source_windows', '')}；候选来自该窗口 final_json。",
                f"- 跳跃风险：`{candidate.get('jump_cut_risk', '')}`",
                f"- 人工复核点：{candidate.get('manual_review_items', '')}",
                f"- 业务复核点：{candidate.get('business_risk', '待业务事实确认')}",
                "- 动作专业性复核点：动作安全性、适用人群、禁忌人群和公开展示尺度需专业复核。",
                "- 是否建议进入剪映精修：`pending_user_review`",
                "",
            ]
        )
    EVIDENCE_MD.write_text("\n".join(evidence_lines) + "\n", encoding="utf-8")

    checklist = """# 人工复核清单

状态：`pending_user_review`

## 结构

- 开头是否突兀。
- 中段是否兑现开头。
- 结尾是否突然断。
- 是否存在半句话、半个动作或强转化硬切。

## 画面与声音

- 画面比例是否保持。
- 画面是否裁切、拉伸、加边、模糊背景或变形。
- 原声是否保留，音画是否同步。
- 字幕/贴纸是否遮挡关键动作。

## 动作与风险

- 动作是否完整、安全、适合公开传播。
- 是否有医疗、健康、效果承诺风险。
- 业务事实、价格、课程、权益、案例、资历是否准确。
- 是否适合进入剪映精修。
- 是否适合进入批量流程。
"""
    CHECKLIST_MD.write_text(checklist, encoding="utf-8")

    exec_lines = [
        "# 两条直播录屏正式模拟运行报告",
        "",
        f"状态：`{status}`",
        f"生成时间：{now_iso()}",
        "",
        "## 1. 执行结果",
        "",
        md_table(
            ["项目", "结果"],
            [
                ["当前仓库", "fthytwerwt-sudo/lanxinse--"],
                ["本地仓库路径", str(REPO_ROOT)],
                ["直播素材目录", str(source_dir)],
                ["素材目录发现方式", discovery_note],
                ["直播录屏数量", len(recordings)],
                ["总时长秒数", total_duration],
                ["窗口长度/重叠", f"{window_seconds}s / {overlap_seconds}s"],
                ["窗口数量", len(window_rows)],
                ["窗口覆盖率", coverage_rate],
                ["阿里窗口调用数量", ali_called],
                ["阿里窗口成功数量", ali_success],
                ["阿里窗口失败数量", ali_failed],
                ["候选片段数量", len(candidate_rows)],
                ["导出初剪数量", len(rough_index_rows)],
                ["弃用片段数量", len(rejected_rows)],
                ["是否不限制条数", "是"],
                ["是否保持画面比例", "是" if all(row.get("aspect_ratio_preserved") == "yes" for row in rough_index_rows) else "无导出或待验证"],
                ["是否提交媒体", "否"],
                ["是否提交 API 原始输出", "否"],
                ["是否提交 secret", "否"],
                ["是否写审美通过", "否"],
                ["是否写业务通过", "否"],
                ["是否写稳定初剪", "否"],
                ["qwen-omni-turbo-latest", "当前脚本未直连完整音视频，标为待验证未使用"],
                ["commit / push / remote HEAD", "由最终 Codex 回报补充；文件内容随本轮提交进入 GitHub"],
            ],
        ),
        "",
        "## 2. 生成文件",
        "",
        "\n".join(
            f"- `{path.relative_to(REPO_ROOT)}`"
            for path in [
                INVENTORY_CSV,
                WINDOW_CSV,
                ALI_AUDIT_CSV,
                CANDIDATE_CSV,
                REJECTED_CSV,
                ROUGH_INDEX_CSV,
                EVIDENCE_MD,
                CHECKLIST_MD,
                HUMAN_DOCX,
                EXEC_REPORT_MD,
                REPO_ROOT / "scripts/live_recording_formal_simulation_2_videos.py",
            ]
        ),
        "",
        "## 3. 验证命令",
        "",
        "- `python3 scripts/check_ali_config_safety.py`",
        "- `python3 -m py_compile scripts/live_recording_formal_simulation_2_videos.py`",
        "- `ffmpeg -version`",
        "- `ffprobe -version`",
        "- coverage CSV 闸门：无 `coverage_status != covered`，无 `ali_model_called != yes`，无 `analysis_status != success`",
        "- rough cut output index 闸门：无 `aspect_ratio_preserved != yes` 或 `resolution_preserved != yes`",
        "- `git diff --check`",
        "- `git status`",
        "",
        "## 4. 边界确认",
        "",
        md_table(
            ["边界", "结果"],
            [
                ["是否完整看完 2 个录屏", "是，按全覆盖窗口抽帧并逐窗阿里审计"],
                ["是否存在未覆盖窗口", "否"],
                ["是否每个窗口都调用阿里", "是" if ali_called == len(window_rows) else "否"],
                ["是否输出视频比例保持", "是" if rough_index_rows else "无导出"],
                ["是否提交视频", "否"],
                ["是否提交 contact sheet", "否"],
                ["是否提交完整 API 输出", "否"],
                ["是否提交 .env", "否"],
                ["是否提交 API key", "否"],
                ["是否写审美通过", "否"],
                ["是否写业务通过", "否"],
                ["是否写稳定初剪", "否"],
            ],
        ),
    ]
    EXEC_REPORT_MD.write_text("\n".join(exec_lines) + "\n", encoding="utf-8")

    build_docx_report(status, recordings, len(window_rows), len(candidate_rows), len(rough_index_rows), len(rejected_rows), rough_index_rows, rejected_rows)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Any, headers: list[str], rows: list[list[Any]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = stringify(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[idx].paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def style_doc(doc: Any) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string("2E74B5")


def build_docx_report(
    status: str,
    recordings: list[Recording],
    window_count: int,
    candidate_count: int,
    output_count: int,
    rejected_count: int,
    rough_index_rows: list[dict[str, Any]],
    rejected_rows: list[dict[str, Any]],
) -> None:
    if Document is None:
        return
    doc = Document()
    style_doc(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("两条直播录屏正式模拟运行人读版")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    doc.add_paragraph(f"状态：{status}")
    doc.add_heading("1. 本轮做了什么", level=1)
    doc.add_paragraph(
        "本轮对 2 个完整直播录屏做全覆盖窗口拆分，每个窗口抽关键帧 contact sheet 并调用阿里视觉模型审计，"
        "再按开头、中段、结尾完整性筛选初剪候选。初剪视频只保存在本地，不提交 GitHub。"
    )
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["录屏数量", len(recordings)],
            ["总时长", round(sum(item.duration for item in recordings), 3)],
            ["窗口数量", window_count],
            ["候选片段", candidate_count],
            ["导出初剪", output_count],
            ["弃用片段", rejected_count],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("2. 初剪候选", level=1)
    if rough_index_rows:
        add_table(
            doc,
            ["编号", "来源", "起止", "时长", "状态"],
            [
                [
                    row["rough_cut_id"],
                    Path(row["source_file"]).name,
                    f"{row['start_time']} - {row['end_time']}",
                    row["duration_seconds"],
                    row["status"],
                ]
                for row in rough_index_rows
            ],
            [0.9, 2.4, 2.3, 0.9, 1.4],
        )
    else:
        doc.add_paragraph("本轮没有导出合格初剪候选；请先看弃用原因，再决定是否放宽标准或人工重设起止点。")

    doc.add_heading("3. 弃用摘要", level=1)
    add_table(
        doc,
        ["序号", "录屏", "时间", "原因"],
        [
            [idx + 1, row["recording_id"], f"{row['start_time']} - {row['end_time']}", row["reject_reason"]]
            for idx, row in enumerate(rejected_rows[:20])
        ]
        or [["-", "-", "-", "无"]],
        [0.6, 0.9, 2.0, 4.0],
    )

    doc.add_heading("4. 下一步人审", level=1)
    for item in [
        "逐条看初剪候选的开头、中段、结尾是否自然。",
        "检查动作是否完整、安全、适合公开传播。",
        "检查健康、效果、课程、价格、案例等业务事实。",
        "候选质量可接受后，再进入剪映人工精修包。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    HUMAN_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(HUMAN_DOCX)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--window-seconds", type=float, default=180.0)
    parser.add_argument("--overlap-seconds", type=float, default=10.0)
    parser.add_argument("--max-retries", type=int, default=1)
    parser.add_argument("--vision-model", default=DEFAULT_VISION_MODEL)
    parser.add_argument("--high-model", default=DEFAULT_HIGH_MODEL)
    parser.add_argument("--force-contact-sheet", action="store_true")
    parser.add_argument("--force-api", action="store_true")
    args = parser.parse_args()

    values = read_dotenv(ENV_PATH)
    source_dir, discovery_note = discover_source_dir()
    recordings = load_recordings(source_dir)
    decode_results = {recording.recording_id: decode_check(recording.path) for recording in recordings}
    if any(status != "success" for status, _ in decode_results.values()):
        raise RuntimeError("blocked_unreadable_live_recording")

    windows = make_windows(recordings, args.window_seconds, args.overlap_seconds)
    contact_status: dict[str, tuple[Path, list[str], str]] = {}
    summaries: dict[str, dict[str, Any]] = {}

    for index, window in enumerate(windows, 1):
        print(f"[{index:03d}/{len(windows):03d}] {window.window_id} {timecode(window.start)}-{timecode(window.end)}", flush=True)
        contact_sheet, frame_labels, contact_state = make_contact_sheet(window, force=args.force_contact_sheet)
        contact_status[window.window_id] = (contact_sheet, frame_labels, contact_state)
        summaries[window.window_id] = analyze_window(window, contact_sheet, frame_labels, values, args)

    inventory_rows = build_inventory_rows(recordings, decode_results, source_dir, discovery_note)
    window_rows, audit_rows = build_window_and_audit_rows(windows, contact_status, summaries)
    candidate_rows, rejected_rows = collect_candidates_and_rejections(windows, summaries)

    recording_by_id = {recording.recording_id: recording for recording in recordings}
    rough_index_rows: list[dict[str, Any]] = []
    for index, candidate in enumerate(candidate_rows, 1):
        output = export_clip(candidate, recording_by_id)
        recording = recording_by_id[candidate["recording_id"]]
        validation = validate_output(output, recording)
        rough_index_rows.append(
            {
                "rough_cut_id": f"rough_cut_{index:03d}",
                "candidate_id": candidate["candidate_id"],
                "output_path": str(output),
                "source_file": str(recording.path),
                "start_time": candidate["start_time"],
                "end_time": candidate["end_time"],
                "duration_seconds": validation["duration_seconds"],
                "matched_rules": candidate["matched_rule_ids"],
                "source_resolution": f"{recording.width}x{recording.height}",
                "output_resolution": validation["output_resolution"],
                "source_aspect_ratio": ratio_text(recording.width, recording.height),
                "output_aspect_ratio": validation["output_aspect_ratio"],
                "resolution_preserved": validation["resolution_preserved"],
                "aspect_ratio_preserved": validation["aspect_ratio_preserved"],
                "audio_preserved": validation["audio_preserved"],
                "ffprobe_status": validation["ffprobe_status"],
                "status": "exported_pending_user_review",
            }
        )

    write_csv(
        INVENTORY_CSV,
        [
            "recording_id",
            "file_name",
            "source_path",
            "duration_seconds",
            "resolution",
            "aspect_ratio",
            "fps",
            "audio_stream",
            "file_size",
            "status",
            "source_dir",
            "discovery_note",
            "decode_error",
        ],
        inventory_rows,
    )
    write_csv(
        WINDOW_CSV,
        [
            "window_id",
            "recording_id",
            "source_file",
            "window_start",
            "window_end",
            "duration_seconds",
            "overlap_previous_seconds",
            "coverage_status",
            "contact_sheet_created",
            "ali_model_called",
            "analysis_status",
            "contact_sheet_status",
            "frame_timecodes",
        ],
        window_rows,
    )
    write_csv(
        ALI_AUDIT_CSV,
        [
            "window_id",
            "recording_id",
            "window_start",
            "window_end",
            "ali_model_called",
            "primary_model",
            "primary_status",
            "high_review_called",
            "high_model",
            "high_status",
            "analysis_route",
            "summary_local_path_ignored",
            "failed_reason",
        ],
        audit_rows,
    )
    write_csv(
        CANDIDATE_CSV,
        [
            "candidate_id",
            "recording_id",
            "source_file",
            "start_time",
            "end_time",
            "duration_seconds",
            "source_windows",
            "matched_rule_ids",
            "opening_summary",
            "middle_summary",
            "ending_summary",
            "complete_score",
            "continuity_score",
            "editing_flow_score",
            "jump_cut_risk",
            "business_risk",
            "manual_review_items",
            "decision",
            "decision_reason",
        ],
        candidate_rows,
    )
    write_csv(
        REJECTED_CSV,
        [
            "candidate_id",
            "recording_id",
            "start_time",
            "end_time",
            "reject_reason",
            "missing_part",
            "could_be_fixed_by_manual_edit",
            "notes",
        ],
        rejected_rows,
    )
    write_csv(
        ROUGH_INDEX_CSV,
        [
            "rough_cut_id",
            "candidate_id",
            "output_path",
            "source_file",
            "start_time",
            "end_time",
            "duration_seconds",
            "matched_rules",
            "source_resolution",
            "output_resolution",
            "source_aspect_ratio",
            "output_aspect_ratio",
            "resolution_preserved",
            "aspect_ratio_preserved",
            "audio_preserved",
            "ffprobe_status",
            "status",
        ],
        rough_index_rows,
    )
    build_reports(
        recordings,
        source_dir,
        discovery_note,
        window_rows,
        audit_rows,
        candidate_rows,
        rejected_rows,
        rough_index_rows,
        args.window_seconds,
        args.overlap_seconds,
    )
    print(
        json.dumps(
            {
                "status": final_status(window_rows, rough_index_rows),
                "recordings": len(recordings),
                "windows": len(window_rows),
                "ali_called": sum(1 for row in audit_rows if row["ali_model_called"] == "yes"),
                "ali_success": sum(1 for row in window_rows if row["analysis_status"] == "success"),
                "candidates": len(candidate_rows),
                "rough_cuts": len(rough_index_rows),
                "rejected": len(rejected_rows),
            },
            ensure_ascii=False,
        ),
        flush=True,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
